#!/usr/bin/python3
# Name:         md2docx.py
# Version:      v06a Shimo-Gion
# Time-stamp:   <2023.04.21-18:05:51-JST>

# md2docx.py
# Copyright (C) 2022-2023  Seiichiro HATA
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <http://www.gnu.org/licenses/>.


# 2022.07.21 v01 Hiroshima
# 2022.08.24 v02 Shin-Hakushima
# 2022.12.25 v03 Yokogawa
# 2023.01.07 v04 Mitaki
# 2023.03.16 v05 Aki-Nagatsuka
# 20XX.XX.XX v06 Shimo-Gion


############################################################
# SETTING


import os
import sys
import argparse
import re
import unicodedata
import datetime
import docx
import chardet
from docx.shared import Cm, Pt
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import socket   # host
import getpass  # user


__version__ = 'v06a Shimo-Gion'


def get_arguments():
    parser = argparse.ArgumentParser(
        formatter_class=argparse.RawDescriptionHelpFormatter,
        description='MarkdownファイルからMS Wordファイルを作ります',
        add_help=False,
        epilog=HELP_EPILOG)
    parser.add_argument(
        '-h', '--help',
        action='help',
        help='ヘルプメッセージを表示します')
    parser.add_argument(
        '-v', '--version',
        action='version',
        version=('%(prog)s ' + __version__),
        help='バージョン番号を表示します')
    parser.add_argument(
        '-T', '--document-title',
        type=str,
        metavar='STRING',
        help='文書の標題')
    parser.add_argument(
        '-p', '--paper-size',
        type=str,
        choices=['A3', 'A3L', 'A3P', 'A4', 'A4L', 'A4P'],
        help='用紙設定（A3、A3L、A3P、A4、A4L、A4P）')
    parser.add_argument(
        '-t', '--top-margin',
        type=float,
        metavar='NUMBER',
        help='上余白（単位cm）')
    parser.add_argument(
        '-b', '--bottom-margin',
        type=float,
        metavar='NUMBER',
        help='下余白（単位cm）')
    parser.add_argument(
        '-l', '--left-margin',
        type=float,
        metavar='NUMBER',
        help='左余白（単位cm）')
    parser.add_argument(
        '-r', '--right-margin',
        type=float,
        metavar='NUMBER',
        help='右余白（単位cm）')
    parser.add_argument(
        '-d', '--document-style',
        type=str,
        choices=['k', 'j'],
        help='文書スタイルの指定（契約、条文）')
    parser.add_argument(
        '-H', '--header-string',
        type=str,
        help='ヘッダーの文字列')
    parser.add_argument(
        '-P', '--page-number',
        type=str,
        help='ページ番号の書式')
    parser.add_argument(
        '-L', '--line-number',
        action='store_true',
        help='行番号を出力します')
    parser.add_argument(
        '-m', '--mincho-font',
        type=str,
        metavar='FONT_NAME',
        help='明朝フォント')
    parser.add_argument(
        '-g', '--gothic-font',
        type=str,
        metavar='FONT_NAME',
        help='ゴシックフォント')
    parser.add_argument(
        '-i', '--ivs-font',
        type=str,
        metavar='FONT_NAME',
        help='異字体（IVS）フォント')
    parser.add_argument(
        '-f', '--font-size',
        type=float,
        metavar='NUMBER',
        help='フォントサイズ（単位pt）')
    parser.add_argument(
        '-s', '--line-spacing',
        type=float,
        metavar='NUMBER',
        help='行間の高さ（単位文字）')
    parser.add_argument(
        '-B', '--space-before',
        type=floats6,
        metavar='NUMBER,NUMBER,...',
        help='タイトル前の空白')
    parser.add_argument(
        '-A', '--space-after',
        type=floats6,
        metavar='NUMBER,NUMBER,...',
        help='タイトル後の空白')
    parser.add_argument(
        '-a', '--auto-space',
        action='store_true',
        help='全角文字と半角文字との間の間隔を微調整します')
    parser.add_argument(
        'md_file',
        help='Markdownファイル')
    parser.add_argument(
        'docx_file',
        default='',
        nargs='?',
        help='MS Wordファイル')
    return parser.parse_args()


def floats6(s):
    if not re.match('^' + RES_NUMBER6 + '$', s):
        raise argparse.ArgumentTypeError
    return s


HELP_EPILOG = '''Markdownの記法:
  行頭指示
    [#+=(数字) ]でセクション番号を変えることができます（独自）
    [v=(数字) ]で段落の上の余白を行数だけ増減します（独自）
    [V=(数字) ]で段落の下の余白を行数だけ増減します（独自）
    [X=(数字) ]で段落の改行幅を行数だけ増減します（独自）
    [<<=(数字) ]で段落1行目の左の余白を文字数だけ増減します（独自）
    [<=(数字) ]で段落の左の余白を文字数だけ増減します（独自）
    [>=(数字) ]で段落の右の余白を文字数だけ増減します（独自）
  行中指示
    [;;]から行末まではコメントアウトされます（独自）
    [<>]は何もせず表示もされません（独自）
    [<br>]で改行されます
    [<pgbr>]で改行されます（独自）
  文字装飾
    [*]で挟まれた文字列は斜体になります
    [**]で挟まれた文字列は太字になります
    [***]で挟まれた文字列は斜体かつ太字になります
    [~~]で挟まれた文字列は打消線が引かれます
    [`]で挟まれた文字列はゴシック体になります
    [//]で挟まれた文字列は斜体になります（独自）
    [--]で挟まれた文字列は文字が小さくなります（独自）
    [++]で挟まれた文字列は文字が大きくなります（独自）
    [^^]で挟まれた文字列は白色になって見えなくなります（独自）
    [^XXYYZZ^]で挟まれた文字列はRGB(XX,YY,ZZ)色になります（独自）
    [^foo^]で挟まれた文字列はfoo色になります（独自）
    [__]で挟まれた文字列は下線が引かれます（独自）
    [_foo_]で挟まれた区間の背景はfoo色になります（独自）
      red(R) darkRed(DR) yellow(Y) darkYellow(DY) green(G) darkGreen(DG)
      cyan(C) darkCyan(DC) blue(B) darkBlue(DB) magenta(M) darkMagenta(DM)
      lightGray(G1) darkGray(G2) black(BK)
    [字N;]（N=0-239）で"字"の異字体（IVS）が使えます（独自）
      ただし、IPAmj明朝フォント等がインストールされている必要があります。
  エスケープ記号
    [\\]をコマンドの前に書くとコマンドが文字列になります
    [\\\\]で"\\"が表示されます
'''

DEFAULT_DOCUMENT_TITLE = ''

DEFAULT_DOCUMENT_STYLE = 'n'

DEFAULT_PAPER_SIZE = 'A4'
PAPER_HEIGHT = {'A3': 29.7, 'A3L': 29.7, 'A3P': 42.0,
                'A4': 29.7, 'A4L': 21.0, 'A4P': 29.7}
PAPER_WIDTH = {'A3': 42.0, 'A3L': 42.0, 'A3P': 29.7,
               'A4': 21.0, 'A4L': 29.7, 'A4P': 21.0}

DEFAULT_TOP_MARGIN = 3.5
DEFAULT_BOTTOM_MARGIN = 2.2
DEFAULT_LEFT_MARGIN = 3.0
DEFAULT_RIGHT_MARGIN = 2.0

DEFAULT_HEADER_STRING = ''

DEFAULT_PAGE_NUMBER = 'n'

DEFAULT_LINE_NUMBER = False

DEFAULT_MINCHO_FONT = 'ＭＳ 明朝'
DEFAULT_GOTHIC_FONT = 'ＭＳ ゴシック'
DEFAULT_IVS_FONT = 'IPAmj明朝'  # IPAmjMincho
DEFAULT_FONT_SIZE = 12.0

DEFAULT_LINE_SPACING = 2.14  # (2.0980+2.1812)/2=2.1396

DEFAULT_SPACE_BEFORE = ''
DEFAULT_SPACE_AFTER = ''
TABLE_SPACE_BEFORE = 0.45
TABLE_SPACE_AFTER = 0.2

DEFAULT_AUTO_SPACE = False

NOT_ESCAPED = '^((?:(?:.*\n)*.*[^\\\\])?(?:\\\\\\\\)*)?'

RES_NUMBER = '(?:[-\\+]?(?:(?:[0-9]+(?:\\.[0-9]+)?)|(?:\\.[0-9]+)))'
RES_NUMBER6 = '(?:' + RES_NUMBER + '?,){,5}' + RES_NUMBER + '?,?'

RES_IMAGE = '! *\\[([^\\[\\]]*)\\] *\\(([^\\(\\)]+)\\)'

FONT_DECORATORS = [
    '\\*\\*\\*',           # italic and bold
    '\\*\\*',              # bold
    '\\*',                 # italic
    '~~',                  # strikethrough
    '`',                   # preformatted
    '//',                  # italic
    '__',                  # underline
    '\\-\\-',              # small
    '\\+\\+',              # large
    '\\^[0-9A-Za-z]*\\^',  # font color
    '_[0-9A-Za-z]+_',      # higilight color
]

RELAX_SYMBOL = '<>'

ORIGINAL_COMMENT_SYMBOL = ';;'

COMMENT_SEPARATE_SYMBOL = ' / '

HORIZONTAL_BAR = '[ー−—－―‐]'

FONT_COLOR = {
    'red':         'FF0000',
    'R':           'FF0000',
    'darkRed':     '7F0000',
    'DR':          '7F0000',
    'yellow':      'FFFF00',
    'Y':           'FFFF00',
    'darkYellow':  '7F7F00',
    'DY':          '7F7F00',
    'green':       '00FF00',
    'G':           '00FF00',
    'darkGreen':   '007F00',
    'DG':          '007F00',
    'cyan':        '00FFFF',
    'C':           '00FFFF',
    'darkCyan':    '007F7F',
    'DC':          '007F7F',
    'blue':        '0000FF',
    'B':           '0000FF',
    'darkBlue':    '00007F',
    'DB':          '00007F',
    'magenta':     'FF00FF',
    'M':           'FF00FF',
    'darkMagenta': '7F007F',
    'DM':          '7F007F',
    'lightGray':   'BFBFBF',
    'G1':          'BFBFBF',
    'darkGray':    '7F7F7F',
    'G2':          '7F7F7F',
    'black':       '000000',
    'BK':          '000000',
    'a000': 'FF5D5D',
    'a010': 'FF603C',
    'a020': 'FF6512',
    'a030': 'E07000',
    'a040': 'BC7A00',
    'a050': 'A08300',
    'a060': '898900',
    'a070': '758F00',
    'a080': '619500',
    'a090': '4E9B00',
    'a100': '38A200',
    'a110': '1FA900',
    'a120': '00B200',
    'a130': '00AF20',
    'a140': '00AC3C',
    'a150': '00AA55',
    'a160': '00A76D',
    'a170': '00A586',
    'a180': '00A2A2',
    'a190': '009FC3',
    'a200': '009AED',
    'a210': '1F8FFF',
    'a220': '4385FF',
    'a230': '5F7CFF',
    'a240': '7676FF',
    'a250': '8A70FF',
    'a260': '9E6AFF',
    'a270': 'B164FF',
    'a280': 'C75DFF',
    'a290': 'E056FF',
    'a300': 'FF4DFF',
    'a310': 'FF50DF',
    'a320': 'FF53C3',
    'a330': 'FF55AA',
    'a340': 'FF5892',
    'a350': 'FF5A79',
}

HIGHLIGHT_COLOR = {
    'red':         WD_COLOR_INDEX.RED,
    'R':           WD_COLOR_INDEX.RED,
    'darkRed':     WD_COLOR_INDEX.DARK_RED,
    'DR':          WD_COLOR_INDEX.DARK_RED,
    'yellow':      WD_COLOR_INDEX.YELLOW,
    'Y':           WD_COLOR_INDEX.YELLOW,
    'darkYellow':  WD_COLOR_INDEX.DARK_YELLOW,
    'DY':          WD_COLOR_INDEX.DARK_YELLOW,
    'green':       WD_COLOR_INDEX.BRIGHT_GREEN,
    'G':           WD_COLOR_INDEX.BRIGHT_GREEN,
    'darkGreen':   WD_COLOR_INDEX.GREEN,
    'DG':          WD_COLOR_INDEX.GREEN,
    'cyan':        WD_COLOR_INDEX.TURQUOISE,
    'C':           WD_COLOR_INDEX.TURQUOISE,
    'darkCyan':    WD_COLOR_INDEX.TEAL,
    'DC':          WD_COLOR_INDEX.TEAL,
    'blue':        WD_COLOR_INDEX.BLUE,
    'B':           WD_COLOR_INDEX.BLUE,
    'darkBlue':    WD_COLOR_INDEX.DARK_BLUE,
    'DB':          WD_COLOR_INDEX.DARK_BLUE,
    'magenta':     WD_COLOR_INDEX.PINK,
    'M':           WD_COLOR_INDEX.PINK,
    'darkMagenta': WD_COLOR_INDEX.VIOLET,
    'DM':          WD_COLOR_INDEX.VIOLET,
    'lightGray':   WD_COLOR_INDEX.GRAY_25,
    'G1':          WD_COLOR_INDEX.GRAY_25,
    'darkGray':    WD_COLOR_INDEX.GRAY_50,
    'G2':          WD_COLOR_INDEX.GRAY_50,
    'black':       WD_COLOR_INDEX.BLACK,
    'BK':          WD_COLOR_INDEX.BLACK,
}


############################################################
# FUNCTION


def get_real_width(s):
    p = ''
    wid = 0.0
    for c in s:
        if (c == '\t'):
            wid = (wid + 8) // 8 * 8
            continue
        w = unicodedata.east_asian_width(c)
        if c == '':
            wid += 0.0
        elif re.match('^[−☐☑]$', c):
            wid += 2.0
        elif re.match('^[´¨―‐∥…‥‘’“”±×÷≠≦≧∞∴♂♀°′″℃§]$', c):
            wid += 2.0
        elif re.match('^[☆★○●◎◇◆□■△▲▽▼※→←↑↓]$', c):
            wid += 2.0
        elif re.match('^[∈∋⊆⊇⊂⊃∪∩∧∨⇒⇔∀∃∠⊥⌒∂∇≡≒≪≫√∽∝∵]$', c):
            wid += 2.0
        elif re.match('^[∫∬Å‰♯♭♪†‡¶◯]$', c):
            wid += 2.0
        elif re.match('^[ΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ]$', c):
            wid += 2.0
        elif re.match('^[αβγδεζηθικλμνξοπρστυφχψω]$', c):
            wid += 2.0
        elif re.match('^[АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ]$', c):
            wid += 2.0
        elif re.match('^[абвгдеёжзийклмнопрстуфхцчшщъыьэюя]$', c):
            wid += 2.0
        elif re.match('^[─│┌┐┘└├┬┤┴┼━┃┏┓┛┗┣┳┫┻╋┠┯┨┷┿┝┰┥┸╂]$', c):
            wid += 2.0
        elif re.match('^[№℡≒≡∫∮∑√⊥∠∟⊿∵∩∪]$', c):
            wid += 2.0
        elif re.match('^[⑴⑵⑶⑷⑸⑹⑺⑻⑼⑽⑾⑿⒀⒁⒂⒃⒄⒅⒆⒇]$', c):
            wid += 2.0
        elif re.match('^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]$', c):
            wid += 2.0
        elif re.match('^[⒈⒉⒊⒋⒌⒍⒎⒏⒐⒑⒒⒓⒔⒕⒖⒗⒘⒙⒚⒛]$', c):
            wid += 2.0
        elif re.match('^[ⅰⅱⅲⅳⅴⅵⅶⅷⅸⅹⅺⅻ]$', c):
            wid += 2.0
        elif re.match('^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫ]$', c):
            wid += 2.0
        elif re.match('^[⒜⒝⒞⒟⒠⒡⒢⒣⒤⒥⒦⒧⒨⒩⒪⒫⒬⒭⒮⒯⒰⒱⒲⒳⒴⒵]$', c):
            wid += 2.0
        elif re.match('^[ⓐⓑⓒⓓⓔⓕⓖⓗⓘⓙⓚⓛⓜⓝⓞⓟⓠⓡⓢⓣⓤⓥⓦⓧⓨⓩ]$', c):
            wid += 2.0
        elif re.match('^[🄐🄑🄒🄓🄔🄕🄖🄗🄘🄙🄚🄛🄜🄝🄞🄟🄠🄡🄢🄣🄤🄥🄦🄧🄨🄩]$', c):
            wid += 2.0
        elif re.match('^[ⒶⒷⒸⒹⒺⒻⒼⒽⒾⒿⓀⓁⓂⓃⓄⓅⓆⓇⓈⓉⓊⓋⓌⓍⓎⓏ]$', c):
            wid += 2.0
        elif re.match('^[㉑㉒㉓㉔㉕㉖㉗㉘㉙㉚㉛㉜㉝㉞㉟㊱㊲㊳㊴㊵㊶㊷㊸㊹㊺㊻㊼㊽㊾㊿]$', c):
            wid += 2.0
        elif re.match('^[🄋➀➁➂➃➄➅➆➇➈➉]$', c):
            wid += 2.0
        elif re.match('^[㋐㋑㋒㋓㋔㋕㋖㋗㋘㋙㋚㋛㋜㋝㋞㋟㋠㋡㋢㋣㋤㋥㋦㋧㋨]$', c):
            wid += 2.0
        elif re.match('^[㋩㋪㋫㋬㋭㋮㋯㋰㋱㋲㋳㋴㋵㋶㋷㋸㋹㋺㋻㋼㋽㋾]$', c):
            wid += 2.0
        elif re.match('^㊀㊁㊂㊃㊄㊅㊆㊇㊈㊉$', c):
            wid += 2.0
        elif (w == 'F'):  # Full alphabet ...
            wid += 2.0
        elif(w == 'H'):   # Half katakana ...
            wid += 1.0
        elif(w == 'W'):   # Chinese character ...
            wid += 2.0
        elif(w == 'Na'):  # Half alphabet ...
            wid += 1.0
        elif(w == 'A'):   # Greek character ...
            wid += 1.0
        elif(w == 'N'):   # Arabic character ...
            wid += 1.0
        if p != '' and p != w:
            wid += 0.5
        p = w
    return wid


def i2c_n_arab(n, md_line=None):
    if n < 0:
        return '△' + str(-n)
    elif n < 10:
        return chr(65296 + n)
    else:
        return str(n)


def i2c_p_arab(n, md_line=None):
    if n < 0:
        return '△(' + str(-n) + ')'
    elif n == 0:
        return '(0)'
    elif n > 0 and n <= 20:
        return chr(9331 + n)
    else:
        return '(' + str(n) + ')'


def i2c_c_arab(n, md_line=None):
    if n == 0:
        return chr(9450)
    elif n > 0 and n <= 20:
        return chr(9311 + n)
    elif n > 0 and n <= 35:
        return chr(12860 + n)
    elif n > 0 and n <= 50:
        return chr(12941 + n)
    else:
        msg = '※ 警告: ' \
            + '丸付き数字番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed circled number'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def i2c_n_kata(n, md_line=None):
    if n == 0:
        return chr(12448 + 83)
    elif n > 0 and n <= 5:
        return chr(12448 + (2 * n))
    elif n > 0 and n <= 17:
        return chr(12448 + (2 * n) - 1)
    elif n > 0 and n <= 20:
        return chr(12448 + (2 * n))
    elif n > 0 and n <= 25:
        return chr(12448 + (1 * n) + 21)
    elif n > 0 and n <= 30:
        return chr(12448 + (3 * n) - 31)
    elif n > 0 and n <= 35:
        return chr(12448 + (1 * n) + 31)
    elif n > 0 and n <= 38:
        return chr(12448 + (2 * n) - 4)
    elif n > 0 and n <= 43:
        return chr(12448 + (1 * n) + 34)
    elif n > 0 and n <= 45:
        return chr(12448 + (3 * n) - 53)
    elif n > 0 and n <= 46:
        return chr(12448 + (1 * n) + 37)
    else:
        msg = '※ 警告: ' \
            + 'カタカナ番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed katakana'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def i2c_p_kata(n, md_line=None):
    if n == 0:
        return '(' + chr(65392 + 45) + ')'
    elif n > 0 and n <= 44:
        return '(' + chr(65392 + n) + ')'
    elif n > 0 and n <= 45:
        return '(' + chr(65392 + n - 55) + ')'
    elif n > 0 and n <= 46:
        return '(' + chr(65392 + n - 1) + ')'
    else:
        msg = '※ 警告: ' \
            + '括弧付きカタカナ番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed parenthesis katakata'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def i2c_c_kata(n, md_line=None):
    if n == 0:
        return chr(13007 + 47)
    elif n > 0 and n <= 47:
        return chr(13007 + n)
    else:
        msg = '※ 警告: ' \
            + '丸付きカタカナ番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed circled katakana'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def i2c_n_alph(n, md_line=None):
    if n == 0:
        return chr(65344 + 26)
    elif n > 0 and n <= 26:
        return chr(65344 + n)
    else:
        msg = '※ 警告: ' \
            + 'アルファベット番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed alphabet'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def i2c_p_alph(n, md_line=None):
    if n == 0:
        return chr(9371 + 26)
    elif n > 0 and n <= 26:
        return chr(9371 + n)
    else:
        msg = '※ 警告: ' \
            + '括弧付きアルファベット番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed parenthesis alphabet'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def i2c_c_alph(n, md_line=None):
    if n == 0:
        return chr(9423 + 26)
    elif n > 0 and n <= 26:
        return chr(9423 + n)
    else:
        msg = '※ 警告: ' \
            + '丸付きアルファベット番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed circled alphabet'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def i2c_n_kanj(n, md_line=None):
    if n >= 0:
        k = str(n)
        if n >= 1000:
            k = re.sub('^(.+)(...)$', '\\1千\\2', k)
        if n >= 100:
            k = re.sub('^(.+)(..)$', '\\1百\\2', k)
        if n >= 10:
            k = re.sub('^(.+)(.)$', '\\1十\\2', k)
        k = re.sub('0', '〇', k)
        k = re.sub('1', '一', k)
        k = re.sub('2', '二', k)
        k = re.sub('3', '三', k)
        k = re.sub('4', '四', k)
        k = re.sub('5', '五', k)
        k = re.sub('6', '六', k)
        k = re.sub('7', '七', k)
        k = re.sub('8', '八', k)
        k = re.sub('9', '九', k)
        k = re.sub('〇$', '', k)
        k = re.sub('〇十', '', k)
        k = re.sub('〇百', '', k)
        k = re.sub('〇千', '', k)
        k = re.sub('一十', '十', k)
        k = re.sub('一百', '百', k)
        k = re.sub('一千', '千', k)
        return k
    else:
        msg = '※ 警告: ' \
            + '漢数字番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed kansuji'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def i2c_p_kanj(n, md_line=None):
    if n == 0:
        return chr(12831 + 10)
    elif n > 0 and n <= 10:
        return chr(12831 + n)
    else:
        msg = '※ 警告: ' \
            + '括弧付き漢数字番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed parenthesis kansuji'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def i2c_c_kanj(n, md_line=None):
    if n == 0:
        return chr(12927 + 10)
    elif n > 0 and n <= 10:
        return chr(12927 + n)
    else:
        msg = '※ 警告: ' \
            + '丸付き漢数字番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed circled kansuji'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def concatenate_string(str1, str2):
    res = '[0-9A-Za-z,\\.\\)}\\]]'
    if re.match('^.*' + res + '$', str1) and re.match('^' + res + '.*$', str2):
        return str1 + ' ' + str2
    else:
        return str1 + str2


############################################################
# CLASS


class Document:

    """A class to handle document"""

    document_title = DEFAULT_DOCUMENT_TITLE
    document_style = DEFAULT_DOCUMENT_STYLE
    paper_size = DEFAULT_PAPER_SIZE
    top_margin = DEFAULT_TOP_MARGIN
    bottom_margin = DEFAULT_BOTTOM_MARGIN
    left_margin = DEFAULT_LEFT_MARGIN
    right_margin = DEFAULT_RIGHT_MARGIN
    header_string = DEFAULT_HEADER_STRING
    page_number = DEFAULT_PAGE_NUMBER
    line_number = DEFAULT_LINE_NUMBER
    mincho_font = DEFAULT_MINCHO_FONT
    gothic_font = DEFAULT_GOTHIC_FONT
    ivs_font = DEFAULT_IVS_FONT
    font_size = DEFAULT_FONT_SIZE
    line_spacing = DEFAULT_LINE_SPACING
    space_before = DEFAULT_SPACE_BEFORE
    space_after = DEFAULT_SPACE_AFTER
    auto_space = DEFAULT_AUTO_SPACE
    original_file = ''

    def __init__(self):
        self.md_file = ''
        self.docx_file = ''
        self.raw_md_lines = []
        self.md_lines = []
        self.all_paragraphs = []
        self.paragraphs = []

    def get_raw_md_lines(self, md_file):
        self.md_file = md_file
        raw_md_lines = []
        try:
            if md_file == '-':
                bd = sys.stdin.buffer.read()
            else:
                bd = open(md_file, 'rb').read()
        except BaseException:
            msg = '※ エラー: ' \
                + '入力ファイル「' + md_file + '」を読み込めません'
            # msg = 'error: ' \
            #     + 'file "' + md_file + '" is not found'
            sys.stderr.write(msg + '\n\n')
            sys.exit(0)
        enc = chardet.detect(bd)['encoding']
        if enc is None:
            enc = 'SHIFT_JIS'
        elif (re.match('^utf[-_]?.*$', enc, re.I)) or \
             (re.match('^shift[-_]?jis.*$', enc, re.I)) or \
             (re.match('^cp932.*$', enc, re.I)) or \
             (re.match('^euc[-_]?(jp|jis).*$', enc, re.I)) or \
             (re.match('^iso[-_]?2022[-_]?jp.*$', enc, re.I)) or \
             (re.match('^ascii.*$', enc, re.I)):
            pass
        else:
            # Windows-1252 (Western Europe)
            # MacCyrillic (Macintosh Cyrillic)
            # ...
            msg = '※ 警告: ' \
                + '文字コード「' + enc + '」から「SHIFT_JIS」に修正しました'
            # msg = 'warning: ' \
            #     + 'detected encoding "' + enc + '" may be wrong'
            sys.stderr.write(msg + '\n\n')
            enc = 'SHIFT_JIS'
        try:
            sd = bd.decode(enc)
        except BaseException:
            msg = '※ エラー: ' \
                + '正しい文字コードを取得できませんでした' \
                + '（Markdownでない？）'
            # msg = 'error: ' \
            #     + 'could not detect correct character code '
            #     + '(not markdown?)'
            sys.stderr.write(msg + '\n\n')
            sys.exit(0)
        sd = re.sub('^' + chr(65279), '', sd)  # remove BOM / unnecessary?
        sd = re.sub('\r\n', '\n', sd)  # unnecessary?
        sd = re.sub('\r', '\n', sd)  # unnecessary?
        for rml in sd.split('\n'):
            rml = re.sub('  $', '\n', rml)
            rml = re.sub('[ \t\u3000]*$', '', rml)
            raw_md_lines.append(rml)
        raw_md_lines.append('')
        # self.raw_md_lines = raw_md_lines
        return raw_md_lines

    def get_md_lines(self, raw_md_lines):
        md_lines = []
        for i, rml in enumerate(raw_md_lines):
            ml = MdLine(i + 1, rml)
            md_lines.append(ml)
        # self.md_lines = md_lines
        return md_lines

    def get_raw_paragraphs(self, md_lines):
        raw_paragraphs = []
        block = []
        for ml in md_lines:
            is_block_end = False
            if ml.raw_text == '':
                is_block_end = True
            pre_text = ''
            if len(block) > 0:
                pre_text = block[-1].raw_text
                cur_text = ml.raw_text
                for pc in [ParagraphChapter, ParagraphSection, ParagraphList]:
                    res_s = '^\\s*' + pc.res_symbol + '\\s+\\S+.*$'
                    res_r = '^\\s*' + pc.res_reviser + '(\\s.*)?$'
                    if re.match(res_s + '|' + res_r, pre_text):
                        if re.match(res_s + '|' + res_r, cur_text):
                            is_block_end = True
            if is_block_end:
                if len(block) == 0:
                    if ml.raw_text != '':
                        block.append(ml)
                    continue
                if re.match('^```.*$', block[0].raw_text):
                    if len(block) == 1:
                        block.append(ml)
                        continue
                    elif not re.match('^.*```$', block[-1].raw_text):
                        block.append(ml)
                        continue
                rp = RawParagraph(block)
                raw_paragraphs.append(rp)
                block = []
            if ml.raw_text != '':
                block.append(ml)
        if len(block) > 0:
            rp = RawParagraph(block)
            raw_paragraphs.append(rp)
            block = []
        # self.raw_paragraphs = raw_paragraphs
        return raw_paragraphs

    def get_paragraphs(self, raw_paragraphs):
        paragraphs = []
        cr = []
        sr = []
        lr = []
        er = []
        hr = []
        sd = []
        res_v = '^v=(' + RES_NUMBER + ')$'
        res_cv = '^V=(' + RES_NUMBER + ')$'
        for rp in raw_paragraphs:
            full_text = rp.full_text
            if rp.paragraph_class == 'empty' or rp.paragraph_class == 'blank':
                cr += rp.chapter_revisers
                sr += rp.section_revisers
                lr += rp.list_revisers
                er += rp.length_revisers
                hr += rp.head_font_revisers + rp.tail_font_revisers
                sd += rp.section_depth_setters
                if rp.paragraph_class == 'blank':
                    nl = full_text.count('\n')
                    er += ['v=' + str(nl)]
            else:
                rp.chapter_revisers = cr + rp.chapter_revisers
                rp.section_revisers = sr + rp.section_revisers
                rp.list_revisers = lr + rp.list_revisers
                for rev in er:
                    if re.match(res_v, rev):
                        rp.length_revisers = [rev] + rp.length_revisers
                    if re.match(res_cv, rev):
                        rev = re.sub('^V=', 'v=', rev)
                        rp.length_revisers = [rev] + rp.length_revisers
                rp.head_font_revisers = hr + rp.head_font_revisers
                rp.section_depth_setters = sd + rp.section_depth_setters
                cr = []
                sr = []
                lr = []
                er = []
                hr = []
                sd = []
                p = rp.get_paragraph()
                paragraphs.append(p)
        # self.paragraphs = paragraphs
        return paragraphs

    def modify_paragraphs(self, paragraphs):
        for i, p in enumerate(paragraphs):
            if i > 0:
                p_prev = paragraphs[i - 1]
            # ARTICLE TITLE (MIMI=EAR)
            if Document.document_style == 'j' and \
               p.paragraph_class == 'section' and \
               p.head_section_depth == 2 and \
               p.tail_section_depth == 2 and \
               i > 0 and \
               p_prev.paragraph_class == 'alignment' and \
               p_prev.alignment == 'left':
                p_prev.length_docx['space before'] \
                    += p.length_conf['space before']
                p.length_docx['space before'] \
                    -= p.length_conf['space before']
        m = len(paragraphs) - 1
        for i, p in enumerate(paragraphs):
            if i > 0:
                p_prev = paragraphs[i - 1]
            if i < m:
                p_next = paragraphs[i + 1]
            # SECTION DEPTH 1
            if p.paragraph_class == 'section' and \
               ParagraphSection._get_section_depths(p.full_text) == (1, 1):
                # BEFORE
                if i > 0:
                    if p_prev.length_docx['space after'] >= 0.1:
                        p_prev.length_docx['space after'] += 0.1
                    elif p_prev.length_docx['space after'] >= 0.0:
                        p_prev.length_docx['space after'] *= 2
                if True:
                    if p.length_docx['space before'] >= 0.1:
                        p.length_docx['space before'] += 0.1
                    elif p.length_docx['space before'] >= 0.0:
                        p.length_docx['space before'] *= 2
                # AFTER
                if True:
                    if p.length_docx['space after'] >= 0.2:
                        p.length_docx['space after'] -= 0.1
                    elif p.length_docx['space after'] >= 0.0:
                        p.length_docx['space after'] /= 2
                if i < m:
                    if p_next.length_docx['space before'] >= 0.2:
                        p_next.length_docx['space before'] -= 0.1
                    elif p_next.length_docx['space before'] >= 0.0:
                        p_next.length_docx['space before'] /= 2
            # TABLE
            if p.paragraph_class == 'table':
                if i > 0:
                    sb = p.length_docx['space before']
                    if sb < 0:
                        msg = '警告: ' \
                            + '段落前の余白「v」の値が小さ過ぎます'
                        # msg = 'warning: ' \
                        #     + '"space before" is too small'
                        p.md_lines[0].append_warning_message(msg)
                    if p_prev.length_docx['space after'] \
                       < sb + TABLE_SPACE_BEFORE:
                        p_prev.length_docx['space after'] = sb
                    p.length_docx['space before'] = 0.0
                if i < m:
                    sa = p.length_docx['space after']
                    if sa < 0:
                        msg = '警告: ' \
                            + '段落前の余白「V」の値が小さ過ぎます'
                        # msg = 'warning: ' \
                        #     + '"space after" is too small'
                        p.md_lines[0].append_warning_message(msg)
                    if p_next.length_docx['space before'] \
                       < sa + TABLE_SPACE_AFTER:
                        p_next.length_docx['space before'] = sa
                    p.length_docx['space after'] = 0.0
        return self.paragraphs

    def configure(self, md_lines, args):
        self._configure_by_md_file(md_lines)
        self._configure_by_args(args)
        Paragraph.mincho_font = Document.mincho_font
        Paragraph.gothic_font = Document.gothic_font
        Paragraph.ivs_font = Document.ivs_font
        Paragraph.font_size = Document.font_size

    def _configure_by_md_file(self, md_lines):
        for ml in md_lines:
            com = ml.comment
            if com is None:
                break
            if re.match('^\\s*#', com):
                continue
            res = '^\\s*([^:：]+)[:：]\\s*(.*)$'
            if not re.match(res, com):
                continue
            nam = re.sub(res, '\\1', com).rstrip()
            val = re.sub(res, '\\2', com).rstrip()
            if False:
                pass
            elif nam == 'document_title' or nam == '書題名':
                Document.document_title = val
            elif nam == 'document_style' or nam == '文書式':
                if val == 'n' or val == '普通' or val == '-':
                    Document.document_style = 'n'
                elif val == 'k' or val == '契約':
                    Document.document_style = 'k'
                elif val == 'j' or val == '条文':
                    Document.document_style = 'j'
                else:
                    msg = '※ 警告: ' \
                        + '「' + nam + '」の値は"普通"、"契約"又は"条文"で' \
                        + 'なければなりません'
                    # msg = 'warning: ' \
                    #     + '"' + nam + '" must be "-", "k" or "j"'
                    sys.stderr.write(msg + '\n\n')
            elif nam == 'paper_size' or nam == '用紙サ':
                val = unicodedata.normalize('NFKC', val)
                if val == 'A3':
                    Document.paper_size = 'A3'
                elif val == 'A3L' or val == 'A3横':
                    Document.paper_size = 'A3L'
                elif val == 'A3P' or val == 'A3縦':
                    Document.paper_size = 'A3P'
                elif val == 'A4':
                    Document.paper_size = 'A4'
                elif val == 'A4L' or val == 'A4横':
                    Document.paper_size = 'A4L'
                elif val == 'A4P' or val == 'A4縦':
                    Document.paper_size = 'A4P'
                else:
                    msg = '※ 警告: ' \
                        + '「' + nam + '」の値は' \
                        + '"A3横"、"A3縦"、"A4横"又は"A4縦"で' \
                        + 'なければなりません'
                    # msg = 'warning: ' \
                    #     + '"' + nam + '" must be "A3", "A3P", "A4" or "A4L"'
                    sys.stderr.write(msg + '\n\n')
            elif (re.match('^(top|bottom|left|right)_margin$', nam) or
                  re.match('^(上|下|左|右)余白$', nam)):
                val = unicodedata.normalize('NFKC', val)
                val = re.sub('\\s*cm$', '', val)
                if re.match('^' + RES_NUMBER + '$', val):
                    if nam == 'top_margin' or nam == '上余白':
                        Document.top_margin = float(val)
                    elif nam == 'bottom_margin' or nam == '下余白':
                        Document.bottom_margin = float(val)
                    elif nam == 'left_margin' or nam == '左余白':
                        Document.left_margin = float(val)
                    elif nam == 'right_margin' or nam == '右余白':
                        Document.right_margin = float(val)
                else:
                    msg = '※ 警告: ' \
                        + '「' + nam + '」の値は整数又は小数で' \
                        + 'なければなりません'
                    # msg = 'warning: ' \
                    #     + '"' + nam + '" must be an integer or a decimal'
                    sys.stderr.write(msg + '\n\n')
            elif nam == 'header_string' or nam == '頭書き':
                Document.header_string = val
            elif nam == 'page_number' or nam == '頁番号':
                val = unicodedata.normalize('NFKC', val)
                if val == 'True' or val == '有':
                    Document.page_number = DEFAULT_PAGE_NUMBER
                elif val == 'False' or val == '無' or val == '-':
                    Document.page_number = ''
                else:
                    Document.page_number = val
            elif nam == 'line_number' or nam == '行番号':
                val = unicodedata.normalize('NFKC', val)
                if val == 'True' or val == '有':
                    Document.line_number = True
                elif val == 'False' or val == '無':
                    Document.line_number = False
                else:
                    msg = '※ 警告: ' \
                        + '「' + nam + '」の値は"有"又は"無"で' \
                        + 'なければなりません'
                    # msg = 'warning: ' \
                    #     + '"' + nam + '" must be "True" or "False"'
                    sys.stderr.write(msg + '\n\n')
            elif nam == 'mincho_font' or nam == '明朝体':
                Document.mincho_font = val
            elif nam == 'gothic_font' or nam == 'ゴシ体':
                Document.gothic_font = val
            elif nam == 'ivs_font' or nam == '異字体':
                Document.ivs_font = val
            elif nam == 'font_size' or nam == '文字サ':
                val = unicodedata.normalize('NFKC', val)
                val = re.sub('\\s*pt$', '', val)
                if re.match('^' + RES_NUMBER + '$', val):
                    Document.font_size = float(val)
                else:
                    msg = '※ 警告: ' \
                        + '「' + nam + '」の値は整数又は小数で' \
                        + 'なければなりません'
                    # msg = 'warning: ' \
                    #     + '"' + nam + '" must be an integer or a decimal'
                    sys.stderr.write(msg + '\n\n')
            elif nam == 'line_spacing' or nam == '行間高':
                val = unicodedata.normalize('NFKC', val)
                val = re.sub('\\s*倍$', '', val)
                if re.match('^' + RES_NUMBER + '$', val):
                    Document.line_spacing = float(val)
                else:
                    msg = '※ 警告: ' \
                        + '「' + nam + '」の値は整数又は小数で' \
                        + 'なければなりません'
                    # msg = 'warning: ' \
                    #     + '"' + nam + '" must be an integer or a decimal'
                    sys.stderr.write(msg + '\n\n')
            elif (re.match('^space_(before|after)$', nam) or
                  re.match('^(前|後)余白$', nam)):
                val = unicodedata.normalize('NFKC', val)
                val = val.replace('、', ',')
                val = val.replace('倍', '')
                val = val.replace(' ', '')
                if re.match('^' + RES_NUMBER6 + '$', val):
                    if nam == 'space_before' or nam == '前余白':
                        Document.space_before = val
                    elif nam == 'space_after'or nam == '後余白':
                        Document.space_after = val
                else:
                    msg = '※ 警告: ' \
                        + '「' + nam + '」の値は' \
                        + '整数又は小数をカンマで区切って並べたもので' \
                        + 'なければなりません'
                    # msg = 'warning: ' \
                    #     + '"' + nam + '" must be 6 integers or decimals'
                    sys.stderr.write(msg + '\n\n')
            elif nam == 'auto_space' or nam == '字間整':
                val = unicodedata.normalize('NFKC', val)
                if val == 'True' or val == '有':
                    Document.auto_space = True
                elif val == 'False' or val == '無':
                    Document.auto_space = False
                else:
                    msg = '※ 警告: ' \
                        + '「' + nam + '」の値は"有"又は"無"で' \
                        + 'なければなりません'
                    # msg = 'warning: ' \
                    #     + '"' + nam + '" must be "True" or "False"'
                    sys.stderr.write(msg + '\n\n')
            elif nam == 'original_file' or nam == '元原稿':
                Document.original_file = val
            else:
                msg = '※ 警告: ' \
                    + '「' + nam + '」という設定項目は存在しません'
                # msg = 'warning: ' \
                #     + 'configuration name "' + nam + '" does not exist'
                sys.stderr.write(msg + '\n\n')

    def _configure_by_args(self, args):
        if args.document_title is not None:
            Document.document_title = args.document_title
        if args.paper_size is not None:
            Document.paper_size = args.paper_size
        if args.top_margin is not None:
            Document.top_margin = args.top_margin
        if args.bottom_margin is not None:
            Document.bottom_margin = args.bottom_margin
        if args.left_margin is not None:
            Document.left_margin = args.left_margin
        if args.right_margin is not None:
            Document.right_margin = args.right_margin
        if args.mincho_font is not None:
            Document.mincho_font = args.mincho_font
        if args.gothic_font is not None:
            Document.gothic_font = args.gothic_font
        if args.ivs_font is not None:
            Document.ivs_font = args.ivs_font
        if args.font_size is not None:
            Document.font_size = args.font_size
        if args.document_style is not None:
            Document.document_style = args.document_style
        if args.header_string is not None:
            Document.header_string = args.header_string
        if args.page_number is not None:
            Document.page_number = args.page_number
        if args.line_number:
            Document.line_number = True
        if args.line_spacing is not None:
            Document.line_spacing = args.line_spacing
        if args.space_before is not None:
            Document.space_before = args.space_before
        if args.space_after is not None:
            Document.space_after = args.space_after
        if args.auto_space:
            Document.auto_space = True

    def get_ms_doc(self):
        size = Document.font_size
        ms_doc = docx.Document()
        ms_sec = ms_doc.sections[0]
        ms_sec.page_height = Cm(PAPER_HEIGHT[Document.paper_size])
        ms_sec.page_width = Cm(PAPER_WIDTH[Document.paper_size])
        ms_sec.top_margin = Cm(Document.top_margin)
        ms_sec.bottom_margin = Cm(Document.bottom_margin)
        ms_sec.left_margin = Cm(Document.left_margin)
        ms_sec.right_margin = Cm(Document.right_margin)
        ms_sec.header_distance = Cm(1.0)
        ms_sec.footer_distance = Cm(1.0)
        ms_doc.styles['Footer'].font.size = Pt(size)      # page number
        ms_doc.styles['Normal'].font.size = Pt(size / 2)  # line number
        ms_doc.styles['List Bullet'].font.size = Pt(size)
        ms_doc.styles['List Bullet 2'].font.size = Pt(size)
        ms_doc.styles['List Bullet 3'].font.size = Pt(size)
        ms_doc.styles['List Number'].font.size = Pt(size)
        ms_doc.styles['List Number 2'].font.size = Pt(size)
        ms_doc.styles['List Number 3'].font.size = Pt(size)
        # HEADER
        if Document.header_string != '':
            hs = Document.header_string
            ms_par = ms_doc.sections[0].header.paragraphs[0]
            ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if re.match('^: (.*) :$', hs):
                hs = re.sub('^: (.*) :', '\\1', hs)
                ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif re.match('^: (.*)$', hs):
                hs = re.sub('^: (.*)', '\\1', hs)
                ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif re.match('^(.*) :$', hs):
                hs = re.sub('(.*) :$', '\\1', hs)
                ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            tex = ''
            is_large = False
            is_small = False
            for c in hs + '\0':
                tex += c
                if re.match(NOT_ESCAPED + '\\-\\-$', tex) or \
                   re.match(NOT_ESCAPED + '\\+\\+$', tex) or \
                   re.match(NOT_ESCAPED + '\0$', tex):
                    ms_run = ms_par.add_run()
                    if is_small:
                        ms_run.font.size = Pt(Document.font_size * 0.8)
                    elif is_large:
                        ms_run.font.size = Pt(Document.font_size * 1.2)
                    else:
                        ms_run.font.size = Pt(Document.font_size * 1.0)
                    oe = OxmlElement('w:t')
                    oe.set(ns.qn('xml:space'), 'preserve')
                    if re.match(NOT_ESCAPED + '\\-\\-$', tex):
                        oe.text = re.sub('\\-\\-$', '', tex)
                        is_small = not is_small
                        is_large = False
                    elif re.match(NOT_ESCAPED + '\\+\\+$', tex):
                        oe.text = re.sub('\\+\\+$', '', tex)
                        is_small = False
                        is_large = not is_large
                    elif re.match(NOT_ESCAPED + '\0$', tex):
                        oe.text = re.sub('\0$', '', tex)
                    tex = ''
                    ms_run._r.append(oe)
        # FOOTER
        if Document.page_number != '':
            pn = Document.page_number
            ms_par = ms_doc.sections[0].footer.paragraphs[0]
            ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if re.match('^: (.*) :$', pn):
                pn = re.sub('^: (.*) :', '\\1', pn)
                ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif re.match('^: (.*)$', pn):
                pn = re.sub('^: (.*)', '\\1', pn)
                ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif re.match('^(.*) :$', pn):
                pn = re.sub('(.*) :$', '\\1', pn)
                ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            tex = ''
            for c in pn + '\0':
                if c != 'n' and c != 'N' and c != '\0':
                    tex += c
                    continue
                if not re.match(NOT_ESCAPED + 'x', tex + 'x') and c != '\0':
                    tex = re.sub('\\\\$', '', tex) + c
                    continue
                # TEXT
                ms_run = ms_par.add_run()
                oe = OxmlElement('w:t')
                oe.set(ns.qn('xml:space'), 'preserve')
                oe.text = tex
                ms_run._r.append(oe)
                tex = ''
                # PAGE OR NUMPAGES
                if c == '\0':
                    continue
                pn = re.sub('^n', '', pn)
                ms_run = ms_par.add_run()
                oe = OxmlElement('w:fldChar')
                oe.set(ns.qn('w:fldCharType'), 'begin')
                ms_run._r.append(oe)
                oe = OxmlElement('w:instrText')
                oe.set(ns.qn('xml:space'), 'preserve')
                if c == 'n':
                    oe.text = 'PAGE'
                elif c == 'N':
                    oe.text = 'NUMPAGES'
                ms_run._r.append(oe)
                oe = OxmlElement('w:fldChar')
                oe.set(ns.qn('w:fldCharType'), 'end')
                ms_run._r.append(oe)
        # LINE NUMBER
        if Document.line_number:
            ms_scp = ms_doc.sections[0]._sectPr
            oe = OxmlElement('w:lnNumType')
            oe.set(ns.qn('w:countBy'), '5')
            oe.set(ns.qn('w:restart'), 'newPage')
            oe.set(ns.qn('w:distance'), '567')  # 567≒20*72/2.54=1cm
            ms_scp.append(oe)
        self.make_styles(ms_doc)
        return ms_doc

    def make_styles(self, ms_doc):
        size = Document.font_size
        line_spacing = Document.line_spacing
        # NORMAL
        ms_doc.styles.add_style('makdo', WD_STYLE_TYPE.PARAGRAPH)
        ms_doc.styles['makdo'].font.name = Document.mincho_font
        ms_doc.styles['makdo'].font.size = Pt(size)
        ms_doc.styles['makdo'].paragraph_format.line_spacing \
            = Pt(line_spacing * size)
        if not Document.auto_space:
            pPr = ms_doc.styles['makdo']._element.get_or_add_pPr()
            # KANJI<->ENGLISH
            oe = OxmlElement('w:autoSpaceDE')
            oe.set(ns.qn('w:val'), '0')
            pPr.append(oe)
            # KANJI<->NUMBER
            oe = OxmlElement('w:autoSpaceDN')
            oe.set(ns.qn('w:val'), '0')
            pPr.append(oe)
        # GOTHIC
        ms_doc.styles.add_style('makdo-g', WD_STYLE_TYPE.PARAGRAPH)
        ms_doc.styles['makdo-g'].font.name = Document.gothic_font
        # IVS
        ms_doc.styles.add_style('makdo-i', WD_STYLE_TYPE.PARAGRAPH)
        ms_doc.styles['makdo-i'].font.name = Document.ivs_font
        # TABLE
        ms_doc.styles.add_style('makdo-t', WD_STYLE_TYPE.PARAGRAPH)
        ms_doc.styles['makdo-t'].paragraph_format.line_spacing = Pt(size * 1.2)
        # ALIGNMENT
        # ms_doc.styles.add_style('makdo-a', WD_STYLE_TYPE.PARAGRAPH)
        # SPACE
        sb = Document.space_before.split(',')
        sa = Document.space_after.split(',')
        for i in range(6):
            n = 'makdo-' + str(i + 1)
            ms_doc.styles.add_style(n, WD_STYLE_TYPE.PARAGRAPH)
            if len(sb) > i and sb[i] != '':
                ms_doc.styles[n].paragraph_format.space_before \
                    = Pt(float(sb[i]) * line_spacing * size)
            if len(sa) > i and sa[i] != '':
                ms_doc.styles[n].paragraph_format.space_after \
                    = Pt(float(sa[i]) * line_spacing * size)

    def write_property(self, ms_doc):
        host = socket.gethostname()
        if host is None:
            host = '-'
        hh = self._get_hash(host)
        user = getpass.getuser()
        if user is None:
            user = '='
        hu = self._get_hash(user)
        tt = Document.document_title
        if Document.document_style == 'n':
            ct = '（普通）'
        elif Document.document_style == 'k':
            ct = '（契約）'
        elif Document.document_style == 'j':
            ct = '（条文）'
        at = hu + '@' + hh + ' (makdo ' + __version__ + ')'
        dt = datetime.datetime.utcnow()
        # utc = datetime.timezone.utc
        # pt = datetime.datetime(1970, 1, 1, 0, 0, 0, tzinfo=utc)
        # TIMEZONE IS NOT SUPPORTED
        # jst = datetime.timezone(datetime.timedelta(hours=9))
        # dt = datetime.datetime.now(jst)
        # pt = datetime.datetime(1970, 1, 1, 9, 0, 0, tzinfo=jst)
        ms_cp = ms_doc.core_properties
        ms_cp.identifier \
            = 'makdo(' + __version__.split()[0] + ');' \
            + hu + '@' + hh + ';' \
            + dt.strftime('%Y-%m-%dT%H:%M:%SZ')
        ms_cp.title = tt               # タイトル
        # ms_cp.subject = ''           # 件名
        # ms_cp.keywords = ''          # タグ
        ms_cp.category = ct            # 分類項目
        # ms_cp.comments = ''          # コメント
        ms_cp.author = at              # 作成者
        # ms_cp.last_modified_by = ''  # 前回保存者
        # ms_cp.revision = 1           # 改訂番号
        # ms_cp.version = ''           # バージョン番号
        ms_cp.created = dt             # コンテンツの作成日時
        ms_cp.modified = dt            # 前回保存時
        # ms_cp.last_printed = pt      # 前回印刷日
        # ms_cp.content_status = ''    # 内容の状態
        # ms_cp.language = ''          # 言語

    @staticmethod
    def _get_hash(st):
        # ''  owicwvnu
        # '-' sojfooxd
        # '=' empzhdhk
        x = 9973
        b = 99999989
        m = 999999999989
        z = int(((4 ** 20) - 1) / (4 - 1) * 2)
        for c in st + ' 2022.05.07 07:31:03':
            x = (x * b + ord(c)) % m
            x = x ^ z
        hs = ''
        for i in range(8):
            hs += chr(x % 26 + 97)
            x = int(x / 26)
        return hs

    def write_document(self, ms_doc):
        for p in self.paragraphs:
            p.write_paragraph(ms_doc)

    def save_docx_file(self, ms_doc, docx_file, md_file):
        if docx_file == '':
            if re.match('^.*\\.md$', md_file):
                docx_file = re.sub('\\.md$', '.docx', md_file)
            else:
                docx_file = md_file + '.docx'
            self.docx_file = docx_file
        if os.path.exists(docx_file):
            if not os.access(docx_file, os.W_OK):
                msg = '※ エラー: ' \
                    + '出力ファイル「' + docx_file + '」に書き込み権限が' \
                    + 'ありません'
                # msg = 'error: ' \
                #     + 'overwriting a unwritable file "' + docx_file + '"'
                sys.stderr.write(msg + '\n\n')
                sys.exit(1)
            if os.path.getmtime(md_file) < os.path.getmtime(docx_file):
                msg = '※ エラー: ' \
                    + '出力ファイル「' + docx_file + '」の方が' \
                    + '入力ファイル「' + md_file + '」よりも新しいです'
                # msg = 'error: ' \
                #     + 'overwriting a newer file "' + docx_file + '"'
                sys.stderr.write(msg + '\n\n')
                sys.exit(1)
            if os.path.exists(docx_file + '~'):
                os.remove(docx_file + '~')
            os.rename(docx_file, docx_file + '~')
        try:
            ms_doc.save(docx_file)
        except BaseException:
            msg = '※ エラー: ' \
                + '出力ファイル「' + docx_file + '」の書込みに失敗しました'
            # msg = 'error: ' \
            #     + 'failed to write output file "' + docx_file + '"'
            sys.stderr.write(msg + '\n\n')
            sys.exit(1)

    def print_warning_messages(self):
        for p in self.paragraphs:
            p.print_warning_messages()


class RawParagraph:

    """A class to handle raw paragraph"""

    raw_paragraph_number = 0

    def __init__(self, md_lines):
        # DECLARATION
        self.raw_paragraph_number = -1
        self.md_lines = []
        self.chapter_revisers = []
        self.section_revisers = []
        self.list_revisers = []
        self.length_revisers = []
        self.head_font_revisers = []
        self.tail_font_revisers = []
        self.full_text = ''
        self.section_depth_setters = []
        self.paragraph_class = ''
        # SUBSTITUTION
        RawParagraph.raw_paragraph_number += 1
        self.raw_paragraph_number = RawParagraph.raw_paragraph_number
        self.md_lines = md_lines
        self.chapter_revisers, \
            self.section_revisers, \
            self.list_revisers, \
            self.length_revisers, \
            self.head_font_revisers, \
            self.tail_font_revisers, \
            self.md_lines \
            = self._get_revisers(self.md_lines)
        self.full_text = self._get_full_text(self.md_lines)
        self.section_depth_setters, self.full_text \
            = self._get_section_depth_setters(self.full_text)
        self.paragraph_class = self._get_paragraph_class()

    @staticmethod
    def _get_revisers(md_lines):
        chapter_revisers = []
        section_revisers = []
        list_revisers = []
        length_revisers = []
        head_font_revisers = []
        tail_font_revisers = []
        res_cr = '^\\s*(' + ParagraphChapter.res_reviser + ') ?(.*)'
        res_sr = '^\\s*(' + ParagraphSection.res_reviser + ') ?(.*)'
        res_lr = '^(\\s*' + ParagraphList.res_reviser + ') ?(.*)'
        res_er = '^\\s*((?:v|V|X|<<|<|>)=' + RES_NUMBER + ') ?(.*)$'
        res_fr = '^\\s*(' + '|'.join(FONT_DECORATORS) + ') ?(.*)$'
        for ml in md_lines:
            # FOR BREAKDOWN
            if re.match('^-+::-*(::-+)?$', ml.text):
                break
            while True:
                if False:
                    pass
                elif re.match(res_cr, ml.text):
                    reviser = re.sub(res_cr, '\\1', ml.text)
                    ml.text = re.sub(res_cr, '\\5', ml.text)
                    chapter_revisers.append(reviser)
                elif re.match(res_sr, ml.text):
                    reviser = re.sub(res_sr, '\\1', ml.text)
                    ml.text = re.sub(res_sr, '\\5', ml.text)
                    section_revisers.append(reviser)
                elif re.match(res_lr, ml.text):
                    reviser = re.sub(res_lr, '\\1', ml.text)
                    ml.text = re.sub(res_lr, '\\3', ml.text)
                    list_revisers.append(reviser)
                elif re.match(res_er, ml.text):
                    reviser = re.sub(res_er, '\\1', ml.text)
                    ml.text = re.sub(res_er, '\\2', ml.text)
                    length_revisers.append(reviser)
                elif re.match(res_fr, ml.text):
                    reviser = re.sub(res_fr, '\\1', ml.text)
                    ml.text = re.sub(res_fr, '\\2', ml.text)
                    head_font_revisers.append(reviser)
                else:
                    break
            if ml.text != '':
                break
        res_fr = '(.*)(' + '|'.join(FONT_DECORATORS) + ')\\s*$'
        for ml in reversed(md_lines):
            while True:
                if False:
                    pass
                elif re.match(res_fr, ml.text):
                    reviser = re.sub(res_fr, '\\2', ml.text)
                    ml.text = re.sub(res_fr, '\\1', ml.text)
                    tail_font_revisers.insert(0, reviser)
                else:
                    break
            if ml.text != '':
                break
        # EXAMPLE "# ###=1"
        full_text = ''
        for ml in md_lines:
            if re.match('^.*\\S.*$', ml.text):
                full_text += ml.text
        res = '^' + \
            '(' + ParagraphSection.res_symbol + ')' + \
            '((\\s+' + ParagraphSection.res_reviser + ')+)' + \
            '\\s*$'
        if re.match(res, full_text):
            symbol = re.sub(res, '\\1', full_text)
            revisers = re.sub(res, '\\4', full_text)
            for ml in md_lines:
                ml.text = ''
            md_lines[0].text = symbol
            for r in revisers.split():
                section_revisers.append(r)
        # self.chapter_revisers = chapter_revisers
        # self.section_revisers = section_revisers
        # self.length_revisers = length_revisers
        # self.head_font_revisers = head_font_revisers
        # self.tail_font_revisers = tail_font_revisers
        # self.md_lines = md_lines
        return chapter_revisers, section_revisers, list_revisers, \
            length_revisers, head_font_revisers, tail_font_revisers, md_lines

    @staticmethod
    def _get_full_text(md_lines):
        full_text = ''
        for ml in md_lines:
            if ml.text != '':
                full_text += ml.text + ' '
        # FOR PARAGRAPH LIST
        list_head_spaces = ''
        res = '^( +)' + ParagraphList.res_symbol + '\\s+(.*)$'
        if re.match(res, full_text):
            list_head_spaces = re.sub(res, '\\1', full_text)
        full_text = re.sub('\t', ' ', full_text)
        full_text = re.sub(' +', ' ', full_text)
        full_text = re.sub('^ ', '', full_text)
        full_text = re.sub(' $', '', full_text)
        # FOR PARAGRAPH LIST
        full_text = list_head_spaces + full_text
        # self.full_text = full_text
        return full_text

    @staticmethod
    def _get_section_depth_setters(full_text):
        max_depth = len(ParagraphSection.states)
        section_depth_setters = []
        res = '^#{1,' + str(max_depth) + '}$'
        if re.match(res, full_text):
            section_depth_setters = [full_text]
            full_text = ''
        # self.section_depth_setters = depth_setters
        # self.full_text = full_text
        return section_depth_setters, full_text

    def _get_paragraph_class(self):
        ft = self.full_text
        hfrs = self.head_font_revisers
        tfrs = self.tail_font_revisers
        if False:
            pass
        elif ParagraphEmpty.is_this_class(ft, hfrs, tfrs):
            return 'empty'
        elif ParagraphBlank.is_this_class(ft, hfrs, tfrs):
            return 'blank'
        elif ParagraphChapter.is_this_class(ft, hfrs, tfrs):
            return 'chapter'
        elif ParagraphSection.is_this_class(ft, hfrs, tfrs):
            return 'section'
        elif ParagraphList.is_this_class(ft, hfrs, tfrs):
            return 'list'
        elif ParagraphTable.is_this_class(ft, hfrs, tfrs):
            return 'table'
        elif ParagraphImage.is_this_class(ft, hfrs, tfrs):
            return 'image'
        elif ParagraphAlignment.is_this_class(ft, hfrs, tfrs):
            return 'alignment'
        elif ParagraphPreformatted.is_this_class(ft, hfrs, tfrs):
            return 'preformatted'
        elif ParagraphPagebreak.is_this_class(ft, hfrs, tfrs):
            return 'pagebreak'
        elif ParagraphBreakdown.is_this_class(ft, hfrs, tfrs):
            return 'breakdown'
        else:
            return 'sentence'

    def get_paragraph(self):
        paragraph_class = self.paragraph_class
        if False:
            pass
        elif paragraph_class == 'empty':
            return ParagraphEmpty(self)
        elif paragraph_class == 'blank':
            return ParagraphBlank(self)
        elif paragraph_class == 'chapter':
            return ParagraphChapter(self)
        elif paragraph_class == 'section':
            return ParagraphSection(self)
        elif paragraph_class == 'list':
            return ParagraphList(self)
        elif paragraph_class == 'table':
            return ParagraphTable(self)
        elif paragraph_class == 'image':
            return ParagraphImage(self)
        elif paragraph_class == 'alignment':
            return ParagraphAlignment(self)
        elif paragraph_class == 'preformatted':
            return ParagraphPreformatted(self)
        elif paragraph_class == 'pagebreak':
            return ParagraphPagebreak(self)
        elif paragraph_class == 'breakdown':
            return ParagraphBreakdown(self)
        else:
            return ParagraphSentence(self)


class Paragraph:

    """A class to handle empty paragraph"""

    paragraph_number = 0

    paragraph_class = None
    res_feature = None

    mincho_font = ''
    gothic_font = ''
    ivs_font = ''
    font_size = -1

    previous_head_section_depth = 0
    previous_tail_section_depth = 0
    is_preformatted = False
    is_large = False
    is_small = False
    is_italic = False
    is_bold = False
    has_strike = False
    has_underline = False
    font_color = ''
    highlight_color = None

    @classmethod
    def is_this_class(cls, full_text,
                      head_font_revisers=[], tail_font_revisers=[]):
        if re.match(cls.res_feature, full_text):
            return True
        return False

    def __init__(self, raw_paragraph):
        # RECEIVE
        self.raw_paragraph_number = raw_paragraph.raw_paragraph_number
        self.md_lines = raw_paragraph.md_lines
        self.chapter_revisers = raw_paragraph.chapter_revisers
        self.section_revisers = raw_paragraph.section_revisers
        self.list_revisers = raw_paragraph.list_revisers
        self.length_revisers = raw_paragraph.length_revisers
        self.head_font_revisers = raw_paragraph.head_font_revisers
        self.tail_font_revisers = raw_paragraph.tail_font_revisers
        self.full_text = raw_paragraph.full_text
        self.section_depth_setters = raw_paragraph.section_depth_setters
        self.paragraph_class = raw_paragraph.paragraph_class
        # DECLARE
        self.paragraph_number = -1
        self.head_section_depth = -1
        self.tail_section_depth = -1
        self.proper_depth = -1
        self.length_revi = {}
        self.length_conf = {}
        self.length_clas = {}
        self.length_docx = {}
        self.alignment = ''
        self.text_to_write = ''
        self.text_to_write_with_reviser = ''
        # SUBSTITUTE
        Paragraph.paragraph_number += 1
        self.paragraph_number = Paragraph.paragraph_number
        self._apply_section_depths_setters(self.section_depth_setters)
        self.head_section_depth, self.tail_section_depth \
            = self._get_section_depths(self.full_text)
        self.proper_depth = self._get_proper_depth(self.full_text)
        self.alignment = self._get_alignment()
        # APPLY REVISERS
        ParagraphChapter._apply_revisers(self.chapter_revisers,
                                         self.md_lines)
        ParagraphSection._apply_revisers(self.section_revisers,
                                         self.md_lines)
        ParagraphList._apply_revisers(self.list_revisers,
                                      self.md_lines)
        ParagraphList.reset_states(self.paragraph_class)
        # GET LENGTH
        self.length_revi = self._get_length_revi()
        self.length_conf = self._get_length_conf()
        self.length_clas = self._get_length_clas()
        self.length_docx = self._get_length_docx()
        # GET TEXT
        self._edit_data()
        self.text_to_write = self._get_text_to_write()
        self.text_to_write_with_reviser \
            = self._get_text_to_write_with_reviser()

    @classmethod
    def _apply_section_depths_setters(cls, section_depth_setters):
        for sds in section_depth_setters:
            depth = len(sds)
            if depth > 0:
                Paragraph.previous_head_section_depth = depth
                Paragraph.previous_tail_section_depth = depth

    @classmethod
    def _get_section_depths(cls, full_text):
        head_section_depth = 0
        tail_section_depth = 0
        # self.head_section_depth = head_section_depth
        # self.tail_section_depth = tail_section_depth
        return head_section_depth, tail_section_depth

    @classmethod
    def _get_proper_depth(cls, full_text):
        proper_depth = 0
        # self.proper_depth = proper_depth
        return proper_depth

    def _get_alignment(self):
        paragraph_class = self.paragraph_class
        head_section_depth = self.head_section_depth
        full_text = self.full_text
        alignment = ''
        if paragraph_class == 'section' and head_section_depth == 1:
            alignment = 'center'
        if paragraph_class == 'alignment':
            if re.match('^:\\s.*\\s:$', full_text):
                alignment = 'center'
            elif re.match('^:\\s.*$', full_text):
                alignment = 'left'
            elif re.match('^.*\\s:$', full_text):
                alignment = 'right'
        # self.alignment = alignment
        return alignment

    @classmethod
    def _apply_revisers(cls, revisers, md_lines):
        res = '^' + cls.res_reviser + '$'
        if cls.paragraph_class == 'chapter':
            char = '$'
        elif cls.paragraph_class == 'section':
            char = '#'
        else:
            return
        for rev in revisers:
            md_line = md_lines[0]
            res_line = '^(.*\\s)?' \
                + rev.replace(char, '\\' + char) \
                + '(\\s.*)?$'
            for ml in md_lines:
                if re.match(res_line, ml.raw_text):
                    md_line = ml
                    break
            if re.match(res, rev):
                trunk = re.sub(res, '\\1', rev)
                branc = re.sub(res, '\\2', rev)
                chval = re.sub(res, '\\3', rev)
                xdepth = len(trunk) - 1
                ydepth = len(branc.replace(char, ''))
                value = int(chval) - 1
                cls._set_state(xdepth, ydepth, value, md_line)

    @classmethod
    def _set_state(cls, xdepth, ydepth, value, md_line):
        paragraph_class_ja = cls.paragraph_class_ja
        paragraph_class = cls.paragraph_class
        states = cls.states
        if xdepth >= len(states):
            msg = '※ 警告: ' + paragraph_class_ja \
                + 'の深さが上限を超えています'
            # msg = 'warning: ' + paragraph_class \
            #     + ' depth exceeds limit'
            md_line.append_warning_message(msg)
        elif ydepth >= len(states[xdepth]):
            msg = '※ 警告: ' + paragraph_class_ja \
                + 'の枝が上限を超えています'
            # msg = 'warning: ' + paragraph_class \
            #     + ' branch exceeds limit'
            md_line.append_warning_message(msg)
        for x in range(len(states)):
            for y in range(len(states[x])):
                if x < xdepth:
                    continue
                elif x == xdepth:
                    if y < ydepth:
                        if states[x][y] == 0:
                            msg = '※ 警告: ' + paragraph_class_ja \
                                + 'の枝が"0"を含んでいます'
                            # msg = 'warning: ' + paragraph_class \
                            #     + ' branch has "0"'
                            md_line.append_warning_message(msg)
                    elif y == ydepth:
                        if value is None:
                            states[x][y] += 1
                        else:
                            states[x][y] = value
                    else:
                        states[x][y] = 0
                else:
                    states[x][y] = 0

    def _get_length_revi(self):
        length_revisers = self.length_revisers
        length_revi \
            = {'space before': 0.0, 'space after': 0.0, 'line spacing': 0.0,
               'first indent': 0.0, 'left indent': 0.0, 'right indent': 0.0}
        res_v = '^v=(' + RES_NUMBER + ')$'
        res_cv = '^V=(' + RES_NUMBER + ')$'
        res_cx = '^X=(' + RES_NUMBER + ')$'
        res_gg = '^<<=(' + RES_NUMBER + ')$'
        res_g = '^<=(' + RES_NUMBER + ')$'
        res_l = '^>=(' + RES_NUMBER + ')$'
        for lr in length_revisers:
            if re.match(res_v, lr):
                length_revi['space before'] += float(re.sub(res_v, '\\1', lr))
            elif re.match(res_cv, lr):
                length_revi['space after'] += float(re.sub(res_cv, '\\1', lr))
            elif re.match(res_cx, lr):
                length_revi['line spacing'] += float(re.sub(res_cx, '\\1', lr))
            elif re.match(res_gg, lr):
                length_revi['first indent'] -= float(re.sub(res_gg, '\\1', lr))
            elif re.match(res_g, lr):
                length_revi['left indent'] -= float(re.sub(res_g, '\\1', lr))
            elif re.match(res_l, lr):
                length_revi['right indent'] -= float(re.sub(res_l, '\\1', lr))
        # self.length_revi = length_revi
        return length_revi

    def _get_length_conf(self):
        paragraph_class = self.paragraph_class
        hd = self.head_section_depth
        td = self.tail_section_depth
        sds = self.section_depth_setters
        has_section_depth_setters = False
        if paragraph_class != 'section' and len(sds) > 0:
            has_section_depth_setters = True
            hd = len(sds[0])
            td = len(sds[-1])
        length_conf \
            = {'space before': 0.0, 'space after': 0.0, 'line spacing': 0.0,
               'first indent': 0.0, 'left indent': 0.0, 'right indent': 0.0}
        sb = (Document.space_before + ',,,,,,,').split(',')
        sa = (Document.space_after + ',,,,,,,').split(',')
        if paragraph_class == 'section':
            if hd <= len(sb) and sb[hd - 1] != '':
                length_conf['space before'] += float(sb[hd - 1])
        if paragraph_class == 'section':
            if td <= len(sa) and sa[td - 1] != '':
                length_conf['space after'] += float(sa[td - 1])
        # self.length_conf = length_conf
        return length_conf

    def _get_length_clas(self):
        paragraph_class = self.paragraph_class
        head_section_depth = self.head_section_depth
        tail_section_depth = self.tail_section_depth
        proper_depth = self.proper_depth
        length_clas \
            = {'space before': 0.0, 'space after': 0.0, 'line spacing': 0.0,
               'first indent': 0.0, 'left indent': 0.0, 'right indent': 0.0}
        if paragraph_class == 'chapter':
            length_clas['first indent'] = -1.0
            length_clas['left indent'] = proper_depth + 0.0
        elif paragraph_class == 'section':
            if head_section_depth > 1:
                length_clas['first indent'] \
                    = head_section_depth - tail_section_depth - 1.0
            if tail_section_depth > 1:
                length_clas['left indent'] = tail_section_depth - 1.0
        elif paragraph_class == 'list':
            length_clas['first indent'] = -1.0
            length_clas['left indent'] = proper_depth + 0.0
            if tail_section_depth > 0:
                length_clas['left indent'] += tail_section_depth - 1.0
        elif paragraph_class == 'table':
            length_clas['space before'] += TABLE_SPACE_BEFORE
            length_clas['space after'] += TABLE_SPACE_AFTER
        elif paragraph_class == 'preformatted':
            if tail_section_depth > 0:
                length_clas['first indent'] = 0.0
                length_clas['left indent'] = tail_section_depth - 0.0
        elif paragraph_class == 'sentence':
            if tail_section_depth > 0:
                length_clas['first indent'] = 1.0
                length_clas['left indent'] = tail_section_depth - 1.0
        if paragraph_class == 'section' or \
           paragraph_class == 'list' or \
           paragraph_class == 'preformatted' or \
           paragraph_class == 'sentence':
            if ParagraphSection.states[1][0] <= 0 and tail_section_depth > 2:
                length_clas['left indent'] -= 1.0
        if Document.document_style == 'j':
            if ParagraphSection.states[1][0] > 0 and tail_section_depth > 2:
                length_clas['left indent'] -= 1.0
        # self.length_clas = length_clas
        return length_clas

    def _get_length_docx(self):
        paragraph_number = self.paragraph_number
        length_revi = self.length_revi
        length_conf = self.length_conf
        length_clas = self.length_clas
        length_docx \
            = {'space before': 0.0, 'space after': 0.0, 'line spacing': 0.0,
               'first indent': 0.0, 'left indent': 0.0, 'right indent': 0.0}
        for ln in length_docx:
            length_docx[ln] \
                = length_revi[ln] + length_conf[ln] + length_clas[ln]
        # LINE SPACING
        ls75 = length_docx['line spacing'] * .75
        ls25 = length_docx['line spacing'] * .25
        if length_docx['line spacing'] <= 0:
            if length_docx['space before'] >= ls75:
                length_docx['space before'] -= ls75
            elif length_docx['space before'] >= 0:
                length_docx['space before'] *= 2
            if length_docx['space after'] >= ls25:
                length_docx['space after'] -= ls25
            elif length_docx['space after'] >= 0:
                length_docx['space after'] *= 2
        else:
            if length_docx['space before'] >= ls75 * 2:
                length_docx['space before'] -= ls75
            elif length_docx['space before'] >= 0:
                length_docx['space before'] /= 2
            if length_docx['space after'] >= ls25 * 2:
                length_docx['space after'] -= ls25
            elif length_docx['space after'] >= 0:
                length_docx['space after'] /= 2
        # self.length_docx = length_docx
        return length_docx

    def _edit_data(self):
        return

    def _edit_data_of_chapter_and_section(self):
        paragraph_class = self.paragraph_class
        paragraph_class_ja = self.paragraph_class_ja
        res = self.res_feature
        md_lines = self.md_lines
        if paragraph_class == 'chapter':
            char = '$'
            paragraph_depth = self.proper_depth
        elif paragraph_class == 'section':
            char = '#'
            paragraph_depth = self.tail_section_depth
        else:
            return
        head_strings = ''
        title = ''
        body = ''
        pdepth = -1
        is_in_body = False
        for ml in md_lines:
            mlt = ml.text
            if not is_in_body:
                while re.match(res, mlt):
                    trunk = re.sub(res, '\\1', mlt)
                    branc = re.sub(res, '\\2', mlt)
                    mlt = re.sub(res, '\\3', mlt)
                    xdepth = len(trunk) - 1
                    ydepth = len(branc.replace(char, ''))
                    if pdepth > 0 and xdepth != pdepth + 1:
                        msg = '※ 警告: ' + paragraph_class_ja \
                            + 'の深さが飛んでいます'
                        # msg = 'warning: ' + paragraph_class \
                        #     + ' depth is not continuous'
                        ml.append_warning_message(msg)
                    pdepth = xdepth
                    head_strings += self._get_head_string(xdepth, ydepth, ml)
                    self._step_state(xdepth, ydepth, ml)
                if mlt != ml.text:
                    title = mlt
                    if re.match('^\\s+', title):
                        msg = '※ 警告: ' + paragraph_class_ja \
                            + 'のタイトルの行頭に空白があります'
                        # msg = 'warning: ' + paragraph_class \
                        #     + ' title has spaces at beginning'
                        ml.append_warning_message(msg)
                    ml.text = ''
                if mlt != '':
                    is_in_body = True
            if body == '' and re.match('^\\s+', ml.text):
                msg = '※ 警告: ' + paragraph_class_ja \
                    + 'の本文の行頭に空白があります'
                # msg = 'warning: ' + paragraph_class \
                #     + ' body has spaces at beginning'
                ml.append_warning_message(msg)
            body += ml.text
        if title + body == '':
            return
        if paragraph_class == 'section' and paragraph_depth == 1:
            md_lines[0].text = title
        elif re.match('^.*\\(.*\\)$', head_strings):
            md_lines[0].text = head_strings + ' ' + title
        else:
            md_lines[0].text = head_strings + '\u3000' + title
        return

    @classmethod
    def _step_state(cls, xdepth, ydepth, md_line):
        cls._set_state(xdepth, ydepth, None, md_line)

    def _get_text_to_write(self):
        md_lines = self.md_lines
        text_to_write = ''
        for ml in md_lines:
            text_to_write = concatenate_string(text_to_write, ml.text)
        # self.text_to_write = text_to_write
        return text_to_write

    def _get_text_to_write_with_reviser(self):
        text_to_write = self.text_to_write
        head_font_revisers = self.head_font_revisers
        tail_font_revisers = self.tail_font_revisers
        text_to_write_with_reviser \
            = ''.join(head_font_revisers) \
            + text_to_write \
            + ''.join(tail_font_revisers)
        return text_to_write_with_reviser

    def write_paragraph(self, ms_doc):
        paragraph_class = self.paragraph_class
        tail_section_depth = self.tail_section_depth
        alignment = self.alignment
        md_lines = self.md_lines
        text_to_write_with_reviser = self.text_to_write_with_reviser
        size = Paragraph.font_size
        xl_size = size * 1.4
        if text_to_write_with_reviser == '':
            return
        if paragraph_class == 'alignment':
            ms_par = self._get_ms_par(ms_doc)
            # ms_par = self._get_ms_par(ms_doc, 'makdo-a')
            oe = OxmlElement('w:wordWrap')
            oe.set(ns.qn('w:val'), '0')
            pPr = ms_par._p.get_or_add_pPr()
            pPr.append(oe)
        elif paragraph_class == 'preformatted':
            ms_par = self._get_ms_par(ms_doc, 'makdo-g')
        else:
            ms_par = self._get_ms_par(ms_doc)
        if alignment == 'left':
            ms_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'center':
            ms_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            ms_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif (paragraph_class == 'section' and
              re.sub('^\\S*\\s*', '', md_lines[0].text) == ''):
            ms_par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif (paragraph_class == 'sentence' and
              not re.match('^.*\n', text_to_write_with_reviser)):
            ms_par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ms_fmt = ms_par.paragraph_format
        if paragraph_class == 'section' and tail_section_depth == 1:
            Paragraph.font_size = xl_size
            self._write_text(text_to_write_with_reviser, ms_par)
            Paragraph.font_size = size
        else:
            self._write_text(text_to_write_with_reviser, ms_par)

    def _get_ms_par(self, ms_doc, par_style='makdo'):
        length_docx = self.length_docx
        size = Paragraph.font_size
        ms_par = ms_doc.add_paragraph(style=par_style)
        if not Document.auto_space:
            pPr = ms_par._p.get_or_add_pPr()
            oe = OxmlElement('w:autoSpaceDE')
            oe.set(ns.qn('w:val'), '0')
            pPr.append(oe)
            oe = OxmlElement('w:autoSpaceDN')
            oe.set(ns.qn('w:val'), '0')
            pPr.append(oe)
        ms_fmt = ms_par.paragraph_format
        ms_fmt.widow_control = False
        if length_docx['space before'] >= 0:
            pt = length_docx['space before'] * Document.line_spacing * size
            ms_fmt.space_before = Pt(pt)
        else:
            ms_fmt.space_before = Pt(0)
            msg = '警告: ' \
                + '段落前の余白「v」の値が小さ過ぎます'
            # msg = 'warning: ' \
            #     + '"space before" is too small'
            self.md_lines[0].append_warning_message(msg)
        if length_docx['space after'] >= 0:
            pt = length_docx['space after'] * Document.line_spacing * size
            ms_fmt.space_after = Pt(pt)
        else:
            ms_fmt.space_after = Pt(0)
            msg = '警告: ' \
                + '段落後の余白「V」の値が小さ過ぎます'
            # msg = 'warning: ' \
            #     + '"space after" is too small'
            self.md_lines[0].append_warning_message(msg)
        ms_fmt.first_line_indent = Pt(length_docx['first indent'] * size)
        ms_fmt.left_indent = Pt(length_docx['left indent'] * size)
        ms_fmt.right_indent = Pt(length_docx['right indent'] * size)
        # ms_fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        ls = Document.line_spacing * (1 + length_docx['line spacing'])
        if ls >= 1.0:
            ms_fmt.line_spacing = Pt(ls * size)
        else:
            ms_fmt.line_spacing = Pt(1.0 * size)
            msg = '警告: ' \
                + '段落後の余白「X」の値が少な過ぎます'
            # msg = 'warning: ' \
            #     + 'too small line spacing'
            self.md_lines[0].append_warning_message(msg)
        ms_fmt.line_spacing = Pt(ls * size)
        return ms_par

    def _write_text(self, text, ms_par):
        lns = text.split('\n')
        text = ''
        res = NOT_ESCAPED + '<br/?>'
        for ln in lns:
            while re.match(res, ln):
                ln = re.sub(res, '\\1\n', ln)
            text += ln + '\n'
        text = re.sub('\n$', '', text)
        text = Paragraph._remove_relax_symbol(text)
        res_img = '(.*(?:\n.*)*)' + RES_IMAGE
        res_ivs = '^(.*)([^\\\\])([0-9]+);$'
        tex = ''
        for c in text + '\0':
            if False:
                pass
            elif re.match(NOT_ESCAPED + '\\*\\*\\*$', tex + c):
                # *** (ITALIC AND BOLD)
                tex = re.sub('\\*\\*\\*$', '', tex + c)
                tex = self._write_string(tex, ms_par)
                Paragraph.is_italic = not Paragraph.is_italic
                Paragraph.is_bold = not Paragraph.is_bold
                continue
            elif re.match(NOT_ESCAPED + '\\*\\*$', tex) and c != '*':
                # ** (BOLD)
                tex = re.sub('\\*\\*$', '', tex)
                tex = self._write_string(tex, ms_par)
                tex += c
                Paragraph.is_bold = not Paragraph.is_bold
                continue
            elif re.match(NOT_ESCAPED + '\\*$', tex) and c != '*':
                # * (ITALIC)
                tex = re.sub('\\*$', '', tex)
                tex = self._write_string(tex, ms_par)
                tex += c
                Paragraph.is_italic = not Paragraph.is_italic
                continue
            elif re.match(NOT_ESCAPED + '~~$', tex + c):
                # ~~ (STRIKETHROUGH)
                tex = re.sub('~~$', '', tex + c)
                tex = self._write_string(tex, ms_par)
                Paragraph.has_strike = not Paragraph.has_strike
                continue
            elif re.match(NOT_ESCAPED + '`$', tex + c):
                # ` (PREFORMATTED)
                tex = re.sub('`$', '', tex + c)
                tex = self._write_string(tex, ms_par)
                Paragraph.is_preformatted = not Paragraph.is_preformatted
                continue
            elif re.match(NOT_ESCAPED + '//$', tex + c):
                # // (ITALIC)
                if not re.match('[a-z]+://', tex + c):
                    # not http:// https:// ftp:// ...
                    tex = re.sub('//$', '', tex + c)
                    tex = self._write_string(tex, ms_par)
                    Paragraph.is_italic = not Paragraph.is_italic
                    continue
            elif re.match(NOT_ESCAPED + '\\-\\-$', tex + c):
                # -- (SMALL)
                tex = re.sub('\\-\\-$', '', tex + c)
                tex = self._write_string(tex, ms_par)
                Paragraph.is_small = not Paragraph.is_small
                continue
            elif re.match(NOT_ESCAPED + '\\+\\+$', tex + c):
                # ++ (LARGE)
                tex = re.sub('\\+\\+$', '', tex + c)
                tex = self._write_string(tex, ms_par)
                Paragraph.is_large = not Paragraph.is_large
                continue
            elif re.match(NOT_ESCAPED + '\\^([0-9A-Za-z]*)\\^$', tex + c):
                # ^...^ (FONT COLOR)
                col = re.sub(NOT_ESCAPED + '\\^([0-9A-Za-z]*)\\^$', '\\2',
                             tex + c)
                if col == '':
                    col = 'FFFFFF'
                elif re.match('^([0-9A-F])([0-9A-F])([0-9A-F])$', col):
                    col = re.sub('^([0-9A-F])([0-9A-F])([0-9A-F])$',
                                 '\\1\\1\\2\\2\\3\\3', col)
                elif col in FONT_COLOR:
                    col = FONT_COLOR[col]
                if re.match('^[0-9A-F]{6}$', col):
                    tex = re.sub('\\^([0-9A-Za-z]*)\\^$', '', tex + c)
                    tex = self._write_string(tex, ms_par)
                    if Paragraph.font_color == '':
                        Paragraph.font_color = col
                    else:
                        Paragraph.font_color = ''
                    continue
            elif re.match(NOT_ESCAPED + '__$', tex + c):
                # __ (UNDERLINE)
                tex = re.sub('__$', '', tex + c)
                tex = self._write_string(tex, ms_par)
                Paragraph.has_underline = not Paragraph.has_underline
                continue
            elif re.match(NOT_ESCAPED + '_([0-9A-Za-z]+)_$', tex + c):
                # _..._ (HIGHLIGHT COLOR)
                col = re.sub(NOT_ESCAPED + '_([0-9A-Za-z]+)_$', '\\2', tex + c)
                if col in HIGHLIGHT_COLOR:
                    tex = re.sub('_([0-9A-Za-z]+)_$', '', tex + c)
                    tex = self._write_string(tex, ms_par)
                    if Paragraph.highlight_color is None:
                        Paragraph.highlight_color = HIGHLIGHT_COLOR[col]
                    else:
                        Paragraph.highlight_color = None
                    continue
            elif re.match(res_img, tex + c):
                # ![...](...)
                comm = re.sub(res_img, '\\2', tex + c)
                path = re.sub(res_img, '\\3', tex + c)
                tex = re.sub(res_img, '\\1', tex + c)
                tex = self._write_string(tex, ms_par)
                self._write_image(comm, path, ms_par)
                continue
            elif re.match(res_ivs, tex + c, flags=re.DOTALL):
                # .[0-9]+; (IVS (IDEOGRAPHIC VARIATION SEQUENCE))
                tmp_t = re.sub(res_ivs, '\\1', tex + c, flags=re.DOTALL)
                ivs_c = re.sub(res_ivs, '\\2', tex + c, flags=re.DOTALL)
                ivs_n = re.sub(res_ivs, '\\3', tex + c, flags=re.DOTALL)
                ivs_u = int('0xE0100', 16) + int(ivs_n)
                if int(ivs_n) <= int('0xE01EF', 16):
                    tex = self._write_string(tmp_t, ms_par)
                    pmf = Paragraph.mincho_font
                    Paragraph.mincho_font = Paragraph.ivs_font
                    self._write_string(ivs_c + chr(ivs_u), ms_par)
                    Paragraph.mincho_font = pmf
                    continue
            tex += c
        tex = re.sub('\0$', '', tex)
        if tex != '':
            tex = self._write_string(tex, ms_par)

    @staticmethod
    def _remove_relax_symbol(text):
        res = NOT_ESCAPED + RELAX_SYMBOL
        while re.match(res, text):
            text = re.sub(res, '\\1', text)
        return text

    @classmethod
    def _write_string(cls, string, ms_par):
        if string == '':
            return ''
        # REMOVE ESCAPE SYMBOL (BACKSLASH)
        string = re.sub('\\\\', '-\\\\', string)
        string = re.sub('-\\\\-\\\\', '-\\\\\\\\', string)
        string = re.sub('-\\\\', '', string)
        size = Paragraph.font_size
        l_size = 1.2 * size
        s_size = 0.8 * size
        ms_run = ms_par.add_run(string)
        if cls.is_large and not cls.is_small:
            ms_run.font.size = Pt(l_size)
        elif not cls.is_large and cls.is_small:
            ms_run.font.size = Pt(s_size)
        else:
            ms_run.font.size = Pt(size)
        if cls.is_bold:
            ms_run.bold = True
        # else:
        #     ms_run.bold = False
        if cls.is_preformatted:
            ms_run.font.name = cls.gothic_font
        else:
            ms_run.font.name = cls.mincho_font
        ms_run._element.rPr.rFonts.set(qn('w:eastAsia'), ms_run.font.name)
        if cls.is_italic:
            ms_run.italic = True
        # else:
        #     ms_run.italic = False
        if cls.has_strike:
            ms_run.font.strike = True
        # else:
        #     ms_run.font.strike = False
        if cls.has_underline:
            ms_run.underline = True
        # else:
        #     ms_run.underline = False
        if cls.font_color != '':
            r = int(re.sub('^(..)(..)(..)$', '\\1', cls.font_color), 16)
            g = int(re.sub('^(..)(..)(..)$', '\\2', cls.font_color), 16)
            b = int(re.sub('^(..)(..)(..)$', '\\3', cls.font_color), 16)
            ms_run.font.color.rgb = RGBColor(r, g, b)
        if cls.highlight_color is not None:
            ms_run.font.highlight_color = cls.highlight_color
        return ''

    def _write_image(self, alte, path, ms_par):
        size = Paragraph.font_size
        l_size = 1.2 * size
        s_size = 0.8 * size
        indent \
            = self.length_docx['first indent'] \
            + self.length_docx['left indent'] \
            + self.length_docx['right indent']
        text_width = PAPER_WIDTH[Document.paper_size] \
            - Document.left_margin - Document.right_margin \
            - (indent * 2.54 / 72)
        text_height = PAPER_HEIGHT[Document.paper_size] \
            - Document.top_margin - Document.bottom_margin
        ms_run = ms_par.add_run()
        res = '^(.*):(' + RES_NUMBER + ')?(?:x(' + RES_NUMBER + ')?)?$'
        cm_w = 0
        cm_h = 0
        if re.match(res, alte):
            st_w = re.sub(res, '\\2', alte)
            if st_w != '':
                cm_w = float(st_w)
                if cm_w < 0:
                    cm_w = text_width * (-1 * cm_w)
            st_h = re.sub(res, '\\3', alte)
            if st_h != '':
                cm_h = float(st_h)
                if cm_h < 0:
                    cm_h = text_height * (-1 * cm_h)
            alte = re.sub(res, '\\1', alte)
        try:
            if cm_w > 0 and cm_h > 0:
                ms_run.add_picture(path, width=Cm(cm_w), height=Cm(cm_h))
            elif cm_w > 0:
                ms_run.add_picture(path, width=Cm(cm_w))
            elif cm_h > 0:
                ms_run.add_picture(path, height=Cm(cm_h))
            elif self.is_large and not self.is_small:
                ms_run.add_picture(path, height=Pt(l_size))
            elif not self.is_large and self.is_small:
                ms_run.add_picture(path, height=Pt(s_size))
            else:
                ms_run.add_picture(path, height=Pt(size))
        except BaseException:
            ms_run.text = '![' + alte + '](' + path + ')'
            msg = '警告: ' \
                + 'インライン画像「' + path + '」が読み込めません'
            # msg = 'warning: can\'t open "' + path + '"'
            r = '^.*! *\\[.*\\] *\\(' + path + '\\).*$'
            for ml in self.md_lines:
                if re.match(r, ml.text):
                    if msg not in ml.warning_messages:
                        ml.append_warning_message(msg)
                        break
            else:
                self.md_lines[0].append_warning_message(msg)

    def print_warning_messages(self):
        for ml in self.md_lines:
            ml.print_warning_messages()


class ParagraphEmpty(Paragraph):

    """A class to handle empty paragraph"""

    paragraph_class = 'empty'
    res_feature = '^$'


class ParagraphBlank(Paragraph):

    """A class to handle blank paragraph"""

    paragraph_class = 'blank'
    res_feature = '^\n( \n)*$'


class ParagraphChapter(Paragraph):

    """A class to handle chapter paragraph"""

    paragraph_class = 'chapter'
    paragraph_class_ja = 'チャプター'
    res_symbol = '(\\$+)((?:\\-\\$+)*)'
    res_feature = '^' + res_symbol + '(?:\\s+(.*(?:.*\n*)*))?$'
    res_reviser = res_symbol + '=([0-9]+)'
    states = [[0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # 第１編
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # 第１章
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # 第１節
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # 第１款
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0]]  # 第１目
    unit_chars = ['編', '章', '節', '款', '目']

    @classmethod
    def _get_proper_depth(cls, full_text):
        if not re.match(cls.res_feature, full_text):
            return 0
        trunk = re.sub(cls.res_feature, '\\1', full_text)
        proper_depth = len(trunk)
        # self.proper_depth = proper_depth
        return proper_depth

    def _edit_data(self):
        self._edit_data_of_chapter_and_section()
        return

    @classmethod
    def _get_head_string(cls, xdepth, ydepth, md_line):
        xvalue_char = '〓'
        unit_char = '〓'
        if xdepth < len(cls.states):
            if ydepth < len(cls.states[xdepth]):
                value = cls.states[xdepth][0]
                if ydepth == 0:
                    value += 1
                xvalue_char = i2c_n_arab(value, md_line)
            unit_char = cls.unit_chars[xdepth]
        head_string = '第' + xvalue_char + unit_char
        for y in range(1, ydepth + 1):
            if y < len(cls.states[xdepth]):
                value = cls.states[xdepth][y] + 1
                if y == ydepth:
                    value += 1
                yvalue_char = i2c_n_arab(value, md_line)
            else:
                yvalue_char = '〓'
            head_string += 'の' + yvalue_char
        return head_string


class ParagraphSection(Paragraph):

    """A class to handle section paragraph"""

    paragraph_class = 'section'
    paragraph_class_ja = 'セクション'
    res_symbol = '(#+)((?:\\-#+)*)'
    res_feature = '^' + res_symbol + '(?:\\s+(.*(?:.*\n*)*))?$'
    res_reviser = res_symbol + '=([0-9]+)'
    states = [[0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # -
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # 第１
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # １
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # (1)
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # ア
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # (ｱ)
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # ａ
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0]]  # (a)

    @classmethod
    def is_this_class(cls, full_text,
                      head_font_revisers=[], tail_font_revisers=[]):
        if re.match(cls.res_feature, full_text):
            if not re.match('^#{15,}', full_text):
                return True
        return False

    @classmethod
    def _get_section_depths(cls, full_text):
        ft = full_text
        head_section_depth = 0
        tail_section_depth = 0
        while re.match(cls.res_feature, ft):
            trunk = re.sub(cls.res_feature, '\\1', ft)
            depth = len(trunk)
            if head_section_depth == 0:
                head_section_depth = depth
            tail_section_depth = depth
            ft = re.sub(cls.res_feature, '\\3', ft)
        Paragraph.previous_head_section_depth = head_section_depth
        Paragraph.previous_tail_section_depth = tail_section_depth
        # self.head_section_depth = head_section_depth
        # self.head_section_depth = tail_section_depth
        return head_section_depth, tail_section_depth

    def _edit_data(self):
        self._edit_data_of_chapter_and_section()
        return

    @classmethod
    def _get_head_string(cls, xdepth, ydepth, md_line):
        # TRUNK
        if xdepth < len(cls.states):
            value = cls.states[xdepth][0]
            if ydepth == 0:
                value += 1
            if xdepth == 0:
                head_string = ''
            elif xdepth == 1:
                if Document.document_style == 'n':
                    head_string = '第' + i2c_n_arab(value, md_line)
                else:
                    head_string = '第' + i2c_n_arab(value, md_line) + '条'
            elif xdepth == 2:
                if Document.document_style != 'j' or cls.states[1][0] == 0:
                    head_string = i2c_n_arab(value, md_line)
                else:
                    head_string = i2c_n_arab(value + 1, md_line)
            elif xdepth == 3:
                head_string = i2c_p_arab(value, md_line)
            elif xdepth == 4:
                head_string = i2c_n_kata(value, md_line)
            elif xdepth == 5:
                head_string = i2c_p_kata(value, md_line)
            elif xdepth == 6:
                head_string = i2c_n_alph(value, md_line)
            elif xdepth == 7:
                head_string = i2c_p_alph(value, md_line)
            else:
                head_string = '〓'
        else:
            head_string = '〓'
        # BRANCH
        for y in range(1, ydepth + 1):
            if y < len(cls.states[xdepth]):
                value = cls.states[xdepth][y] + 1
                if y == ydepth:
                    value += 1
                yvalue_char = i2c_n_arab(value, md_line)
            else:
                yvalue_char = '〓'
            head_string += 'の' + yvalue_char
        return head_string


class ParagraphList(Paragraph):

    """A class to handle list paragraph"""

    paragraph_class = 'list'
    paragraph_class_ja = 'リスト'
    res_symbol = '([-\\+\\*]|[0-9]+\\.)()'
    res_feature = '^\\s*' + res_symbol + '\\s+(.*)$'
    res_reviser = '[0-9]+\\.=([0-9]+)'
    states = [[0],  # ①
              [0],  # ㋐
              [0],  # ⓐ
              [0]]  # ㊀

    @classmethod
    def _get_section_depths(cls, full_text):
        head_section_depth = Paragraph.previous_tail_section_depth
        tail_section_depth = Paragraph.previous_tail_section_depth
        # self.head_section_depth = head_section_depth
        # self.tail_section_depth = tail_section_depth
        return head_section_depth, tail_section_depth

    @staticmethod
    def _get_proper_depth(full_text):
        full_text = re.sub('\u3000', '  ', full_text)
        full_text = re.sub('\t', '  ', full_text)
        full_text = re.sub('  ', ' ', full_text)
        spaces = re.sub('^( *).*$', '\\1', full_text)
        proper_depth = len(spaces) + 1
        # self.proper_depth = proper_depth
        return proper_depth

    @classmethod
    def _apply_revisers(cls, revisers, md_lines):
        for rev in revisers:
            res = '(\\s*).*=([0-9]+)'
            if re.match(res, rev):
                str_d = re.sub(res, '\\1', rev)
                str_v = re.sub(res, '\\2', rev)
                depth = len(re.sub('\\s\\s', ' ', str_d))
                value = int(str_v)
                cls.states[depth][0] = value - 1
                for d in range(depth + 1, len(cls.states)):
                    cls.states[d][0] = 0

    def _edit_data(self):
        res = '^\\s*' + ParagraphList.res_symbol + '\\s*'
        states = ParagraphList.states
        ml = self.md_lines
        proper_depth = self.proper_depth
        n = 0
        while n < len(ml) and ml[n].text == '':
            n += 1
        line = self.md_lines[n].text
        is_numbering = False
        if re.match('\\s*[0-9].\\s', line):
            is_numbering = True
        line = re.sub(res, '', line)
        if not is_numbering:
            if proper_depth == 1:
                head_strings = '・'
                # head_strings = '• '  # U+2022 Bullet
            elif proper_depth == 2:
                head_strings = '○'
                # head_strings = '◦ '  # U+25E6 White Bullet
            elif proper_depth == 3:
                head_strings = '△'
                # head_strings = '‣ '  # U+2023 Triangular Bullet
            elif proper_depth == 4:
                head_strings = '◇'
                # head_strings = '⁃ '  # U+2043 Hyphen Bullet
            else:
                head_strings = '〓'
        else:
            if proper_depth == 1:
                head_strings = i2c_c_arab(states[0][0] + 1)
            elif proper_depth == 2:
                head_strings = i2c_c_kata(states[1][0] + 1)
            elif proper_depth == 3:
                head_strings = i2c_c_alph(states[2][0] + 1)
            elif proper_depth == 4:
                head_strings = i2c_c_kanj(states[3][0] + 1)
            else:
                head_strings = '〓'
            if proper_depth <= len(states):
                states[proper_depth - 1][0] += 1
                for d in range(proper_depth, len(states)):
                    states[d][0] = 0
        self.md_lines[n].text = head_strings + '\u3000' + line

    @classmethod
    def reset_states(cls, paragraph_class):
        if paragraph_class != 'list':
            for s in cls.states:
                s[0] = 0
        return


class ParagraphTable(Paragraph):

    """A class to handle table paragraph"""

    paragraph_class = 'table'
    res_feature = '^\\|.*\\|$'

    def write_paragraph(self, ms_doc):
        size = Paragraph.font_size
        s_size = 0.8 * size
        tab = self._get_table_data(self.md_lines)
        conf_row, ali_list, wid_list = self._get_table_alignment_and_width(tab)
        if conf_row >= 0:
            tab.pop(conf_row)
        row = len(tab)
        col = len(tab[0])
        ms_tab = ms_doc.add_table(row, col, style='Table Grid')
        # ms_tab.autofit = True
        for i in range(len(tab)):
            ms_tab.rows[i].height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for j in range(len(tab[0])):
            ms_tab.columns[j].width = Pt((wid_list[j] + 2) * s_size)
        ms_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i in range(len(tab)):
            # ms_tab.rows[i].height = Pt(1.5 * size)
            for j in range(len(tab[i])):
                cell = tab[i][j]
                if ali_list[j] != WD_TABLE_ALIGNMENT.LEFT:
                    cell = re.sub('^\\s+', '', cell)
                if ali_list[j] != WD_TABLE_ALIGNMENT.RIGHT:
                    cell = re.sub('\\s+$', '', cell)
                ms_cell = ms_tab.cell(i, j)
                ms_cell.width = Pt((wid_list[j] + 2) * s_size)
                ms_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                ms_par = ms_cell.paragraphs[0]
                ms_par.style = 'makdo-t'
                Paragraph.font_size = s_size
                self._write_text(cell, ms_par)
                Paragraph.font_size = size
                ms_fmt = ms_par.paragraph_format
                if i < conf_row:
                    ms_fmt.alignment = WD_TABLE_ALIGNMENT.CENTER
                else:
                    ms_fmt.alignment = ali_list[j]

    @staticmethod
    def _get_table_data(md_lines):
        tab = []
        line = ''
        for ml in md_lines:
            if re.match('^\\\\?$', ml.text):
                continue
            if re.match('^.*\\\\$', line):
                line = re.sub('\\\\$', '', line)
                line += re.sub('^\\s*', '', ml.text)
            else:
                line += ml.text
            if re.match('^.*\\\\$', line):
                continue
            line = re.sub('^\\|', '', line)
            line = re.sub('\\|$', '', line)
            tab.append(line.split('|'))
            line = ''
        m = 0
        for rw in tab:
            if m < len(rw) - 1:
                m = len(rw) - 1
        for rw in tab:
            while len(rw) - 1 < m:
                rw.append('')
        # for i in range(len(tab)):
        #     for j in range(len(tab[i])):
        #         tab[i][j] = re.sub('^\\s+', '', tab[i][j])
        #         tab[i][j] = re.sub('\\s+$', '', tab[i][j])
        return tab

    @staticmethod
    def _get_table_alignment_and_width(tab):
        conf_row = -1
        for i in range(len(tab)):
            for j in range(len(tab[i])):
                if not re.match('^ *:?-*:? *$', tab[i][j]):
                    break
            else:
                conf_row = i
                break
        ali_list = []
        wid_list = []
        if conf_row >= 0:
            for s in tab[conf_row]:
                s = s.replace(' ', '')
                if re.match('^:-*:$', s):
                    ali_list.append(WD_TABLE_ALIGNMENT.CENTER)
                elif re.match('^-+:$', s):
                    ali_list.append(WD_TABLE_ALIGNMENT.RIGHT)
                else:
                    ali_list.append(WD_TABLE_ALIGNMENT.LEFT)
                wid_list.append(float(len(s)) / 2)
        else:
            for i in range(len(tab)):
                while len(ali_list) < len(tab[i]):
                    ali_list.append(WD_TABLE_ALIGNMENT.LEFT)
                while len(wid_list) < len(tab[i]):
                    wid_list.append(0.0)
                for j in range(len(tab[i])):
                    s = tab[i][j]
                    w = float(get_real_width(s)) / 2
                    if wid_list[j] < w:
                        wid_list[j] = w
        return conf_row, ali_list, wid_list


class ParagraphImage(Paragraph):

    """A class to handle image paragraph"""

    paragraph_class = 'image'
    res_feature = '^(?:\\s*' + RES_IMAGE + '\\s*)+$'

    def write_paragraph(self, ms_doc):
        ttwwr = self.text_to_write_with_reviser
        ttwwr = re.sub('\\s*(' + RES_IMAGE + ')\\s*', '\\1\n', ttwwr)
        ttwwr = re.sub('\n+', '\n', ttwwr)
        ttwwr = re.sub('^\n+', '', ttwwr)
        ttwwr = re.sub('\n+$', '', ttwwr)
        is_large = False
        is_small = False
        text_width = PAPER_WIDTH[Document.paper_size] \
            - Document.left_margin - Document.right_margin
        text_height = PAPER_HEIGHT[Document.paper_size] \
            - Document.top_margin - Document.bottom_margin
        res = '^(.*):(' + RES_NUMBER + ')?(?:x(' + RES_NUMBER + ')?)?$'
        for text in ttwwr.split('\n'):
            alte = re.sub(RES_IMAGE, '\\1', text)
            path = re.sub(RES_IMAGE, '\\2', text)
            cm_w = 0
            cm_h = 0
            if re.match(res, alte):
                st_w = re.sub(res, '\\2', alte)
                if st_w != '':
                    cm_w = float(st_w)
                    if cm_w < 0:
                        cm_w = text_width * (-1 * cm_w)
                st_h = re.sub(res, '\\3', alte)
                if st_h != '':
                    cm_h = float(st_h)
                    if cm_h < 0:
                        cm_h = text_height * (-1 * cm_h)
                alte = re.sub(res, '\\1', alte)
            try:
                if cm_w > 0 and cm_h > 0:
                    ms_doc.add_picture(path, width=Cm(cm_w), height=Cm(cm_h))
                elif cm_w > 0:
                    ms_doc.add_picture(path, width=Cm(cm_w))
                elif cm_h > 0:
                    ms_doc.add_picture(path, height=Cm(cm_h))
                else:
                    ms_doc.add_picture(path)
                ms_doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except BaseException:
                e = ms_doc.paragraphs[-1]._element
                e.getparent().remove(e)
                ms_par = self._get_ms_par(ms_doc)
                ms_par.add_run(text)
                ms_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                msg = '警告: ' \
                    + '画像「' + path + '」が読み込めません'
                # msg = 'warning: can\'t open "' + path + '"'
                r = '^.*! *\\[.*\\] *\\(' + path + '\\).*$'
                for ml in self.md_lines:
                    if re.match(r, ml.text):
                        if msg not in ml.warning_messages:
                            ml.append_warning_message(msg)
                            break
                else:
                    self.md_lines[0].append_warning_message(msg)


class ParagraphAlignment(Paragraph):

    """A class to handle alignment paragraph"""

    paragraph_class = 'alignment'
    res_feature = '^(?::|: .*|.* :)$'

    def _edit_data(self):
        md_lines = self.md_lines
        for ml in md_lines:
            if self.alignment == 'left' or self.alignment == 'center':
                ml.text = re.sub('^: ', '', ml.text)
            if self.alignment == 'center' or self.alignment == 'right':
                ml.text = re.sub(' :$', '', ml.text)
            if ml.text == ':':
                ml.text = ''

    def _get_text_to_write(self):
        md_lines = self.md_lines
        alignment = self.alignment
        text_to_write = ''
        for ml in md_lines:
            if alignment == 'left':
                if not re.match('^: .*$', ml.raw_text):
                    continue
            elif alignment == 'center':
                if not re.match('^: .* :$', ml.raw_text):
                    continue
            elif alignment == 'right':
                if not re.match('^.* :$', ml.raw_text):
                    continue
            text_to_write += ml.text + '\n'
        text_to_write = re.sub('\n$', '', text_to_write)
        return text_to_write


class ParagraphPreformatted(Paragraph):

    """A class to handle preformatted paragraph"""

    paragraph_class = 'preformatted'

    @classmethod
    def is_this_class(cls, full_text,
                      head_font_revisers, tail_font_revisers):
        if re.match('^```.*$', ''.join(head_font_revisers)) and \
           re.match('^.*```$', ''.join(tail_font_revisers)):
            return True
        return False

    @classmethod
    def _get_section_depths(cls, full_text):
        head_section_depth = Paragraph.previous_tail_section_depth
        tail_section_depth = Paragraph.previous_tail_section_depth
        # self.head_section_depth = head_section_depth
        # self.tail_section_depth = tail_section_depth
        return head_section_depth, tail_section_depth

    def _edit_data(self):
        self.head_font_revisers.pop(0)
        self.head_font_revisers.pop(0)
        self.head_font_revisers.pop(0)
        self.tail_font_revisers.pop(-1)
        self.tail_font_revisers.pop(-1)
        self.tail_font_revisers.pop(-1)
        self.md_lines[0].text = re.sub('\\s', '', self.md_lines[0].text)
        return

    def _get_text_to_write(self):
        md_lines = self.md_lines
        text_to_write = ''
        for i in range(len(md_lines)):
            if i == 0:
                if md_lines[i].text != '':
                    text_to_write += '[' + md_lines[i].text + ']\n'
            else:
                text_to_write += md_lines[i].text + '\n'
        text_to_write = re.sub('\n$', '', text_to_write)
        text_to_write = '`' + text_to_write + '`'
        return text_to_write


class ParagraphPagebreak(Paragraph):

    """A class to handle preformatted paragraph"""

    paragraph_class = 'pagebreak'
    res_feature = '^(?:<div style="break-.*: page;"></div>|<pgbr/?>)$'

    def write_paragraph(self, ms_doc):
        ms_doc.add_page_break()


class ParagraphBreakdown(Paragraph):

    """A class to handle breakdown paragraph"""

    paragraph_class = 'breakdown'
    res_feature = NOT_ESCAPED + '::.*$'


class ParagraphSentence(Paragraph):

    """A class to handle sentence paragraph"""

    paragraph_class = 'sentence'

    @classmethod
    def _get_section_depths(cls, full_text):
        head_section_depth = Paragraph.previous_tail_section_depth
        tail_section_depth = Paragraph.previous_tail_section_depth
        # self.head_section_depth = head_section_depth
        # self.tail_section_depth = tail_section_depth
        return head_section_depth, tail_section_depth


class MdLine:

    """A class to handle markdown line"""

    is_in_comment = False

    def __init__(self, line_number, raw_text):
        self.line_number = line_number
        self.raw_text = raw_text
        self.md_text, self.comment = self.separate_comment()
        self.text = self.md_text
        self.warning_messages = []

    def separate_comment(self):
        ori_sym = ORIGINAL_COMMENT_SYMBOL
        com_sep = COMMENT_SEPARATE_SYMBOL
        rt = self.raw_text
        text = ''
        comment = None
        if MdLine.is_in_comment:
            comment = ''
        tmp = ''
        for i, c in enumerate(rt):
            tmp += c
            if not MdLine.is_in_comment:
                if re.match(NOT_ESCAPED + '<!--$', tmp):
                    tmp = re.sub('<!--$', '', tmp)
                    text += tmp
                    tmp = ''
                    if comment is None:
                        comment = ''
                    MdLine.is_in_comment = True
            else:
                if re.match(NOT_ESCAPED + '-->$', tmp):
                    tmp = re.sub('-->$', '', tmp)
                    comment += tmp + com_sep
                    tmp = ''
                    MdLine.is_in_comment = False
            if not MdLine.is_in_comment:
                if re.match(NOT_ESCAPED + ori_sym + '$', tmp):
                    tmp = re.sub(ori_sym + '$', '', tmp)
                    text += tmp
                    tmp = ''
                    if comment is None:
                        comment = ''
                    comment += rt[i + 1:] + com_sep
                    break
        else:
            if tmp != '':
                if not MdLine.is_in_comment:
                    text += tmp
                    tmp = ''
                else:
                    if comment is None:
                        comment = ''
                    comment += tmp + com_sep
                    tmp = ''
        if comment is not None:
            comment = re.sub(com_sep + '$', '', comment)
        text = re.sub(NOT_ESCAPED + '<br/?>$', '\\1\n', text)
        # TRACK CHANGES
        res = NOT_ESCAPED + '<!?\\+>'
        while re.match(res, text):
            text = re.sub(res, '\\1', text)
        text = re.sub('  $', '\n', text)
        text = re.sub(' *$', '', text)
        # self.text = text
        # self.comment = comment
        return text, comment

    def append_warning_message(self, warning_message):
        self.warning_messages.append(warning_message)

    def print_warning_messages(self):
        for wm in self.warning_messages:
            msg = wm + '\n' \
                + '  (line ' + str(self.line_number) + ') ' + self.raw_text
            sys.stderr.write(msg + '\n\n')


############################################################
# MAIN


def main():

    args = get_arguments()

    doc = Document()

    doc.raw_md_lines = doc.get_raw_md_lines(args.md_file)

    doc.md_lines = doc.get_md_lines(doc.raw_md_lines)

    doc.configure(doc.md_lines, args)

    doc.raw_paragraphs = doc.get_raw_paragraphs(doc.md_lines)
    doc.paragraphs = doc.get_paragraphs(doc.raw_paragraphs)
    doc.paragraphs = doc.modify_paragraphs(doc.paragraphs)

    ms_doc = doc.get_ms_doc()

    doc.write_property(ms_doc)

    doc.write_document(ms_doc)

    doc.save_docx_file(ms_doc, args.docx_file, args.md_file)

    doc.print_warning_messages()

    sys.exit(0)


if __name__ == '__main__':
    main()
