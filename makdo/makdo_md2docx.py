#!/usr/bin/python3
# Name:         md2docx.py
# Version:      v07 Furuichibashi
# Time-stamp:   <2024.11.25-22:54:05-JST>

# md2docx.py
# Copyright (C) 2022-2024  Seiichiro HATA
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <http://www.gnu.org/licenses/>.


# 2022.07.21 v01 Hiroshima
# 2022.08.24 v02 Shin-Hakushima
# 2022.12.25 v03 Yokogawa
# 2023.01.07 v04 Mitaki
# 2023.03.16 v05 Aki-Nagatsuka
# 2023.06.07 v06 Shimo-Gion
# 2024.04.02 v07 Furuichibashi


# USAGE
# from makdo_md2docx import Md2Docx
# m2d = Md2Docx('xxx.md')
# m2d.set_document_title('aaa')
# m2d.set_document_style('bbb')
# m2d.set_paper_size('ccc')
# m2d.set_top_margin('ddd')
# m2d.set_bottom_margin('eee')
# m2d.set_left_margin('fff')
# m2d.set_right_margin('ggg')
# m2d.set_header_string('hhh')
# m2d.set_page_number('hhh')
# m2d.set_line_number('iii')
# m2d.set_mincho_font('jjj')
# m2d.set_gothic_font('kkk')
# m2d.set_ivs_font('lll')
# m2d.set_font_size('mmm')
# m2d.set_line_spacing('nnn')
# m2d.set_space_before('ooo')
# m2d.set_space_after('ppp')
# m2d.set_auto_space('qqq')
# m2d.set_version_number('rrr')
# m2d.set_content_status('sss')
# m2d.set_has_completed('ttt')
# m2d.save('xxx.docx')


############################################################
# NOTICE


# LibreOfficeには、一つのファイルの中で、
# ○●□■を半角文字と認識したり、全角文字と認識したりする
# バグがあるので、注意が必要です。
# （改行の次の行頭は、半角文字と認識するようです。）


############################################################
# POLICY

# document -> paragraph -> text -> chars -> imm


############################################################
# SETTING


import sys
import os
import argparse     # Python Software Foundation License
import re
import chardet      # GNU Lesser General Public License v2 or later (LGPLv2+)
import unicodedata
import datetime     # Zope Public License
import docx         # MIT License
from docx.shared import Cm, Pt
# from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
# from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement, ns
# from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
# from docx.enum.text import WD_UNDERLINE
from docx.enum.section import WD_SECTION
import socket   # host
import getpass  # user


__version__ = 'v07 Furuichibashi'


def get_arguments():
    parser = argparse.ArgumentParser(
        formatter_class=argparse.RawDescriptionHelpFormatter,
        description='MarkdownファイルからMS Wordファイルを作ります',
        add_help=False,
        epilog=HELP_EPILOG)
    parser.add_argument(
        '-h', '--help',
        action='help',
        help='ヘルプメッセージを表示します')
    parser.add_argument(
        '-v', '--version',
        action='version',
        version=('%(prog)s ' + __version__),
        help='バージョン番号を表示します')
    parser.add_argument(
        '-T', '--document-title',
        type=str,
        metavar='STRING',
        help='文書の標題')
    parser.add_argument(
        '-d', '--document-style',
        type=str,
        choices=['k', 'j'],
        help='文書スタイルの指定（契約、条文）')
    parser.add_argument(
        '-p', '--paper-size',
        type=str,
        choices=['A3', 'A3L', 'A3P', 'A4', 'A4L', 'A4P'],
        help='用紙設定（A3、A3L、A3P、A4、A4L、A4P）')
    parser.add_argument(
        '-t', '--top-margin',
        type=float,
        metavar='NUMBER',
        help='上余白（単位cm）')
    parser.add_argument(
        '-b', '--bottom-margin',
        type=float,
        metavar='NUMBER',
        help='下余白（単位cm）')
    parser.add_argument(
        '-l', '--left-margin',
        type=float,
        metavar='NUMBER',
        help='左余白（単位cm）')
    parser.add_argument(
        '-r', '--right-margin',
        type=float,
        metavar='NUMBER',
        help='右余白（単位cm）')
    parser.add_argument(
        '-H', '--header-string',
        type=str,
        help='ヘッダーの文字列')
    parser.add_argument(
        '-P', '--page-number',
        type=str,
        help='ページ番号の書式')
    parser.add_argument(
        '-L', '--line-number',
        action='store_true',
        help='行番号を出力します')
    parser.add_argument(
        '-m', '--mincho-font',
        type=str,
        metavar='FONT_NAME or ASCII_FONT_NAME/KANJI_FONT_NAME',
        help='明朝フォント')
    parser.add_argument(
        '-g', '--gothic-font',
        type=str,
        metavar='FONT_NAME or ASCII_FONT_NAME/KANJI_FONT_NAME',
        help='ゴシックフォント')
    parser.add_argument(
        '-i', '--ivs-font',
        type=str,
        metavar='FONT_NAME',
        help='異字体（IVS）フォント')
    # parser.add_argument(
    #     '--math_font',
    #     type=str,
    #     help=argparse.SUPPRESS)
    parser.add_argument(
        '-f', '--font-size',
        type=float,
        metavar='NUMBER',
        help='フォントサイズ（単位pt）')
    parser.add_argument(
        '-s', '--line-spacing',
        type=float,
        metavar='NUMBER',
        help='行間隔（単位文字）')
    parser.add_argument(
        '-B', '--space-before',
        type=floats6,
        metavar='NUMBER,NUMBER,...',
        help='セクションタイトル前の空白')
    parser.add_argument(
        '-A', '--space-after',
        type=floats6,
        metavar='NUMBER,NUMBER,...',
        help='セクションタイトル後の空白')
    parser.add_argument(
        '-a', '--auto-space',
        action='store_true',
        help='全角文字と半角文字との間の間隔を微調整します')
    parser.add_argument(
        '--version-number',
        type=str,
        metavar='VERSION_NUMBER',
        help='バージョン番号')
    parser.add_argument(
        '--content-status',
        type=str,
        metavar='CONTENT_STATUS',
        help='文書の状態')
    parser.add_argument(
        '-c', '--has-completed',
        action='store_true',
        help='備考書（コメント）などを消して完成させます')
    parser.add_argument(
        'md_file',
        help='Markdownファイル')
    parser.add_argument(
        'docx_file',
        default='',
        nargs='?',
        help='MS Wordファイル')
    return parser.parse_args()


def floats6(s):
    if not re.match('^' + RES_NUMBER6 + '$', s):
        msg = 'invalid 6 floats separated by commas value: \'' + s + '\''
        raise argparse.ArgumentTypeError(msg)
    return s


# def positive_integer(s):
#     if not re.match('[1-9][0-9]*', s):
#         msg = 'invalid positive integer value: \'' + s + '\''
#         raise argparse.ArgumentTypeError(msg)
#     return int(s)


HELP_EPILOG = '''Markdownの記法:
  段落指示
    [<pgbr>]で改ページされます（独自）
    [-----]（ハイフン5個以上）で水平線が挿入されます（独自）
  行頭指示
    [$+=(数字) ]でチャプター番号を変えることができます（独自）
    [$+ (タイトル)]でチャプターが挿入されます（独自）
    [#+=(数字) ]でセクション番号を変えることができます（独自）
    [#+ (タイトル)]でセクションが挿入されます
    [v=(数字) ]で段落の上の余白を行数だけ増減します（独自）
    [V=(数字) ]で段落の下の余白を行数だけ増減します（独自）
    [X=(数字) ]で段落の行間隔を行数だけ増減します（独自）
    [x=(数字) ]で段落の文字の間の間隔を文字数だけ増減します（独自）
    [<<=(数字) ]で段落1行目の左の余白を文字数だけ増減します（独自）
    [<=(数字) ]で段落の左の余白を文字数だけ増減します（独自）
    [>=(数字) ]で段落の右の余白を文字数だけ増減します（独自）
    ["" ]で段落の備考を付記することができます（独自）
  行中指示
    [->]から[<-]まで変更履歴の削除文字列になります（独自）
    [+>]から[<+]まで変更履歴の加筆文字列になります（独自）
    [<>]は何もせず表示もされません（独自）
    [<br>]で改行されます
  文字装飾
    [*]で挟まれた文字列は斜体になります
    [**]で挟まれた文字列は太字になります
    [***]で挟まれた文字列は斜体かつ太字になります
    [~~]で挟まれた文字列は打消線が引かれます
    [[|]と[|]]で囲まれた文字列は文字が枠で囲まれます（独自）
    [`]で挟まれた文字列はゴシック体になります
    [//]で挟まれた文字列は斜体になります（独自）
    [__]で挟まれた文字列は下線が引かれます（独自）
    [_foo_]で挟まれた文字列は特殊な下線が引かれます（独自）
      $(単語だけ) =(二重線) .(点線) #(太線) -(破線) .-(点破線) ..-(点々破線)
      ~(波線) .#(点太線) -#(破太線) .-#(点破太線) ..-#(点々破太線) ~#(波太線)
      -+(破長線) ~=(波二重線) -+#(破長太線)
    [---]で挟まれた文字列は文字がとても小さくなります（独自）
    [--]で挟まれた文字列は文字が小さくなります（独自）
    [++]で挟まれた文字列は文字が大きくなります（独自）
    [+++]で挟まれた文字列は文字がとても大きくなります（独自）
    [@N@]で囲まれた文字列は文字がNポイントの大きさになります（独自）
    [<<<]と[>>>]に挟まれた文字列は文字幅がとても広がります（独自）
    [<<]と[>>]に挟まれた文字列は文字幅が広がります（独自）
    [>>]と[<<]に挟まれた文字列は文字幅が狭まります（独自）
    [>>>]と[<<<]に挟まれた文字列は文字幅がとても狭まります（独自）
    [^^]で挟まれた文字列は白色になって見えなくなります（独自）
    [^XXYYZZ^]で挟まれた文字列はRGB(XX,YY,ZZ)色になります（独自）
    [^foo^]で挟まれた文字列はfoo色になります（独自）
      red(R) darkRed(DR) yellow(Y) darkYellow(DY) green(G) darkGreen(DG)
      cyan(C) darkCyan(DC) blue(B) darkBlue(DB) magenta(M) darkMagenta(DM)
      lightGray(G1) darkGray(G2) black(BK)
    [_foo_]で挟まれた文字列の背景はfoo色になります（独自）
      red(R) darkRed(DR) yellow(Y) darkYellow(DY) green(G) darkGreen(DG)
      cyan(C) darkCyan(DC) blue(B) darkBlue(DB) magenta(M) darkMagenta(DM)
      lightGray(G1) darkGray(G2) black(BK)
    [@foo@]で囲まれた文字列のフォントはfooになります（独自）
      "ＭＳ 明朝" "ＭＳ Ｐ明朝" "ＭＳ ゴシック" "ＭＳ Ｐゴシック"
      "游明朝" "游明朝 Light" "游明朝 Demibold"
      "游ゴシック" "游ゴシック Light" "游ゴシック Medium"
      "ヒラギノ明朝 Pro" "ヒラギノ明朝 ProN"
      "ヒラギノ丸ゴ Pro" "ヒラギノ丸ゴ ProN"
      "ヒラギノ角ゴ Pro" "ヒラギノ角ゴ ProN"
      "ヒラギノ角ゴ Std" "ヒラギノ角ゴ StdN"
      "IPA明朝" "IPA P明朝" "IPAex明朝"
      "IPAゴシック" "IPA Pゴシック" "IPAexゴシック"
      "Noto Serif JP" "Noto Serif JP Light" "Noto Serif JP Medium"
      "Noto Sans JP" "Noto Sans JP Light" "Noto Sans JP Medium"
    [字N;]（N=0-239）で"字"の異字体（IVS）が使えます（独自）
      ただし、"IPAmj明朝"フォント等がインストールされている必要があります
      参考：https://moji.or.jp/mojikiban/font/
            https://moji.or.jp/mojikibansearch/basic
    [^{foo}]でfooが上付文字（累乗等）になります（独自）
    [_{foo}]でfooが下付文字（添字等）になります（独自）
    [<foo/bar>]でfooの上にbarというルビ（ふりがな）が振られます（独自）
    [<N>]（Nは数字）で漢字N文字幅の空白が入ります（独自）
    [\\[]と[\\]]でLaTeX形式の文字列を挟むと数式が書けます（独自）
    [{{]と[}}]でPython風のスクリプトを挟むと簡単な計算できます（独自）
      変数の利用: "x=1" "x=2.3" "x=4+5" "x=6-y" "x=y*z"
      二項演算:   "x^y" "x+y" "x-y" "x*y" "x/y" "x//y" "x%y"
      複合計算:   "1+(2-3)" "4+(5-(6*7))" "x+(8.9-y)" "x*(y/z)"
      本文に印字: "print(1)" "print(2+3)" "print(x)" "print(x-y)"
    [{N{]と[}N}]（Nは1-9）でスクリプトを挟むと後の結果を利用できます（独自）
      "{{ print(x) }}{{ x=1 }}"はエラーになりますが、
      "{2{ print(x) }2}{{ x=1 }}"では"1"が印字されます
  エスケープ記号
    [\\]をコマンドの前に書くとコマンドが文字列になります
    [\\\\]で"\\"が表示されます
'''

DEFAULT_DOCUMENT_TITLE = ''

DEFAULT_DOCUMENT_STYLE = 'n'

DEFAULT_PAPER_SIZE = 'A4'
PAPER_HEIGHT = {'A3': 29.7, 'A3L': 29.7, 'A3P': 42.0,
                'A4': 29.7, 'A4L': 21.0, 'A4P': 29.7}
PAPER_WIDTH = {'A3': 42.0, 'A3L': 42.0, 'A3P': 29.7,
               'A4': 21.0, 'A4L': 29.7, 'A4P': 21.0}

DEFAULT_TOP_MARGIN = 3.5
DEFAULT_BOTTOM_MARGIN = 2.2
DEFAULT_LEFT_MARGIN = 3.0
DEFAULT_RIGHT_MARGIN = 2.3  # 21.0 - (2.54/72*12*37) - 3.0 = 2.3366666...
# DEFAULT_RIGHT_MARGIN = 1.9  # 21.0 - (2.54/72*12*38) - 3.0 = 1.9133333...
# DEFAULT_RIGHT_MARGIN = 2.0

DEFAULT_HEADER_STRING = ''

DEFAULT_PAGE_NUMBER = ': n :'

DEFAULT_LINE_NUMBER = False

DEFAULT_MINCHO_FONT = 'Times New Roman / ＭＳ 明朝'
DEFAULT_GOTHIC_FONT = 'ＭＳ ゴシック'
DEFAULT_IVS_FONT = 'IPAmj明朝'  # IPAmjMincho
# DEFAULT_MATH_FONT = 'Cambria Math'  # 'Liberation Serif'
DEFAULT_LINE_NUMBER_FONT = 'Calibri'
DEFAULT_FONT_SIZE = 12.0

DEFAULT_LINE_SPACING = 2.14  # (2.0980+2.1812)/2=2.1396
TABLE_LINE_SPACING = 1.5

DEFAULT_CHAR_SPACING = 0.0
# DEFAULT_CHAR_SPACING = 0.0208  # 5/12/20=.0208333...

DEFAULT_SPACE_BEFORE = ''
DEFAULT_SPACE_AFTER = ''
TABLE_SPACE_BEFORE = 0.45
TABLE_SPACE_AFTER = 0.20
IMAGE_SPACE_BEFORE = 0.68
IMAGE_SPACE_AFTER = 0.00

DEFAULT_AUTO_SPACE = False

DEFAULT_VERSION_NUMBER = ''

DEFAULT_CONTENT_STATUS = ''

DEFAULT_HAS_COMPLETED = False

BASIC_TABLE_CELL_HEIGHT = 1.5
BASIC_TABLE_CELL_WIDTH = 1.5  # >= 1.1068

NOT_ESCAPED = '^((?:(?:.|\n)*?[^\\\\])??(?:\\\\\\\\)*?)??'
# NOT_ESCAPED = '^((?:(?:.|\n)*[^\\\\])?(?:\\\\\\\\)*)?'

RES_NUMBER = '(?:[-\\+]?(?:(?:[0-9]+(?:\\.[0-9]+)?)|(?:\\.[0-9]+)))'
RES_NUMBER6 = '(?:' + RES_NUMBER + '?,){,5}' + RES_NUMBER + '?,?'

RES_IMAGE = '! *\\[([^\\[\\]]*)\\] *\\(([^\\(\\)]+)\\)'

FONT_DECORATORS_INVISIBLE = [
    '\\*\\*\\*',                # italic and bold
    '\\*\\*',                   # bold
    '\\*',                      # italic
    '//',                       # italic
    '\\^[0-9A-Za-z]{0,11}\\^',  # font color
]
FONT_DECORATORS_VISIBLE = [
    '\\-\\-\\-',                # xsmall
    '\\-\\-',                   # small
    '\\+\\+\\+',                # xlarge
    '\\+\\+',                   # large
    '>>>',                      # xnarrow or reset
    '>>',                       # narrow or reset
    '<<<',                      # xwide or reset
    '<<',                       # wide or reset
    '~~',                       # strikethrough
    '\\[\\|', '\\|\\]',         # frame
    '_[\\$=\\.#\\-~\\+]{,4}_',  # underline
    '_[0-9A-Za-z]{1,11}_',      # higilight color
    '`',                        # preformatted
    '@' + RES_NUMBER + '@',     # font scale
    '@[^@]{1,66}@',             # font
]
FONT_DECORATORS = FONT_DECORATORS_INVISIBLE + FONT_DECORATORS_VISIBLE

RELAX_SYMBOL = '<>'

TAB_WIDTH = 4

HORIZONTAL_BAR = '[ー−—－―‐]'

UNDERLINE = {
    '':     'single',
    '$':    'words',
    '=':    'double',
    '.':    'dotted',
    '#':    'thick',
    '-':    'dash',
    '.-':   'dotDash',
    '..-':  'dotDotDash',
    '~':    'wave',
    '.#':   'dottedHeavy',
    '-#':   'dashedHeavy',
    '.-#':  'dashDotHeavy',
    '..-#': 'dashDotDotHeavy',
    '~#':   'wavyHeavy',
    '-+':   'dashLong',
    '~=':   'wavyDouble',
    '-+#':  'dashLongHeavy',
}
# WD_UNDERLINE = {
#     '':     WD_UNDERLINE.SINGLE,
#     '$':    WD_UNDERLINE.WORDS,
#     '=':    WD_UNDERLINE.DOUBLE,
#     '.':    WD_UNDERLINE.DOTTED,
#     '#':    WD_UNDERLINE.THICK,
#     '-':    WD_UNDERLINE.DASH,
#     '.-':   WD_UNDERLINE.DOT_DASH,
#     '..-':  WD_UNDERLINE.DOT_DOT_DASH,
#     '~':    WD_UNDERLINE.WAVY,
#     '.#':   WD_UNDERLINE.DOTTED_HEAVY,
#     '-#':   WD_UNDERLINE.DASH_HEAVY,
#     '.-#':  WD_UNDERLINE.DOT_DASH_HEAVY,
#     '..-#': WD_UNDERLINE.DOT_DOT_DASH_HEAVY,
#     '~#':   WD_UNDERLINE.WAVY_HEAVY,
#     '-+':   WD_UNDERLINE.DASH_LONG,
#     '~=':   WD_UNDERLINE.WAVY_DOUBLE,
#     '-+#':  WD_UNDERLINE.DASH_LONG_HEAVY,
# }

FONT_COLOR = {
    'red':         'FF0000',
    'R':           'FF0000',
    'darkRed':     '7F0000',
    'DR':          '7F0000',
    'yellow':      'FFFF00',
    'Y':           'FFFF00',
    'darkYellow':  '7F7F00',
    'DY':          '7F7F00',
    'green':       '00FF00',
    'G':           '00FF00',
    'darkGreen':   '007F00',
    'DG':          '007F00',
    'cyan':        '00FFFF',
    'C':           '00FFFF',
    'darkCyan':    '007F7F',
    'DC':          '007F7F',
    'blue':        '0000FF',
    'B':           '0000FF',
    'darkBlue':    '00007F',
    'DB':          '00007F',
    'magenta':     'FF00FF',
    'M':           'FF00FF',
    'darkMagenta': '7F007F',
    'DM':          '7F007F',
    'lightGray':   'BFBFBF',
    'G1':          'BFBFBF',
    'darkGray':    '7F7F7F',
    'G2':          '7F7F7F',
    'black':       '000000',
    'BK':          '000000',
    'a000': 'FF5D5D',
    'a010': 'FF603C',
    'a020': 'FF6512',
    'a030': 'E07000',
    'a040': 'BC7A00',
    'a050': 'A08300',
    'a060': '898900',
    'a070': '758F00',
    'a080': '619500',
    'a090': '4E9B00',
    'a100': '38A200',
    'a110': '1FA900',
    'a120': '00B200',
    'a130': '00AF20',
    'a140': '00AC3C',
    'a150': '00AA55',
    'a160': '00A76D',
    'a170': '00A586',
    'a180': '00A2A2',
    'a190': '009FC3',
    'a200': '009AED',
    'a210': '1F8FFF',
    'a220': '4385FF',
    'a230': '5F7CFF',
    'a240': '7676FF',
    'a250': '8A70FF',
    'a260': '9E6AFF',
    'a270': 'B164FF',
    'a280': 'C75DFF',
    'a290': 'E056FF',
    'a300': 'FF4DFF',
    'a310': 'FF50DF',
    'a320': 'FF53C3',
    'a330': 'FF55AA',
    'a340': 'FF5892',
    'a350': 'FF5A79',
}

HIGHLIGHT_COLOR = {
    'red':         'red',
    'R':           'red',
    'darkRed':     'darkRed',
    'DR':          'darkRed',
    'yellow':      'yellow',
    'Y':           'yellow',
    'darkYellow':  'darkYellow',
    'DY':          'darkYellow',
    'green':       'green',
    'G':           'green',
    'darkGreen':   'darkGreen',
    'DG':          'darkGreen',
    'cyan':        'cyan',
    'C':           'cyan',
    'darkCyan':    'darkCyan',
    'DC':          'darkCyan',
    'blue':        'blue',
    'B':           'blue',
    'darkBlue':    'darkBlue',
    'DB':          'darkBlue',
    'magenta':     'magenta',
    'M':           'magenta',
    'darkMagenta': 'darkMagenta',
    'DM':          'darkMagenta',
    'lightGray':   'lightGray',
    'G1':          'lightGray',
    'darkGray':    'darkGray',
    'G2':          'darkGray',
    'black':       'black',
    'BK':          'black',
}
# WD_HIGHLIGHT_COLOR = {
#     'red':         WD_COLOR_INDEX.RED,
#     'R':           WD_COLOR_INDEX.RED,
#     'darkRed':     WD_COLOR_INDEX.DARK_RED,
#     'DR':          WD_COLOR_INDEX.DARK_RED,
#     'yellow':      WD_COLOR_INDEX.YELLOW,
#     'Y':           WD_COLOR_INDEX.YELLOW,
#     'darkYellow':  WD_COLOR_INDEX.DARK_YELLOW,
#     'DY':          WD_COLOR_INDEX.DARK_YELLOW,
#     'green':       WD_COLOR_INDEX.BRIGHT_GREEN,
#     'G':           WD_COLOR_INDEX.BRIGHT_GREEN,
#     'darkGreen':   WD_COLOR_INDEX.GREEN,
#     'DG':          WD_COLOR_INDEX.GREEN,
#     'cyan':        WD_COLOR_INDEX.TURQUOISE,
#     'C':           WD_COLOR_INDEX.TURQUOISE,
#     'darkCyan':    WD_COLOR_INDEX.TEAL,
#     'DC':          WD_COLOR_INDEX.TEAL,
#     'blue':        WD_COLOR_INDEX.BLUE,
#     'B':           WD_COLOR_INDEX.BLUE,
#     'darkBlue':    WD_COLOR_INDEX.DARK_BLUE,
#     'DB':          WD_COLOR_INDEX.DARK_BLUE,
#     'magenta':     WD_COLOR_INDEX.PINK,
#     'M':           WD_COLOR_INDEX.PINK,
#     'darkMagenta': WD_COLOR_INDEX.VIOLET,
#     'DM':          WD_COLOR_INDEX.VIOLET,
#     'lightGray':   WD_COLOR_INDEX.GRAY_25,
#     'G1':          WD_COLOR_INDEX.GRAY_25,
#     'darkGray':    WD_COLOR_INDEX.GRAY_50,
#     'G2':          WD_COLOR_INDEX.GRAY_50,
#     'black':       WD_COLOR_INDEX.BLACK,
#     'BK':          WD_COLOR_INDEX.BLACK,
# }

UNIX_TIME = datetime.datetime.timestamp(datetime.datetime.now())


############################################################
# FUNCTION


def get_real_width(s):
    p = ''
    wid = 0.0
    for c in s:
        if (c == '\t'):
            wid += (int(wid / TAB_WIDTH) + 1) * TAB_WIDTH
            continue
        w = unicodedata.east_asian_width(c)
        if c == '':
            wid += 0.0
        elif re.match('^[☐☑]$', c):
            wid += 2.0
        elif re.match('^[´¨―‐∥…‥‘’“”±×÷≠≦≧∞∴♂♀°′″℃§]$', c):
            wid += 2.0
        elif re.match('^[☆★○●◎◇◆□■△▲▽▼※→←↑↓]$', c):
            wid += 2.0
        elif re.match('^[∈∋⊆⊇⊂⊃∪∩∧∨⇒⇔∀∃∠⊥⌒∂∇≡≒≪≫√∽∝∵]$', c):
            wid += 2.0
        elif re.match('^[∫∬Å‰♯♭♪†‡¶◯]$', c):
            wid += 2.0
        elif re.match('^[ΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ]$', c):
            wid += 2.0
        elif re.match('^[αβγδεζηθικλμνξοπρστυφχψω]$', c):
            wid += 2.0
        elif re.match('^[АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ]$', c):
            wid += 2.0
        elif re.match('^[абвгдеёжзийклмнопрстуфхцчшщъыьэюя]$', c):
            wid += 2.0
        elif re.match('^[─│┌┐┘└├┬┤┴┼━┃┏┓┛┗┣┳┫┻╋┠┯┨┷┿┝┰┥┸╂]$', c):
            wid += 2.0
        elif re.match('^[№℡≒≡∫∮∑√⊥∠∟⊿∵∩∪]$', c):
            wid += 2.0
        elif re.match('^[⑴⑵⑶⑷⑸⑹⑺⑻⑼⑽⑾⑿⒀⒁⒂⒃⒄⒅⒆⒇]$', c):
            wid += 2.0
        elif re.match('^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳]$', c):
            wid += 2.0
        elif re.match('^[⒈⒉⒊⒋⒌⒍⒎⒏⒐⒑⒒⒓⒔⒕⒖⒗⒘⒙⒚⒛]$', c):
            wid += 2.0
        elif re.match('^[ⅰⅱⅲⅳⅴⅵⅶⅷⅸⅹⅺⅻ]$', c):
            wid += 2.0
        elif re.match('^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫ]$', c):
            wid += 2.0
        elif re.match('^[⒜⒝⒞⒟⒠⒡⒢⒣⒤⒥⒦⒧⒨⒩⒪⒫⒬⒭⒮⒯⒰⒱⒲⒳⒴⒵]$', c):
            wid += 2.0
        elif re.match('^[ⓐⓑⓒⓓⓔⓕⓖⓗⓘⓙⓚⓛⓜⓝⓞⓟⓠⓡⓢⓣⓤⓥⓦⓧⓨⓩ]$', c):
            wid += 2.0
        elif re.match('^[🄐🄑🄒🄓🄔🄕🄖🄗🄘🄙🄚🄛🄜🄝🄞🄟🄠🄡🄢🄣🄤🄥🄦🄧🄨🄩]$', c):
            wid += 2.0
        elif re.match('^[ⒶⒷⒸⒹⒺⒻⒼⒽⒾⒿⓀⓁⓂⓃⓄⓅⓆⓇⓈⓉⓊⓋⓌⓍⓎⓏ]$', c):
            wid += 2.0
        elif re.match('^[㉑㉒㉓㉔㉕㉖㉗㉘㉙㉚㉛㉜㉝㉞㉟㊱㊲㊳㊴㊵㊶㊷㊸㊹㊺㊻㊼㊽㊾㊿]$', c):
            wid += 2.0
        elif re.match('^[🄋➀➁➂➃➄➅➆➇➈➉]$', c):
            wid += 2.0
        elif re.match('^[㋐㋑㋒㋓㋔㋕㋖㋗㋘㋙㋚㋛㋜㋝㋞㋟㋠㋡㋢㋣㋤㋥㋦㋧㋨]$', c):
            wid += 2.0
        elif re.match('^[㋩㋪㋫㋬㋭㋮㋯㋰㋱㋲㋳㋴㋵㋶㋷㋸㋹㋺㋻㋼㋽㋾]$', c):
            wid += 2.0
        elif re.match('^[㊀㊁㊂㊃㊄㊅㊆㊇㊈㊉]$', c):
            wid += 2.0
        elif (w == 'F'):  # Full alphabet ...
            wid += 2.0
        elif(w == 'H'):   # Half katakana ...
            wid += 1.0
        elif(w == 'W'):   # Chinese character ...
            wid += 2.0
        elif(w == 'Na'):  # Half alphabet ...
            wid += 1.0
        elif(w == 'A'):   # Greek character ...
            wid += 1.0
        elif(w == 'N'):   # Arabic character ...
            wid += 1.0
        if p != '' and p != w:
            wid += 0.5
        p = w
    return wid


def n2c_n_arab(n, md_line=None):
    if n >= 0 and n <= 9:
        # ０１２３４５６７８９
        return chr(65296 + n)
    elif n >= 0:
        # 101112...
        return str(n)
    else:
        msg = '※ 警告: ' \
            + '数字番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed arabic number'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def n2c_p_arab(n, md_line=None):
    if n >= 0 and n == 0:
        # (0)
        return '(0)'
    elif n >= 0 and n <= 20:
        # ⑴⑵⑶⑷⑸⑹⑺⑻⑼⑽⑾⑿⒀⒁⒂⒃⒄⒅⒆⒇
        return chr(9331 + n)
    elif n >= 0:
        # (21)(22)(23)...
        return '(' + str(n) + ')'
    else:
        msg = '※ 警告: ' \
            + '括弧付き数字番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed parenthesis arabic number'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def n2c_c_arab(n, md_line=None):
    if n >= 0 and n == 0:
        # ⓪
        return chr(9450)
    elif n >= 0 and n <= 20:
        # ①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳
        return chr(9311 + n)
    elif n >= 0 and n <= 35:
        # ㉑㉒㉓㉔㉕㉖㉗㉘㉙㉚㉛㉜㉝㉞㉟
        return chr(12860 + n)
    elif n >= 0 and n <= 50:
        # ㊱㊲㊳㊴㊵㊶㊷㊸㊹㊺㊻㊼㊽㊾㊿
        return chr(12941 + n)
    else:
        msg = '※ 警告: ' \
            + '丸付き数字番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed circled arabic number'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def n2c_n_kata(n, md_line=None):
    if n >= 1 and n <= 5:
        # アイウエオ
        return chr(12448 + (2 * n))
    elif n >= 1 and n <= 17:
        # カキクケコサシスセソタチ
        return chr(12448 + (2 * n) - 1)
    elif n >= 1 and n <= 20:
        # ツテト
        return chr(12448 + (2 * n))
    elif n >= 1 and n <= 25:
        # ナニヌネノ
        return chr(12448 + (1 * n) + 21)
    elif n >= 1 and n <= 30:
        # ハヒフヘホ
        return chr(12448 + (3 * n) - 31)
    elif n >= 1 and n <= 35:
        # マミムメモ
        return chr(12448 + (1 * n) + 31)
    elif n >= 1 and n <= 38:
        # ヤユヨ
        return chr(12448 + (2 * n) - 4)
    elif n >= 1 and n <= 43:
        # ラリルレロ
        return chr(12448 + (1 * n) + 34)
    elif n >= 1 and n <= 48:
        # ワヰヱヲン
        return chr(12448 + (1 * n) + 35)
    else:
        msg = '※ 警告: ' \
            + 'カタカナ番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed katakana'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def n2c_p_kata(n, md_line=None):
    if n >= 1 and n <= 44:
        # (ｱ)(ｲ)(ｳ)(ｴ)(ｵ)(ｶ)(ｷ)(ｸ)(ｹ)(ｺ)(ｻ)(ｼ)(ｽ)(ｾ)(ｿ)
        # (ﾀ)(ﾁ)(ﾂ)(ﾃ)(ﾄ)(ﾅ)(ﾆ)(ﾇ)(ﾈ)(ﾉ)(ﾊ)(ﾋ)(ﾌ)(ﾍ)(ﾎ)
        # (ﾏ)(ﾐ)(ﾑ)(ﾒ)(ﾓ)(ﾔ)(ﾕ)(ﾖ)(ﾗ)(ﾘ)(ﾙ)(ﾚ)(ﾛ)(ﾜ)
        return '(' + chr(65392 + n) + ')'
    elif n >= 1 and n <= 45:
        # (ｦ)
        return '(' + chr(65392 + n - 55) + ')'
    elif n >= 1 and n <= 46:
        # (ﾝ)
        return '(' + chr(65392 + n - 1) + ')'
    else:
        msg = '※ 警告: ' \
            + '括弧付きカタカナ番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed parenthesis katakata'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def n2c_c_kata(n, md_line=None):
    if n >= 1 and n <= 47:
        # ㋐㋑㋒㋓㋔㋕㋖㋗㋘㋙㋚㋛㋜㋝㋞㋟㋠㋡㋢㋣㋤㋥㋦㋧㋨
        # ㋩㋪㋫㋬㋭㋮㋯㋰㋱㋲㋳㋴㋵㋶㋷㋸㋹㋺㋻㋼㋽㋾
        return chr(13007 + n)
    else:
        msg = '※ 警告: ' \
            + '丸付きカタカナ番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed circled katakana'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def n2c_n_alph(n, md_line=None):
    if n >= 1 and n <= 26:
        # ａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ
        return chr(65344 + n)
    else:
        msg = '※ 警告: ' \
            + 'アルファベット番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed alphabet'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def n2c_p_alph(n, md_line=None):
    if n >= 1 and n <= 26:
        # ⒜⒝⒞⒟⒠⒡⒢⒣⒤⒥⒦⒧⒨⒩⒪⒫⒬⒭⒮⒯⒰⒱⒲⒳⒴⒵
        return chr(9371 + n)
    else:
        msg = '※ 警告: ' \
            + '括弧付きアルファベット番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed parenthesis alphabet'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def n2c_c_alph(n, md_line=None):
    if n >= 1 and n <= 26:
        # ⓐⓑⓒⓓⓔⓕⓖⓗⓘⓙⓚⓛⓜⓝⓞⓟⓠⓡⓢⓣⓤⓥⓦⓧⓨⓩ
        return chr(9423 + n)
    else:
        msg = '※ 警告: ' \
            + '丸付きアルファベット番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed circled alphabet'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def n2c_n_kanj(n, md_line=None):
    if n >= 0:
        k = str(n)
        if n >= 10000:
            k = re.sub('^(.+)(....)$', '\\1万\\2', k)
        if n >= 1000:
            k = re.sub('^(.+)(...)$', '\\1千\\2', k)
        if n >= 100:
            k = re.sub('^(.+)(..)$', '\\1百\\2', k)
        if n >= 10:
            k = re.sub('^(.+)(.)$', '\\1十\\2', k)
        k = re.sub('0', '〇', k)
        k = re.sub('1', '一', k)
        k = re.sub('2', '二', k)
        k = re.sub('3', '三', k)
        k = re.sub('4', '四', k)
        k = re.sub('5', '五', k)
        k = re.sub('6', '六', k)
        k = re.sub('7', '七', k)
        k = re.sub('8', '八', k)
        k = re.sub('9', '九', k)
        k = re.sub('(.+)〇$', '\\1', k)
        k = re.sub('〇十', '', k)
        k = re.sub('〇百', '', k)
        k = re.sub('〇千', '', k)
        k = re.sub('一十', '十', k)
        k = re.sub('一百', '百', k)
        k = re.sub('一千', '千', k)
        return k
    else:
        msg = '※ 警告: ' \
            + '漢数字番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed kansuji'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def n2c_p_kanj(n, md_line=None):
    # ㈠㈡㈢㈣㈤㈥㈦㈧㈨㈩
    if n >= 1 and n <= 10:
        return chr(12831 + n)
    else:
        msg = '※ 警告: ' \
            + '括弧付き漢数字番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed parenthesis kansuji'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def n2c_c_kanj(n, md_line=None):
    # ㊀㊁㊂㊃㊄㊅㊆㊇㊈㊉
    if n >= 1 and n <= 10:
        return chr(12927 + n)
    else:
        msg = '※ 警告: ' \
            + '丸付き漢数字番号は範囲を超えています'
        # msg = 'warning: ' \
        #     + 'overflowed circled kansuji'
        if md_line is None:
            sys.stderr.write(msg + '\n\n')
        else:
            md_line.append_warning_message(msg)
        return '〓'


def concatenate_text(str1, str2):
    res = '[0-9A-Za-z,\\.\\)}\\]]'
    if re.match('^.*' + res + '$', str1) and re.match('^' + res + '.*$', str2):
        return str1 + ' ' + str2
    else:
        return str1 + str2


############################################################
# CLASS


class IO:

    """A class to handle input and output"""

    def __init__(self):
        self.inputed_md_file = None
        self.inputed_docx_file = None
        self.md_file = None
        self.docx_file = None
        self.ms_doc = None

    def set_md_file(self, inputed_md_file):
        md_file = inputed_md_file
        if not self._verify_input_file(md_file):
            return False
        self.inputed_md_file = inputed_md_file
        self.md_file = md_file
        return True

    def read_md_file(self):
        mf = MdFile(self.md_file)
        self.formal_md_lines = mf.read_file()
        return self.formal_md_lines

    def set_docx_file(self, inputed_docx_file):
        inputed_md_file = self.inputed_md_file
        md_file = self.md_file
        docx_file = inputed_docx_file
        if docx_file == '':
            if inputed_md_file == '-':
                msg = '※ エラー: ' \
                    + '出力ファイルの指定がありません'
                # msg = 'error: ' \
                #     + 'no output file name'
                sys.stderr.write(msg + '\n\n')
                if __name__ == '__main__':
                    sys.exit(201)
                return False
            elif re.match('^.*\\.md$', inputed_md_file):
                docx_file = re.sub('\\.md$', '.docx', inputed_md_file)
            else:
                docx_file = inputed_md_file + '.docx'
        if not self._verify_output_file(docx_file):
            return False
        if not self._verify_older(md_file, docx_file):
            return False
        self.inputed_docx_file = inputed_docx_file
        self.docx_file = docx_file
        return True

    def save_docx_file(self):
        ms_doc = self.ms_doc
        df = DocxFile(self.docx_file)
        df.write_file(ms_doc)
        return

    @staticmethod
    def _verify_input_file(input_file):
        if input_file == '-':
            return True
        if not os.path.exists(input_file):
            msg = '※ エラー: ' \
                + '入力ファイル「' + input_file + '」がありません'
            # msg = 'error: ' \
            #     + 'no input file "' + input_file + '"'
            sys.stderr.write(msg + '\n\n')
            if __name__ == '__main__':
                sys.exit(101)
            return False
        if not os.path.isfile(input_file):
            msg = '※ エラー: ' \
                + '入力「' + input_file + '」はファイルではありません'
            # msg = 'error: ' \
            #     + 'not a file "' + input_file + '"'
            sys.stderr.write(msg + '\n\n')
            if __name__ == '__main__':
                sys.exit(102)
            return False
        if not os.access(input_file, os.R_OK):
            msg = '※ エラー: ' \
                + '入力ファイル「' + input_file + '」に読込権限が' \
                + 'ありません'
            # msg = 'error: ' \
            #     + 'unreadable "' + input_file + '"'
            sys.stderr.write(msg + '\n\n')
            if __name__ == '__main__':
                sys.exit(103)
            return False
        return True

    @staticmethod
    def _verify_output_file(output_file):
        if output_file == '-':
            return True
        if not os.path.exists(output_file):
            return True
        if not os.path.isfile(output_file):
            msg = '※ エラー: ' \
                + '出力「' + output_file + '」はファイルではありません'
            # msg = 'error: ' \
            #     + 'not a file "' + output_file + '"'
            sys.stderr.write(msg + '\n\n')
            if __name__ == '__main__':
                sys.exit(202)
            return False
        if not os.access(output_file, os.W_OK):
            msg = '※ エラー: ' \
                + '出力ファイル「' + output_file + '」に書込権限が' \
                + 'ありません'
            # msg = 'error: ' \
            #     + 'unwritable "' + output_file + '"'
            sys.stderr.write(msg + '\n\n')
            if __name__ == '__main__':
                sys.exit(203)
            return False
        return True

    @staticmethod
    def _verify_older(input_file, output_file):
        if input_file != '-' and os.path.exists(input_file) and \
           output_file != '-' and os.path.exists(output_file):
            if os.path.getmtime(input_file) < os.path.getmtime(output_file):
                msg = '※ エラー: ' \
                    + '出力ファイルの方が入力ファイルよりも新しいです'
                # msg = 'error: ' \
                #     + 'overwriting newer file'
                sys.stderr.write(msg + '\n\n')
                if __name__ == '__main__':
                    sys.exit(301)
                return False
        return True

    def get_ms_doc(self):
        f_size = Form.font_size
        ms_doc = docx.Document()
        ms_sec = ms_doc.sections[0]
        ms_sec.page_height = Cm(PAPER_HEIGHT[Form.paper_size])
        ms_sec.page_width = Cm(PAPER_WIDTH[Form.paper_size])
        ms_sec.top_margin = Cm(Form.top_margin)
        ms_sec.bottom_margin = Cm(Form.bottom_margin)
        ms_sec.left_margin = Cm(Form.left_margin)
        ms_sec.right_margin = Cm(Form.right_margin)
        ms_sec.header_distance = Cm(1.0)
        ms_sec.footer_distance = Cm(1.0)
        # NORMAL (LINE NUMBER)
        ms_doc.styles['Normal'].font.size = Pt(f_size / 2)
        XML.set_font(ms_doc.styles['Normal'], DEFAULT_LINE_NUMBER_FONT)
        # LIST
        ms_doc.styles['List Bullet'].font.size = Pt(f_size)
        ms_doc.styles['List Bullet 2'].font.size = Pt(f_size)
        ms_doc.styles['List Bullet 3'].font.size = Pt(f_size)
        ms_doc.styles['List Number'].font.size = Pt(f_size)
        ms_doc.styles['List Number 2'].font.size = Pt(f_size)
        ms_doc.styles['List Number 3'].font.size = Pt(f_size)
        # HEADER
        # XML.set_font(ms_doc.styles['Header'], Form.mincho_font)
        # ms_doc.styles['Header'].font.size = Pt(f_size)
        if Form.header_string != '':
            # MDLINE
            ml = MdLine(-1, Form.header_string)
            # RAWPARAGRAPH
            pn = RawParagraph.raw_paragraph_number
            rp = RawParagraph([ml])
            RawParagraph.raw_paragraph_number = pn
            rp.raw_paragraph_number = -1
            rp.paragraph_class = 'alignment'
            # PARAGRAPH
            pn = Paragraph.paragraph_number
            p = rp.get_paragraph()
            Paragraph.paragraph_number = pn
            p.paragraph_number = -1
            # WRITE
            ms_par = ms_doc.sections[0].header.paragraphs[0]
            if p.alignment == 'right':
                ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif p.alignment == 'center':
                ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.write_text(ms_par, p.chars_state, p.text_to_write_with_reviser)
            Paragraph.bridge_chars_state.initialize()
        # FOOTER
        # XML.set_font(ms_doc.styles['Footer'], Form.mincho_font)
        # ms_doc.styles['Footer'].font.size = Pt(f_size)
        if Form.page_number != '':
            # MDLINE
            ml = MdLine(-1, Form.page_number)
            # RAWPARAGRAPH
            pn = RawParagraph.raw_paragraph_number
            rp = RawParagraph([ml])
            RawParagraph.raw_paragraph_number = pn
            rp.raw_paragraph_number = -2
            rp.paragraph_class = 'alignment'
            # PARAGRAPH
            pn = Paragraph.paragraph_number
            p = rp.get_paragraph()
            Paragraph.paragraph_number = pn
            p.paragraph_number = -2
            # WRITE
            ms_par = ms_doc.sections[0].footer.paragraphs[0]
            if p.alignment == 'right':
                ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif p.alignment == 'center':
                ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.write_text(ms_par, p.chars_state, p.text_to_write_with_reviser,
                         'footer')
            Paragraph.bridge_chars_state.initialize()
        # LINE NUMBER
        if Form.line_number:
            # MS WORD
            # if 'line number' not in ms_doc.styles:
            #     ms_doc.styles.add_style('line number',
            #                             WD_STYLE_TYPE.CHARACTER)
            # XML.set_font(ms_doc.styles['line number'],
            #              DEFAULT_LINE_NUMBER_FONT)
            # ms_doc.styles['line number'].font.size = Pt(f_size / 2)
            # LIBREOFFICE (ENGLISH)
            if 'Line Numbering' not in ms_doc.styles:
                ms_doc.styles.add_style('Line Numbering',
                                        WD_STYLE_TYPE.CHARACTER)
            XML.set_font(ms_doc.styles['Line Numbering'],
                         DEFAULT_LINE_NUMBER_FONT)
            ms_doc.styles['Line Numbering'].font.size = Pt(f_size / 2)
            # LIBREOFFICE (JAPANESE)
            if '行番号付け' not in ms_doc.styles:
                ms_doc.styles.add_style('行番号付け',
                                        WD_STYLE_TYPE.CHARACTER)
            XML.set_font(ms_doc.styles['行番号付け'],
                         DEFAULT_LINE_NUMBER_FONT)
            ms_doc.styles['行番号付け'].font.size = Pt(f_size / 2)
            opts = {}
            opts['w:countBy'] = '5'
            opts['w:restart'] = 'newPage'
            opts['w:distance'] = '567'  # 567≒20*72/2.54=1cm
            XML.add_tag(ms_doc.sections[0]._sectPr, 'w:lnNumType', opts)
        self._make_styles(ms_doc)
        return ms_doc

    def _make_styles(self, ms_doc):
        f_size = Form.font_size
        line_spacing = Form.line_spacing
        # NORMAL
        ms_doc.styles.add_style('makdo', WD_STYLE_TYPE.PARAGRAPH)
        XML.set_font(ms_doc.styles['makdo'], Form.mincho_font)
        ms_doc.styles['makdo'].font.size = Pt(f_size)
        ms_doc.styles['makdo'].paragraph_format.line_spacing \
            = Pt(line_spacing * f_size)
        if not Form.auto_space:
            ms_ppr = ms_doc.styles['makdo']._element.get_or_add_pPr()
            # KANJI<->ENGLISH
            XML.add_tag(ms_ppr, 'w:autoSpaceDE', {'w:val': '0'})
            # KANJI<->NUMBER
            XML.add_tag(ms_ppr, 'w:autoSpaceDN', {'w:val': '0'})
        # GOTHIC
        ms_doc.styles.add_style('makdo-g', WD_STYLE_TYPE.PARAGRAPH)
        XML.set_font(ms_doc.styles['makdo-g'], Form.gothic_font)
        # IVS
        ms_doc.styles.add_style('makdo-i', WD_STYLE_TYPE.PARAGRAPH)
        XML.set_font(ms_doc.styles['makdo-i'], Form.ivs_font)
        # TABLE
        ms_doc.styles.add_style('makdo-t', WD_STYLE_TYPE.PARAGRAPH)
        # ALIGNMENT
        # ms_doc.styles.add_style('makdo-a', WD_STYLE_TYPE.PARAGRAPH)
        # SECTION
        sb = Form.space_before.split(',')
        sa = Form.space_after.split(',')
        for i in range(6):
            n = 'makdo-' + str(i + 1)
            ms_doc.styles.add_style(n, WD_STYLE_TYPE.PARAGRAPH)
            if len(sb) > i and sb[i] != '':
                ms_doc.styles[n].paragraph_format.space_before \
                    = Pt(float(sb[i]) * line_spacing * f_size)
            if len(sa) > i and sa[i] != '':
                ms_doc.styles[n].paragraph_format.space_after \
                    = Pt(float(sa[i]) * line_spacing * f_size)
        # HORIZONTAL LINE
        ms_doc.styles.add_style('makdo-h', WD_STYLE_TYPE.PARAGRAPH)
        ms_doc.styles['makdo-h'].paragraph_format.line_spacing = 0
        ms_doc.styles['makdo-h'].font.size = Pt(f_size * 0.5)
        # MATH
        ms_doc.styles.add_style('makdo-m', WD_STYLE_TYPE.PARAGRAPH)
        # ms_doc.styles['makdo-m'].font.name = DEFAULT_MATH_FONT
        ms_doc.styles['makdo-m'].font.size = Pt(f_size)
        # REMARKS
        ms_doc.styles.add_style('makdo-r', WD_STYLE_TYPE.PARAGRAPH)
        ms_doc.styles['makdo-r'].paragraph_format.line_spacing = Pt(10.5)
        ms_doc.styles['makdo-r'].paragraph_format.space_before = Pt(10.5)
        ms_doc.styles['makdo-r'].paragraph_format.space_after = Pt(10.5)
        text_width = PAPER_WIDTH[Form.paper_size] \
            - Form.left_margin - Form.right_margin
        half_width = text_width / 2
        ms_doc.styles['makdo-r'].paragraph_format.first_line_indent = 0
        ms_doc.styles['makdo-r'].paragraph_format.left_indent = 0
        ms_doc.styles['makdo-r'].paragraph_format.right_indent = Cm(half_width)
        XML.set_font(ms_doc.styles['makdo-r'], Form.gothic_font)
        ms_doc.styles['makdo-r'].font.size = Pt(10.5)
        ms_doc.styles['makdo-r'].font.color.rgb = RGBColor(255, 255, 0)
        ms_doc.styles['makdo-r'].font.highlight_color = WD_COLOR_INDEX.BLUE


class MdFile:

    """A class to handle md file"""

    def __init__(self, md_file):
        # DECLARE
        self.md_file = None
        self.raw_data = None
        self.encoding = None
        self.decoded_data = None
        self.cleansed_data = None
        self.splited_data = None
        self.formal_md_lines = None
        # SUBSTITUTE
        self.md_file = md_file
        self.raw_data = self._get_raw_data(self.md_file)
        self.encoding = self._get_encoding(self.raw_data)
        self.decoded_data = self._decode_data(self.encoding, self.raw_data)
        self.cleansed_data = self._cleanse_data(self.decoded_data)
        self.splited_data = self._split_data(self.cleansed_data)
        self.formal_md_lines = self.splited_data

    def read_file(self):
        return self.formal_md_lines

    @staticmethod
    def _get_raw_data(md_file):
        if md_file is None:
            return ''
        try:
            if md_file == '-':
                raw_data = sys.stdin.buffer.read()
            else:
                raw_data = open(md_file, 'rb').read()
        except BaseException:
            msg = '※ エラー: ' \
                + '入力ファイル「' + md_file + '」の読込みに失敗しました'
            # msg = 'error: ' \
            #     + 'failed to read input file "' + md_file + '"'
            sys.stderr.write(msg + '\n\n')
            raise BaseException('failed to read input file')
            if __name__ == '__main__':
                sys.exit(104)
            return ''
        return raw_data

    @staticmethod
    def _get_encoding(raw_data):
        encoding = 'SHIFT_JIS'
        if raw_data != '':
            encoding = chardet.detect(raw_data)['encoding']
        if encoding is None:
            encoding = 'SHIFT_JIS'
        elif (re.match('^utf[-_]?.*$', encoding, re.I)) or \
             (re.match('^shift[-_]?jis.*$', encoding, re.I)) or \
             (re.match('^cp932.*$', encoding, re.I)) or \
             (re.match('^euc[-_]?(jp|jis).*$', encoding, re.I)) or \
             (re.match('^iso[-_]?2022[-_]?jp.*$', encoding, re.I)) or \
             (re.match('^ascii.*$', encoding, re.I)):
            pass
        else:
            # Windows-1252 (Western Europe)
            # MacCyrillic (Macintosh Cyrillic)
            # ...
            encoding = 'SHIFT_JIS'
            msg = '※ 警告: ' \
                + '文字コードを「SHIFT_JIS」に修正しました'
            # msg = 'warning: ' \
            #     + 'changed encoding to "SHIFT_JIS"'
            sys.stderr.write(msg + '\n\n')
        return encoding

    @staticmethod
    def _decode_data(encoding, raw_data):
        try:
            decoded_data = raw_data.decode(encoding)
        except BaseException:
            try:
                decoded_data = raw_data.decode('utf-8')
            except BaseException:
                msg = '※ エラー: ' \
                    + 'データを読みません（Markdownでないかも？）'
                # msg = 'error: ' \
                #     + 'can\'t read data (maybe not Markdown?)'
                sys.stderr.write(msg + '\n\n')
                raise BaseException('failed to read data')
                if __name__ == '__main__':
                    sys.exit(105)
                return ''
        return decoded_data

    @staticmethod
    def _cleanse_data(decoded_data):
        tmp_data = decoded_data
        bom = chr(65279)  # BOM (byte order mark)
        tmp_data = re.sub('^' + bom, '', tmp_data)  # unnecessary?
        tmp_data = re.sub('\r\n', '\n', tmp_data)  # unnecessary?
        tmp_data = re.sub('\r', '\n', tmp_data)  # unnecessary?
        # ISOLATE CONFIGURATIONS
        res = '^(<!--(?:.|\n)*?-->)\n*((?:.|\n)*)$'
        tmp_data = re.sub(res, '\\1\n\n\\2', tmp_data)
        cleansed_data = tmp_data
        return cleansed_data

    @staticmethod
    def _split_data(cleansed_data):
        splited_data = cleansed_data.split('\n')
        splited_data.append('')
        return splited_data


class DocxFile:

    """A class to handle docx file"""

    def __init__(self, docx_file):
        # DECLARE
        self.docx_file = None
        self.ms_doc = None
        # SUBSTITUTE
        self.docx_file = docx_file

    def write_file(self, ms_doc):
        self.ms_doc = ms_doc
        docx_file = self.docx_file
        self._save_old_file(docx_file)
        self._write_new_file(ms_doc, docx_file)

    @staticmethod
    def _save_old_file(output_file):
        if output_file is None:
            return False
        if output_file == '-':
            return True
        backup_file = output_file + '~'
        if os.path.exists(output_file):
            if os.path.exists(backup_file):
                os.remove(backup_file)
            if os.path.exists(backup_file):
                msg = '※ エラー: ' \
                    + '古いファイル「' + backup_file + '」を削除できません'
                # msg = 'error: ' \
                #     + 'can\'t remove "' + backup_file + '"'
                sys.stderr.write(msg + '\n\n')
                raise BaseException('failed to remove backup file')
                if __name__ == '__main__':
                    sys.exit(204)
                return False
            os.rename(output_file, backup_file)
        if os.path.exists(output_file):
            msg = '※ エラー: ' \
                + '古いファイル「' + output_file + '」を改名できません'
            # msg = 'error: ' \
            #     + 'can\'t rename "' + output_file + '"'
            sys.stderr.write(msg + '\n\n')
            raise BaseException('failed to rename old file')
            if __name__ == '__main__':
                sys.exit(205)
            return False
        return True

    @staticmethod
    def _write_new_file(ms_doc, docx_file):
        if docx_file is None:
            return False
        if docx_file == '-':
            output_file = '/dev/stdout'
        else:
            output_file = docx_file
        try:
            ms_doc.save(output_file)
        except BaseException:
            msg = '※ エラー: ' \
                + '出力ファイル「' + docx_file + '」の書込みに失敗しました'
            # msg = 'error: ' \
            #     + 'failed to write output file "' + docx_file + '"'
            sys.stderr.write(msg + '\n\n')
            raise BaseException('failed to write output file')
            if __name__ == '__main__':
                sys.exit(206)
            return False
        return True


class Form:

    """A class to handle form"""

    document_title = DEFAULT_DOCUMENT_TITLE
    document_style = DEFAULT_DOCUMENT_STYLE
    paper_size = DEFAULT_PAPER_SIZE
    top_margin = DEFAULT_TOP_MARGIN
    bottom_margin = DEFAULT_BOTTOM_MARGIN
    left_margin = DEFAULT_LEFT_MARGIN
    right_margin = DEFAULT_RIGHT_MARGIN
    header_string = DEFAULT_HEADER_STRING
    page_number = DEFAULT_PAGE_NUMBER
    line_number = DEFAULT_LINE_NUMBER
    mincho_font = DEFAULT_MINCHO_FONT
    gothic_font = DEFAULT_GOTHIC_FONT
    ivs_font = DEFAULT_IVS_FONT
    font_size = DEFAULT_FONT_SIZE
    line_spacing = DEFAULT_LINE_SPACING
    space_before = DEFAULT_SPACE_BEFORE
    space_after = DEFAULT_SPACE_AFTER
    auto_space = DEFAULT_AUTO_SPACE
    version_number = DEFAULT_VERSION_NUMBER
    content_status = DEFAULT_CONTENT_STATUS
    has_completed = DEFAULT_HAS_COMPLETED
    created_time = ''
    modified_time = ''

    def __init__(self):
        # DECLARE
        self.md_lines = None
        self.args = None

    def configure(self):
        # BY MD FILE
        self._configure_by_md_file(self.md_lines)
        # BY ARGUMENTS
        self._configure_by_args(self.args)
        # PRINT MESSAGES
        if self.has_completed:
            msg = '※ 警告: ' \
                + '備考書（コメント）は削除されます'
            # msg = 'warning: ' \
            #     + 'remarks(comments) is removed'
            sys.stderr.write(msg + '\n\n')
        # DOCUMENT TITLE
        if Form.document_title == '':
            Form.document_title = hex(int(UNIX_TIME * 1000000))
        # FOR LIBREOFFICE (NOT SUPPORT "SECTIONPAGES")
        has_two_or_more_sections = False
        for i in range(len(self.md_lines)):
            if not re.match('^\\s*<Pgbr>\\s*$', self.md_lines[i].text):
                continue
            if i > 0:
                if self.md_lines[i - 1].text != '':
                    continue
            if i < len(self.md_lines) - 1:
                if self.md_lines[i + 1].text != '':
                    continue
            has_two_or_more_sections = True
        if not has_two_or_more_sections:
            while re.match(NOT_ESCAPED + 'N', Form.page_number):
                Form.page_number \
                    = re.sub(NOT_ESCAPED + 'N', '\\1M', Form.page_number)
        elif re.match(NOT_ESCAPED + '(N|M)', Form.page_number):
            msg = '※ 警告: ' \
                + '"<Page>"を含む場合、' \
                + 'Libreofficeでは総ページ番号を適切に表示できません'
            # msg = 'warning: ' \
            #     + 'If "<Page>" is present, ' \
            #     + 'Libreoffice can\'t display total page numbers properly'
            sys.stderr.write(msg + '\n\n')

    @staticmethod
    def _configure_by_md_file(md_lines):
        for i, ml in enumerate(md_lines):
            if i == 0 and not re.match('^<!--.*$', ml.raw_text):
                break  # NO CONFIGURATIONS
            if i > 0 and re.match('^.*-->$', md_lines[i - 1].raw_text):
                break  # END OF CONFIGURATIONS
            if ml.text != '':
                break  # BEGINNING OF TEXT
            com = ml.comment
            if re.match('^\\s*#', com):
                continue
            res = '^\\s*([^:：]+)[:：]\\s*(.*)$'
            if not re.match(res, com):
                continue
            nam = re.sub(res, '\\1', com).rstrip()
            val = re.sub(res, '\\2', com).rstrip()
            if False:
                pass
            elif nam == 'document_title' or nam == '書題名':
                Form.set_document_title(val, nam)
            elif nam == 'document_style' or nam == '文書式':
                Form.set_document_style(val, nam)
            elif nam == 'paper_size' or nam == '用紙サ':
                Form.set_paper_size(val, nam)
            elif nam == 'top_margin' or nam == '上余白':
                Form.set_top_margin(val, nam)
            elif nam == 'bottom_margin' or nam == '下余白':
                Form.set_bottom_margin(val, nam)
            elif nam == 'left_margin' or nam == '左余白':
                Form.set_left_margin(val, nam)
            elif nam == 'right_margin' or nam == '右余白':
                Form.set_right_margin(val, nam)
            elif nam == 'header_string' or nam == '頭書き':
                Form.set_header_string(val, nam)
            elif nam == 'page_number' or nam == '頁番号':
                Form.set_page_number(val, nam)
            elif nam == 'line_number' or nam == '行番号':
                Form.set_line_number(val, nam)
            elif nam == 'mincho_font' or nam == '明朝体':
                Form.set_mincho_font(val, nam)
            elif nam == 'gothic_font' or nam == 'ゴシ体':
                Form.set_gothic_font(val, nam)
            elif nam == 'ivs_font' or nam == '異字体':
                Form.set_ivs_font(val, nam)
            elif nam == 'font_size' or nam == '文字サ':
                Form.set_font_size(val, nam)
            elif nam == 'line_spacing' or nam == '行間隔':
                Form.set_line_spacing(val, nam)
            elif nam == 'space_before' or nam == '前余白':
                Form.set_space_before(val, nam)
            elif nam == 'space_after' or nam == '後余白':
                Form.set_space_after(val, nam)
            elif nam == 'auto_space' or nam == '字間整':
                Form.set_auto_space(val, nam)
            elif nam == 'version_number' or nam == '版番号':
                Form.set_version_number(val, nam)
            elif nam == 'content_status' or nam == '書状態':
                Form.set_content_status(val, nam)
            elif nam == 'has_completed' or nam == '完成稿':
                Form.set_has_completed(val, nam)
            elif nam == 'created_time' or nam == '作成時':
                Form.set_created_time(val, nam)
            elif nam == 'modified_time' or nam == '更新時':
                Form.set_modified_time(val, nam)
            else:
                msg = '※ 警告: ' \
                    + '「' + nam + '」という設定項目は存在しません'
                # msg = 'warning: ' \
                #     + 'configuration name "' + nam + '" does not exist'
                sys.stderr.write(msg + '\n\n')

    @staticmethod
    def _configure_by_args(args):
        if args is not None:
            if args.document_title is not None:
                Form.set_document_title(args.document_title)
            if args.document_style is not None:
                Form.set_document_style(args.document_style)
            if args.paper_size is not None:
                Form.set_paper_size(args.paper_size)
            if args.top_margin is not None:
                Form.set_top_margin(str(args.top_margin))
            if args.bottom_margin is not None:
                Form.set_bottom_margin(str(args.bottom_margin))
            if args.left_margin is not None:
                Form.set_left_margin(str(args.left_margin))
            if args.right_margin is not None:
                Form.set_right_margin(str(args.right_margin))
            if args.header_string is not None:
                Form.set_header_string(args.header_string)
            if args.page_number is not None:
                Form.set_page_number(args.page_number)
            if args.line_number:
                Form.set_line_number(str(args.line_number))
            if args.mincho_font is not None:
                Form.set_mincho_font(args.mincho_font)
            if args.gothic_font is not None:
                Form.set_gothic_font(args.gothic_font)
            if args.ivs_font is not None:
                Form.set_ivs_font(args.ivs_font)
            if args.font_size is not None:
                Form.set_font_size(str(args.font_size))
            if args.line_spacing is not None:
                Form.set_line_spacing(str(args.line_spacing))
            if args.space_before is not None:
                Form.set_space_before(args.space_before)
            if args.space_after is not None:
                Form.set_space_after(args.space_after)
            if args.auto_space:
                Form.set_auto_space(str(args.auto_space))
            if args.version_number is not None:
                Form.set_version_number(args.version_number)
            if args.content_status is not None:
                Form.set_content_status(args.content_status)
            if args.has_completed:
                Form.set_has_completed(str(args.has_completed))

    @staticmethod
    def set_document_title(value, item='document_title'):
        if value is None:
            return False
        Form.document_title = value
        return True

    @staticmethod
    def set_document_style(value, item='document_style'):
        if value is None:
            return False
        if value == 'n' or value == '普通' or value == '-':
            Form.document_style = 'n'
            return True
        if value == 'k' or value == '契約':
            Form.document_style = 'k'
            return True
        if value == 'j' or value == '条文':
            Form.document_style = 'j'
            return True
        msg = '※ 警告: ' \
            + '「' + item + '」の値は' \
            + '"普通"、"契約"又は"条文"でなければなりません'
        # msg = 'warning: ' \
        #     + '"' + nam + '" must be "n", "k" or "j"'
        sys.stderr.write(msg + '\n\n')
        return False

    @staticmethod
    def set_paper_size(value, item='paper_size'):
        if value is None:
            return False
        value = unicodedata.normalize('NFKC', value)
        if value == 'A3':
            Form.paper_size = 'A3'
            return True
        elif value == 'A3L' or value == 'A3横':
            Form.paper_size = 'A3L'
            return True
        elif value == 'A3P' or value == 'A3縦':
            Form.paper_size = 'A3P'
            return True
        elif value == 'A4':
            Form.paper_size = 'A4'
            return True
        elif value == 'A4L' or value == 'A4横':
            Form.paper_size = 'A4L'
            return True
        elif value == 'A4P' or value == 'A4縦':
            Form.paper_size = 'A4P'
            return True
        msg = '※ 警告: ' \
            + '「' + item + '」の値は' \
            + '"A3横"、"A3縦"、"A4横"又は"A4縦"でなければなりません'
        # msg = 'warning: ' \
        #     + '"' + item + '" must be "A3", "A3P", "A4" or "A4L"'
        sys.stderr.write(msg + '\n\n')
        return False

    @staticmethod
    def set_top_margin(value, item='top_margin'):
        return Form._set_margin(value, item)

    @staticmethod
    def set_bottom_margin(value, item='bottom_margin'):
        return Form._set_margin(value, item)

    @staticmethod
    def set_left_margin(value, item='left_margin'):
        return Form._set_margin(value, item)

    @staticmethod
    def set_right_margin(value, item='right_margin'):
        return Form._set_margin(value, item)

    @staticmethod
    def _set_margin(value, item):
        if value is None:
            return False
        value = unicodedata.normalize('NFKC', value)
        value = re.sub('\\s*cm$', '', value)
        if re.match('^' + RES_NUMBER + '$', value):
            if item == 'top_margin' or item == '上余白':
                Form.top_margin = float(value)
                return True
            if item == 'bottom_margin' or item == '下余白':
                Form.bottom_margin = float(value)
                return True
            if item == 'left_margin' or item == '左余白':
                Form.left_margin = float(value)
                return True
            if item == 'right_margin' or item == '右余白':
                Form.right_margin = float(value)
                return True
        msg = '※ 警告: ' \
            + '「' + item + '」の値は' \
            + '整数又は小数でなければなりません'
        # msg = 'warning: ' \
        #     + '"' + item + '" must be an integer or a decimal'
        sys.stderr.write(msg + '\n\n')
        return False

    @staticmethod
    def set_header_string(value, item='header_string'):
        if value is None:
            return False
        Form.header_string = value
        return True

    @staticmethod
    def set_page_number(value, item='page_number'):
        if value is None:
            return False
        value = unicodedata.normalize('NFKC', value)
        if value == 'True' or value == '有':
            Form.page_number = DEFAULT_PAGE_NUMBER
            return True
        elif value == 'False' or value == 'None' or value == '無':
            Form.page_number = ''
            return True
        else:
            Form.page_number = value
            return True

    @staticmethod
    def set_line_number(value, item='line_number'):
        if value is None:
            return False
        value = unicodedata.normalize('NFKC', value)
        if value == 'True' or value == '有':
            Form.line_number = True
            return True
        elif value == 'False' or value == '無':
            Form.line_number = False
            return True
        msg = '※ 警告: ' \
            + '「' + item + '」の値は' \
            + '"有"又は"無"でなければなりません'
        # msg = 'warning: ' \
        #     + '"' + item + '" must be "True" or "False"'
        sys.stderr.write(msg + '\n\n')
        return False

    @staticmethod
    def set_mincho_font(value, item='mincho_font'):
        if value is None:
            return False
        Form.mincho_font = value
        return True

    @staticmethod
    def set_gothic_font(value, item='gothic_font'):
        if value is None:
            return False
        Form.gothic_font = value
        return True

    @staticmethod
    def set_ivs_font(value, item='ivs_font'):
        if value is None:
            return False
        Form.ivs_font = value
        return True

    @staticmethod
    def set_font_size(value, item='font_size'):
        if value is None:
            return False
        value = unicodedata.normalize('NFKC', value)
        value = re.sub('\\s*pt$', '', value)
        if re.match('^' + RES_NUMBER + '$', value):
            Form.font_size = float(value)
            return True
        msg = '※ 警告: ' \
            + '「' + item + '」の値は' \
            + '整数又は小数でなければなりません'
        # msg = 'warning: ' \
        #     + '"' + item + '" must be an integer or a decimal'
        sys.stderr.write(msg + '\n\n')
        return False

    @staticmethod
    def set_line_spacing(value, item='line_spacing'):
        if value is None:
            return False
        value = unicodedata.normalize('NFKC', value)
        value = re.sub('\\s*倍$', '', value)
        if re.match('^' + RES_NUMBER + '$', value):
            Form.line_spacing = float(value)
            return True
        msg = '※ 警告: ' \
            + '「' + item + '」の値は' \
            + '整数又は小数でなければなりません'
        # msg = 'warning: ' \
        #     + '"' + item + '" must be an integer or a decimal'
        sys.stderr.write(msg + '\n\n')
        return False

    @staticmethod
    def set_space_before(value, item='space_before'):
        return Form._set_space(value, item)

    @staticmethod
    def set_space_after(value, item='space_after'):
        return Form._set_space(value, item)

    @staticmethod
    def _set_space(value, item):
        if value is None:
            return False
        value = unicodedata.normalize('NFKC', value)
        value = value.replace('、', ',')
        value = value.replace('倍', '')
        value = value.replace(' ', '')
        if re.match('^' + RES_NUMBER6 + '$', value):
            if item == 'space_before' or item == '前余白':
                Form.space_before = value
                return True
            elif item == 'space_after' or item == '後余白':
                Form.space_after = value
                return True
        msg = '※ 警告: ' \
            + '「' + item + '」の値は' \
            + '整数又は小数をカンマで区切って並べたものでなければなりません'
        # msg = 'warning: ' \
        #     + '"' + item + '" must be 6 integers or decimals'
        sys.stderr.write(msg + '\n\n')
        return False

    @staticmethod
    def set_auto_space(value, item='auto_space'):
        if value is None:
            return False
        value = unicodedata.normalize('NFKC', value)
        if value == 'True' or value == '有':
            Form.auto_space = True
            return True
        elif value == 'False' or value == '無':
            Form.auto_space = False
            return True
        msg = '※ 警告: ' \
            + '「' + item + '」の値は' \
            + '"有"又は"無"でなければなりません'
        # msg = 'warning: ' \
        #     + '"' + item + '" must be "True" or "False"'
        sys.stderr.write(msg + '\n\n')

    @staticmethod
    def set_version_number(value, item='version_number'):
        if value is None:
            return False
        Form.version_number = value
        return True

    @staticmethod
    def set_content_status(value, item='content_status'):
        if value is None:
            return False
        Form.content_status = value
        return True

    @staticmethod
    def set_has_completed(value, item='has_completed'):
        if value is None:
            return False
        value = unicodedata.normalize('NFKC', value)
        if value == 'True' or value == '真':
            Form.has_completed = True
            return True
        elif value == 'False' or value == '偽':
            Form.has_completed = False
            return True
        msg = '※ 警告: ' \
            + '「' + item + '」の値は' \
            + '"真"又は"偽"でなければなりません'
        # msg = 'warning: ' \
        #     + '"' + item + '" must be "True" or "False"'
        sys.stderr.write(msg + '\n\n')

    @staticmethod
    def set_created_time(value, item='created_time'):
        if value is None:
            return False
        Form.created_time = value
        return True

    @staticmethod
    def set_modified_time(value, item='modified_time'):
        if value is None:
            return False
        Form.modified_time = value
        return True


class CharsState:

    """A class to keep character state"""

    def __init__(self):
        self.initialize()

    def initialize(self):
        self.mincho_font = Form.mincho_font
        self.gothic_font = Form.gothic_font
        self.ivs_font = Form.ivs_font
        self.font_size = Form.font_size
        self.font_scale = 1.0
        self.font_width = 1.0
        self.is_italic = False
        self.is_bold = False
        self.has_strike = False
        self.has_frame = False
        self.is_preformatted = False
        self.underline = None
        self.font_color = None
        self.highlight_color = None
        self.sub_or_sup = ''
        self.track_changes = ''  # ''|'del'|'ins'
        self.char_spacing = 0.0

    def copy(self):
        copy = CharsState()
        for v in vars(copy):
            vars(copy)[v] = vars(self)[v]
        return copy

    def apply_font_decorators(self, font_decorators):
        for fd in font_decorators:
            if False:
                pass
            elif fd == '*' or fd == '//':
                self.apply_is_italic_font_decorator(fd)
            elif fd == '**':
                self.apply_is_bold_font_decorator(fd)
            elif fd == '***':
                self.apply_is_italic_font_decorator(fd)
                self.apply_is_bold_font_decorator(fd)
            elif fd == '~~':
                self.apply_has_strike_font_decorator(fd)
            elif fd == '[|' or fd == '|]':
                self.apply_has_frame_font_decorator(fd)
            elif fd == '`':
                self.apply_is_preformatted_font_decorator(fd)
            elif fd == '---' or fd == '--' or fd == '++' or fd == '+++':
                self.apply_font_scale_font_decorator(fd)
            elif fd == '>>>' or fd == '>>' or fd == '<<' or fd == '<<<':
                self.apply_font_width_font_decorator(fd)
            elif re.match('^_[\\$=\\.#\\-~\\+]{,4}_$', fd):
                self.apply_underline_font_decorator(fd)
            elif re.match('^@' + RES_NUMBER + '@$', fd):
                self.apply_font_scale_font_decorator(fd)
            elif re.match('^@[^@]{1,66}@$', fd):
                self.apply_font_name_font_decorator(fd)
            elif re.match('^\\^[0-9A-Za-z]{0,11}\\^$', fd):
                self.apply_font_color_font_decorator(fd)
            elif re.match('^_[0-9A-Za-z]{1,11}_$', fd):
                self.apply_highlight_color_font_decorator(fd)
            elif fd == '_{' or fd == '^{' or fd == '}':
                self.apply_sub_or_sup_font_decorator(fd)
            elif fd == '->' or fd == '<-' or fd == '+>' or fd == '<+':
                self.apply_track_changes_font_decorator(fd)

    def apply_is_italic_font_decorator(self, font_decorator='*'):
        self.is_italic = not self.is_italic

    def apply_is_bold_font_decorator(self, font_decorator='**'):
        self.is_bold = not self.is_bold

    def apply_has_strike_font_decorator(self, font_decorator='~~'):
        self.has_strike = not self.has_strike

    def apply_has_frame_font_decorator(self, font_decorator):
        if font_decorator == '[|':
            self.has_frame = True
        elif font_decorator == '|]':
            self.has_frame = False

    def apply_is_preformatted_font_decorator(self, font_decorator='`'):
        self.is_preformatted = not self.is_preformatted

    def apply_font_scale_font_decorator(self, font_decorator):
        res = '^@(' + RES_NUMBER + ')@$'
        if re.match(res, font_decorator):
            c_size = float(re.sub(res, '\\1', font_decorator))
            if c_size > 0:
                font_scale = c_size / self.font_size
                if self.font_scale == font_scale:
                    self.font_scale = 1.0
                else:
                    self.font_scale = font_scale
        elif font_decorator == '---':
            if self.font_scale == 0.6:
                self.font_scale = 1.0
            else:
                self.font_scale = 0.6
        elif font_decorator == '--':
            if self.font_scale == 0.8:
                self.font_scale = 1.0
            else:
                self.font_scale = 0.8
        elif font_decorator == '++':
            if self.font_scale == 1.2:
                self.font_scale = 1.0
            else:
                self.font_scale = 1.2
        elif font_decorator == '+++':
            if self.font_scale == 1.4:
                self.font_scale = 1.0
            else:
                self.font_scale = 1.4
        else:
            self.font_scale = 1.0

    def apply_font_width_font_decorator(self, font_decorator):
        if font_decorator == '>>>':
            if self.font_width == 1.4:
                self.font_width = 1.0
            else:
                self.font_width = 0.6
        elif font_decorator == '>>':
            if self.font_width == 1.2:
                self.font_width = 1.0
            else:
                self.font_width = 0.8
        elif font_decorator == '<<':
            if self.font_width == 0.8:
                self.font_width = 1.0
            else:
                self.font_width = 1.2
        elif font_decorator == '<<<':
            if self.font_width == 0.6:
                self.font_width = 1.0
            else:
                self.font_width = 1.4
        else:
            self.font_width = 1.0

    def apply_underline_font_decorator(self, font_decorator='__'):
        underline = re.sub('^_(.*)_$', '\\1', font_decorator)
        if underline in UNDERLINE:
            if self.underline == UNDERLINE[underline]:
                self.underline = None
            else:
                self.underline = UNDERLINE[underline]
        else:
            self.underline = None

    def apply_font_size_font_decorator(self, font_decorator):
        font_size = float(re.sub('^@(.*)@$', '\\1', font_decorator))
        if self.font_size == font_size:
            self.font_size = Form.font_size
        else:
            self.font_size = font_size

    def apply_font_name_font_decorator(self, font_decorator):
        font = re.sub('^@(.*)@$', '\\1', font_decorator)
        if self.mincho_font != font or self.gothic_font != font:
            self.mincho_font = font
            self.gothic_font = font
        else:
            self.mincho_font = Form.mincho_font
            self.gothic_font = Form.gothic_font

    def apply_font_color_font_decorator(self, font_decorator):
        color = re.sub('^\\^(.*)\\^$', '\\1', font_decorator)
        if color == '':
            color = 'FFFFFF'
            self.font_color = 'FFFFFF'
        elif re.match('^([0-9A-F])([0-9A-F])([0-9A-F])$', color):
            color = re.sub('^([0-9A-F])([0-9A-F])([0-9A-F])$',
                           '\\1\\1\\2\\2\\3\\3', color)
        elif color in FONT_COLOR:
            color = FONT_COLOR[color]
        if self.font_color == color:
            self.font_color = None
        else:
            self.font_color = color

    def apply_highlight_color_font_decorator(self, font_decorator):
        color = re.sub('^_(.*)_$', '\\1', font_decorator)
        if color in HIGHLIGHT_COLOR:
            color = HIGHLIGHT_COLOR[color]
            if self.highlight_color == color:
                self.highlight_color = None
            else:
                self.highlight_color = color
        else:
            self.highlight_color = None

    def apply_sub_or_sup_font_decorator(self, font_decorator):
        if font_decorator == '_{':
            self.sub_or_sup = 'sub'
        elif font_decorator == '^{':
            self.sub_or_sup = 'sup'
        elif font_decorator == '}':
            self.sub_or_sup = ''
        else:
            self.sub_or_sup = ''

    def apply_track_changes_font_decorator(self, font_decorator):
        if font_decorator == '->' and self.track_changes == '':
            self.track_changes = 'del'
        elif font_decorator == '<-' and self.track_changes == 'del':
            self.track_changes = ''
        elif font_decorator == '+>' and self.track_changes == '':
            self.track_changes = 'ins'
        elif font_decorator == '<+' and self.track_changes == 'ins':
            self.track_changes = ''
        else:
            self.track_changes = ''


class XML:

    """A class to handle xml"""

    @staticmethod
    def add_tag(oe0, tag, opts={}, text=None):
        oe1 = OxmlElement(tag)
        for item in opts:
            value = opts[item]
            oe1.set(ns.qn(item), value)
        if text is not None:
            oe1.text = text
        oe0.append(oe1)
        return oe1

    @staticmethod
    def write_chars(oe0, chars_state, chars):
        if chars == '':
            return ''
        chars = XML._prepare_chars(chars)
        if chars_state.track_changes == 'del':
            oe1 = XML.add_tag(oe0, 'w:del', {'w:id': '1'})
            tag = 'w:delText'
        elif chars_state.track_changes == 'ins':
            oe1 = XML.add_tag(oe0, 'w:ins', {'w:id': '1'})
            tag = 'w:t'
        else:
            oe1 = oe0
            tag = 'w:t'
        oe2 = XML.add_tag(oe1, 'w:r')
        XML._decorate_chars(oe2, chars_state)
        res = '^([^\t\n]*)([\t\n\0])((?:.|\n)*)$'
        chars += '\0'
        while re.match(res, chars):
            rest = re.sub(res, '\\1', chars)
            char = re.sub(res, '\\2', chars)
            chars = re.sub(res, '\\3', chars)
            oe3 = XML.add_tag(oe2, tag, {'xml:space': 'preserve'}, rest)
            # oe3 = XML.add_tag(oe2, tag, {}, rest)
            if char == '\t':
                oe3 = XML.add_tag(oe2, 'w:tab', {})
            elif char == '\n':
                oe3 = XML.add_tag(oe2, 'w:br', {})
            elif char == '\0':
                pass
        return ''

    @staticmethod
    def _prepare_chars(chars):
        # REMOVE RELAX SYMBOL ("<>" -> "" / "\<\>" -> "\<\>")
        d = []
        for i in range(len(chars)):
            if re.match(NOT_ESCAPED + '<$', chars[:i]):
                if re.match('^>', chars[i:]):
                    d.append(i)
        us = list(chars)
        for i in d[::-1]:
            us.pop(i)
            us.pop(i - 1)
        chars = ''.join(us)
        # REMOVE ESCAPE SYMBOL (BACKSLASH)
        chars = re.sub('\\\\', '-\\\\', chars)
        chars = re.sub('-\\\\-\\\\', '-\\\\\\\\', chars)
        chars = re.sub('-\\\\', '', chars)
        # TRANSFORM
        # chars = chars.replace('&', '&amp;')
        # chars = chars.replace('>', '&gt;')
        # chars = chars.replace('"', '&quot;')
        # chars = chars.replace('<', '&lt;')
        # RETURN
        return chars

    @staticmethod
    def write_page_number(oe0, chars_state, char):
        oe1 = XML.add_tag(oe0, 'w:r')
        oe2 = XML._decorate_chars(oe1, chars_state)
        oe2 = XML.add_tag(oe1, 'w:fldChar', {'w:fldCharType': 'begin'})
        #
        oe1 = XML.add_tag(oe0, 'w:r')
        oe2 = XML._decorate_chars(oe1, chars_state)
        opts = {}
        # opts = {'xml:space': 'preserve'}
        if char == 'n':
            oe2 = XML.add_tag(oe1, 'w:instrText', opts, 'PAGE')
        elif char == 'N':
            # "SECTIONPAGES" IS NOT SUPPORTOD BY LIBREOFFICE
            oe2 = XML.add_tag(oe1, 'w:instrText', opts, 'SECTIONPAGES')
        elif char == 'M':
            oe2 = XML.add_tag(oe1, 'w:instrText', opts, 'NUMPAGES')
        #
        oe1 = XML.add_tag(oe0, 'w:r')
        oe2 = XML._decorate_chars(oe1, chars_state)
        oe2 = XML.add_tag(oe1, 'w:fldChar', {'w:fldCharType': 'end'})
        #
        return ''

    @staticmethod
    def _decorate_chars(oe0, chars_state):
        c_size = round(chars_state.font_size * chars_state.font_scale, 1)
        oe1 = XML.add_tag(oe0, 'w:rPr', {})
        # FONT
        if chars_state.is_preformatted:
            font = chars_state.gothic_font
        else:
            font = chars_state.mincho_font
        af, kf = XML._get_ascii_and_kanji_font(font)
        opt = {'w:ascii': af, 'w:hAnsi': af, 'w:eastAsia': kf}
        oe2 = XML.add_tag(oe1, 'w:rFonts', opt)
        # ITALIC
        if chars_state.is_italic:
            oe2 = XML.add_tag(oe1, 'w:i', {})
        # BOLD
        if chars_state.is_bold:
            oe2 = XML.add_tag(oe1, 'w:b', {})
        # STRIKETHROUGH
        if chars_state.has_strike:
            oe2 = XML.add_tag(oe1, 'w:strike', {})
        # FRAME
        if chars_state.has_frame:
            # 'w:val': 'single', 'w:sz': '4', 'w:space': '0', 'w:color': 'auto'
            oe2 = XML.add_tag(oe1, 'w:bdr', {'w:val': 'single'})
        # UNDERLINE
        if chars_state.underline is not None:
            oe2 = XML.add_tag(oe1, 'w:u', {'w:val': chars_state.underline})
        # FONT SIZE
        oe2 = XML.add_tag(oe1, 'w:sz', {'w:val': str(c_size * 2)})
        # oe2 = XML.add_tag(oe1, 'w:szCs', {'w:val': str(c_size * 2)})
        # FONT WIDTH
        if chars_state.font_width != 1.00:
            fw = round(chars_state.font_width * 100)
            if fw > 0:
                oe2 = XML.add_tag(oe1, 'w:w', {'w:val': str(fw)})
        # FONT COLOR
        if chars_state.font_color is not None:
            oe2 = XML.add_tag(oe1, 'w:color',
                              {'w:val': chars_state.font_color})
        # HIGHTLIGHT COLOR
        if chars_state.highlight_color is not None:
            opt = {'w:val': chars_state.highlight_color}
            oe2 = XML.add_tag(oe1, 'w:highlight', opt)
        # SUBSCRIPT
        if chars_state.sub_or_sup == 'sub':
            oe2 = XML.add_tag(oe1, 'w:vertAlign', {'w:val': 'subscript'})
        # SUPERSCRIPT
        if chars_state.sub_or_sup == 'sup':
            oe2 = XML.add_tag(oe1, 'w:vertAlign', {'w:val': 'superscript'})
        # SPACING
        cs_char = DEFAULT_CHAR_SPACING + chars_state.char_spacing
        cs_int = int(round(cs_char * Form.font_size * 20))
        oe2 = XML.add_tag(oe1, 'w:spacing', {'w:val': str(cs_int)})

    @staticmethod
    def set_font(style_or_run, font):
        af, kf = XML._get_ascii_and_kanji_font(font)
        style_or_run.font.name = af
        style_or_run.element.rPr.rFonts.set(ns.qn('w:eastAsia'), kf)
        # style_or_run._element.rPr.rFonts.set(ns.qn('w:eastAsia'), kf)

    @staticmethod
    def _get_ascii_and_kanji_font(font):
        fs = (font + '/').split('/')
        af = fs[0]
        af = re.sub('^\\s+', '', af)
        af = re.sub('\\s+$', '', af)
        kf = fs[1]
        kf = re.sub('^\\s+', '', kf)
        kf = re.sub('\\s+$', '', kf)
        if af == '' or af == '=':
            return kf, kf
        if kf == '' or kf == '=':
            return af, af
        return af, kf


class Math:

    """A class to write math expressions"""

    symbols = {
        '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ',
        '\\epsilon': 'ϵ', '\\zeta': 'ζ', '\\eta': 'η', '\\theta': 'θ',
        '\\iota': 'ι', '\\kappa': 'κ', '\\lambda': 'λ', '\\mu': 'μ',
        '\\nu': 'ν', '\\xi': 'ξ', '\\omicron': 'o', '\\pi': 'π',
        '\\rho': 'ρ', '\\sigma': 'σ', '\\tau': 'τ', '\\upsilon': 'υ',
        '\\phi': 'ϕ', '\\chi': 'χ', '\\psi': 'ψ', '\\omega': 'ω',
        '\\varepsilon': 'ε', '\\vartheta': 'ϑ', '\\varpi': 'ϖ',
        '\\varrho': 'ϱ', '\\varsigma': 'ς', '\\varphi': 'φ',
        '\\Alpha': 'A', '\\Beta': 'B', '\\Gamma': 'Γ', '\\Delta': 'Δ',
        '\\Epsilon': 'E', '\\Zeta': 'Z', '\\Eta': 'H', '\\Theta': 'Θ',
        '\\Iota': 'I', '\\Kappa': 'K', '\\Lambda': 'Λ', '\\Mu': 'M',
        '\\Nu': 'N', '\\Xi': 'Ξ', '\\Omicron': 'O', '\\Pi': 'Π',
        '\\Rho': 'P', '\\Sigma': 'Σ', '\\Tau': 'T', '\\Upsilon': 'Υ',
        '\\Phi': 'Φ', '\\Chi': 'X', '\\Psi': 'Ψ', '\\Omega': 'Ω',
        '\\partial': '∂',
        '\\pm': '±', '\\mp': '∓', '\\times': '×', '\\div': '÷',
        '\\cdot': '⋅',
        '\\equiv': '≡', '\\neq': '≠', '\\fallingdotseq': '≒',
        '\\geqq': '≧', '\\leqq': '≦', '\\gg': '≫', '\\ll': '≪',
        '\\in': '∈', '\\ni': '∋',
        '\\notin': '∉', '\\notni': '∌',
        '\\subset': '⊂', '\\supset': '⊃',
        '\\subseteq': '⊆', '\\supseteq': '⊇',
        '\\nsubseteq': '⊈', '\\nsupseteq': '⊉',
        '\\subsetneq': '⊊', '\\supsetneq': '⊋',
        '\\cap': '∩', '\\cup': '∪',
        '\\emptyset': '∅', '\\varnothing': '∅',
        '\\mathbb{N}': 'ℕ', '\\mathbb{Z}': 'ℤ', '\\mathbb{R}': 'ℝ',
        '\\mathbb{C}': 'ℂ', '\\mathbb{K}': '𝕂',
        '\\forall': '∀', '\\exists': '∃',
        '\\therefore': '∴', '\\because': '∵',
        '\\to': '→', '\\infty': '∞',
    }

    @classmethod
    def write_chars(cls, oe0, chars_state, chars):
        # ADD VARIABLES
        chars_state.is_italic = True
        chars_state.must_break_line = False
        # ADD MATH TAG
        oe1 = XML.add_tag(oe0, 'm:oMath')
        # PREPARE
        chars = cls._prepare_chars(chars)
        # WRITE
        is_italic = chars_state.is_italic
        chars_state.is_italic = True
        chars = cls._write_math_exp(oe1, chars_state, chars)
        chars_state.is_italic = is_italic
        # RETURN
        return chars

    @classmethod
    def _prepare_chars(cls, chars):
        # FONT WIDTH
        chars \
            = re.sub('(\\\\scalebox{' + RES_NUMBER + '})\\[1\\]', '\\1', chars)
        chars = re.sub('^\\\\\\[(.*)\\\\\\]$', '{\\1}', chars)
        chars = cls._envelop_command(chars)
        chars = cls._replace_symbol(chars)
        chars = chars.replace(' ', '')  # ' ' -> ''
        chars = cls._prepare_func(chars)
        chars = cls._close_paren(chars)
        chars = cls._envelop_all(chars)
        chars = cls._cancel_multi_paren(chars)
        return chars

    @staticmethod
    def _envelop_command(chars):
        # TEX COMMAND
        imm = ''
        res1 = NOT_ESCAPED + '(\\\\[A-Za-z]+)$'
        res9 = '^[^A-Za-z]$'
        for c in chars + '\0':
            # ALPHABET COMMAND
            if re.match(res9, c):
                imm = re.sub(res1, '\\1{\\2}', imm)
            imm = re.sub(NOT_ESCAPED + '(\\\\\\\\)$', '\\1{\\2}', imm)
            # FONT SIZE
            # imm = re.sub('{{\\\\tiny}$', '{\\\\tiny}{', imm)
            # imm = re.sub('{{\\\\scriptsize}$', '{\\\\scriptsize}{', imm)
            # imm = re.sub('{{\\\\footnotesize}$', '{\\\\footnotesize}{', imm)
            # imm = re.sub('{{\\\\small}$', '{\\\\small}{', imm)
            # imm = re.sub('{{\\\\normalsize}$', '{\\\\normalsize}{', imm)
            # imm = re.sub('{{\\\\large}$', '{\\\\large}{', imm)
            # imm = re.sub('{{\\\\Large}$', '{\\\\Large}{', imm)
            # imm = re.sub('{{\\\\LARGE}$', '{\\\\LARGE}{', imm)
            # imm = re.sub('{{\\\\huge}$', '{\\\\huge}{', imm)
            # imm = re.sub('{{\\\\Huge}$', '{\\\\Huge}{', imm)
            # FONT WIDTH
            # imm = re.sub('{{\\\\scalebox}$', '{\\\\scalebox}{', imm)
            # SPACE
            imm = re.sub('\\\\%$', '%0', imm)                   # "%"  -> "%0"
            imm = re.sub(NOT_ESCAPED + '\\\\,$', '\\1%1', imm)  # "\," -> "%1"
            imm = re.sub(NOT_ESCAPED + '\\\\:$', '\\1%2', imm)  # "\:" -> "%2"
            imm = re.sub(NOT_ESCAPED + '\\\\;$', '\\1%3', imm)  # "\;" -> "%3"
            imm = re.sub(NOT_ESCAPED + '\\\\ $', '\\1%4', imm)  # "\ " -> "%4"
            imm = re.sub(NOT_ESCAPED + '\\\\!$', '\\1%5', imm)  # "\!" -> "%5"
            # PARENTHESES
            imm = re.sub('{\\\\[Bb]igg?}', '', imm)
            imm = re.sub('{\\\\(?:left|right)}', '', imm)
            imm = re.sub('\\($', '{(-}', imm)      # "("  -> "{(-}"
            imm = re.sub('\\)$', '{-)}', imm)      # ")"  -> "{-)}"
            imm = re.sub(NOT_ESCAPED + '\\\\{$',
                         '\\1{(=}', imm)           # "\{" -> "{(=}"
            imm = re.sub(NOT_ESCAPED + '\\\\}$',
                         '\\1{=)}', imm)           # "\}" -> "{=)}"
            imm = re.sub('\\[$', '{[}', imm)       # "["  -> "{[}"
            imm = re.sub('\\]$', '{]}', imm)       # "]"  -> "{]}"
            # TEX COMMAND OPTION
            sqrt = '{\\\\sqrt}' + '{\\[}([^\\[\\]]*' \
                + ('(?:\\[[^\\[\\]]*' * 3) + ('\\][^\\[\\]]*)*' * 3) \
                + '){\\]}$'
            imm = re.sub(NOT_ESCAPED + sqrt, '\\1{\\\\sqrt}{[\\2]}', imm)
            # DEL AND INS
            imm = re.sub(NOT_ESCAPED + '\\->$', '\\1{{->}{', imm)
            imm = re.sub(NOT_ESCAPED + '<\\-$', '\\1}{<-}}', imm)
            imm = re.sub(NOT_ESCAPED + '\\+>$', '\\1{{+>}{', imm)
            imm = re.sub(NOT_ESCAPED + '<\\+$', '\\1}{<+}}', imm)
            # SUB, SUP (NO PARENTHESES)
            oc = '^([^ \\\\_\\^\\(\\){}\\[\\]\0])$'
            if re.match(NOT_ESCAPED + '(_|\\^)$', imm + c):
                if imm[-1] != '}':
                    if re.match(oc, imm[-1]):
                        imm = re.sub('(.)$', '{\\1}', imm)
                    else:
                        imm += '{}'
            if re.match(NOT_ESCAPED + '(_|\\^)$', imm):
                if c != '{':
                    if re.match(oc, c):
                        imm += '{' + c + '}'
                        c = ''
                    else:
                        imm += '{}'
            # ADD CHAR
            if c != '\0':
                imm += c
        chars = imm
        return chars

    @staticmethod
    def _replace_symbol(chars):
        for com in Math.symbols:
            chars = re.sub('{\\' + com + '}', Math.symbols[com], chars)
        return chars

    @classmethod
    def _prepare_func(cls, chars):
        imm = ''
        for c in chars + '\0':
            nubs = cls.__get_nubs(imm)
            tmps = []
            while tmps != nubs:
                tmps = []
                for n in nubs:
                    tmps.append(n)
                # CONTINUE
                res = '^.*{\\\\(' \
                    + 'sum|prod|int|iint|iiint|oint|sin|cos|tan|log|lim' \
                    + ')}+$'
                if (len(nubs) >= 3) and re.match(res, nubs[-3]):
                    continue
                res = '^.*{\\\\(' \
                    + 'sum|prod|int|iint|iiint|oint' \
                    + ')}+$'
                if (len(nubs) >= 5) and re.match(res, nubs[-5]):
                    continue
                # CONBINATION, PERMUTATION
                if (len(nubs) >= 4) and \
                   re.match('^{{}(_{.*})}$', nubs[-4]) and nubs[-2] == '_':
                    nubs[-4] = re.sub('^{{}(_{.*})}$', '\\1', nubs[-4])
                    nubs[-4], nubs[-3] = nubs[-3], nubs[-4]
                    nubs[-4], nubs[-1] = cls._close_func(nubs[-4], nubs[-1])
                # SUBSCRIPT, SUPERSCRIPT
                elif (c != '_') and (c != '^'):
                    if (len(nubs) >= 5) and \
                       ((nubs[-4] == '_') and (nubs[-2] == '^')):
                        nubs[-5], nubs[-1] \
                            = cls._close_func(nubs[-5], nubs[-1])
                    elif (len(nubs) >= 5) and \
                         ((nubs[-4] == '^') and (nubs[-2] == '_')):
                        nubs[-4], nubs[-2] = nubs[-2], nubs[-4]
                        nubs[-3], nubs[-1] = nubs[-1], nubs[-3]
                        nubs[-5], nubs[-1] \
                            = cls._close_func(nubs[-5], nubs[-1])
                    elif (len(nubs) >= 3) and \
                         ((nubs[-2] == '_') or (nubs[-2] == '^')):
                        nubs[-3], nubs[-1] \
                            = cls._close_func(nubs[-3], nubs[-1])
                # LINEBREAK, MATHRM, MATHBF, STRIKE, FRAME, UNDERLINE, EXP, VEC
                res = '^{\\\\(?:' \
                    + '\\\\|mathrm|mathbf|sout|boxed|underline|exp|vec' \
                    + ')}$'
                if (len(nubs) >= 2) and re.match(res, nubs[-2]):
                    nubs[-2], nubs[-1] = cls._close_func(nubs[-2], nubs[-1])
                # TEXTCOLOR, COLORBOX, FRACTION, BINOMIAL
                res = '^{\\\\(?:textcolor|colorbox|frac|binom)}$'
                if (len(nubs) >= 3) and re.match(res, nubs[-3]):
                    nubs[-3], nubs[-1] = cls._close_func(nubs[-3], nubs[-1])
                # SQRT
                if (len(nubs) >= 2) and (nubs[-2] == '{\\sqrt}'):
                    if not re.match('{\\[.*\\]}', nubs[-1]):
                        nubs.insert(-1, '{[]}')
                if (len(nubs) >= 3) and (nubs[-3] == '{\\sqrt}'):
                    nubs[-3], nubs[-1] = cls._close_func(nubs[-3], nubs[-1])
                # SIN, COS, TAN
                res = '^.*{\\\\(?:sin|cos|tan)}+$'
                if (len(nubs) >= 2) and re.match(res, nubs[-2]):
                    if nubs[-1] != '^':
                        nubs.insert(-1, '^')
                        nubs.insert(-1, '{}')
                if (len(nubs) >= 4) and re.match(res, nubs[-4]):
                    nubs[-4], nubs[-1] = cls._close_func(nubs[-4], nubs[-1])
                # LOG, LIMIT
                if (len(nubs) >= 2) and \
                   re.match('^{\\\\(?:log|lim)}$', nubs[-2]):
                    if nubs[-1] != '_':
                        nubs.insert(-1, '_')
                        nubs.insert(-1, '{}')
                if (len(nubs) >= 4) and \
                   re.match('^{\\\\(?:log|lim)}$', nubs[-4]):
                    nubs[-4], nubs[-1] = cls._close_func(nubs[-4], nubs[-1])
                # SIGMA, PI, INTEGRAL, LINE INTEGRAL
                if (len(nubs) >= 2) and \
                   re.match('^{\\\\(?:sum|prod|(?:|i|ii|o)int)}$', nubs[-2]):
                    if nubs[-1] != '_':
                        nubs.insert(-1, '_')
                        nubs.insert(-1, '{}')
                if (len(nubs) >= 4) and \
                   re.match('^{\\\\(?:sum|prod|(?:|i|ii|o)int)}$', nubs[-4]):
                    if nubs[-1] != '^':
                        nubs.insert(-1, '^')
                        nubs.insert(-1, '{}')
                if (len(nubs) >= 6) and \
                   re.match('^{\\\\(?:sum|prod|(?:|i|ii|o)int)}$', nubs[-6]):
                    nubs[-6], nubs[-1] = cls._close_func(nubs[-6], nubs[-1])
                # MATRIX
                if '{\\Ybmx}' in nubs:
                    if (len(nubs) >= 1) and (nubs[-1] == '{\\\\}'):
                        nubs[-1] = '{\\Ylmx}'
                if (len(nubs) >= 2) and \
                   nubs[-2] == '{\\begin}' and \
                   re.match('^{.*matrix}$', nubs[-1]):
                    nubs[-2] = '{\\Ybmx}'
                if (len(nubs) >= 2) and \
                   nubs[-2] == '{\\end}' and \
                   re.match('^{.*matrix}$', nubs[-1]):
                    b = None
                    for i, n in enumerate(nubs):
                        if n == '{\\Ybmx}':
                            b = i
                    if b is not None:
                        nubs[b] = '{{\\Xbmx}'
                        nubs[b + 1] += '{'
                        nubs[-1] = '}{\\Xemx}}'
                        s = ''
                        for i in range(b + 2, len(nubs) - 2):
                            if nubs[i] == '&':
                                nubs[i] = '}{'
                            if nubs[i] == '{\\Ylmx}':
                                nubs[i] = '}{\\Xlmx}{'
                            s += nubs[i]
                            nubs[i] = ''
                        nubs[-2] = s
                        if re.match('^.*{$', nubs[-2]) and \
                           re.match('^}.*$', nubs[-1]):
                            nubs[-2] = re.sub('{$', '', nubs[-2])
                            nubs[-1] = re.sub('^}', '', nubs[-1])
                # FONT SIZE
                res = '^{\\\\(?:' \
                    + 'tiny|scriptsize|footnotesize|small|' \
                    + 'normalsize|large|Large|LARGE|huge|Huge' \
                    + ')}$'
                if (len(nubs) >= 2) and re.match(res, nubs[-2]):
                    # nubs[-2], nubs[-1] = cls._close_func(nubs[-2], nubs[-1])
                    pass
                # FONT WIDTH
                if (len(nubs) >= 3) \
                   and re.match('{\\\\scalebox}', nubs[-3]) \
                   and re.match('{' + RES_NUMBER + '}', nubs[-2]):
                    # nubs[-3], nubs[-1] = cls._close_func(nubs[-3], nubs[-1])
                    pass
                # PARENTHESES
                if (len(nubs) >= 1) and (nubs[-1] == '{-)}'):
                    for i in range(len(nubs) - 1, -1, -1):
                        if nubs[i] == '{(-}':
                            nubs[i], nubs[-1] \
                                = cls._close_func(nubs[i], nubs[-1])
                if (len(nubs) >= 1) and (nubs[-1] == '{=)}'):
                    for i in range(len(nubs) - 1, -1, -1):
                        if nubs[i] == '{(=}':
                            nubs[i], nubs[-1] \
                                = cls._close_func(nubs[i], nubs[-1])
                if (len(nubs) >= 1) and (nubs[-1] == '{]}'):
                    for i in range(len(nubs) - 1, -1, -1):
                        if nubs[i] == '{[}':
                            nubs[i], nubs[-1] \
                                = cls._close_func(nubs[i], nubs[-1])
                # REMAKE
                imm = ''.join(nubs)
                nubs = cls.__get_nubs(imm)
            if c != '\0':
                imm += c
        chars = imm
        return chars

    @staticmethod
    def _close_func(beg_str, end_str):
        oc = '^([^ \\\\_\\^\\(\\){}\\[\\]\0])$'
        beg_str = '{' + beg_str
        if not re.match('^{.*}$', end_str):
            if re.match(oc, end_str):
                end_str = '{' + end_str + '}}'
            else:
                end_str = '{}}' + end_str
        else:
            end_str = end_str + '}'
        return beg_str, end_str

    @staticmethod
    def __get_nubs(imm):
        nubs = []
        nub = ''
        dep = 0
        for n, c in enumerate(imm[::-1] + '\0'):
            if c == '{':
                dep -= 1
            if c == '}':
                dep += 1
            if c != '\0':
                nub = c + nub
            if nub != '' and (dep == 0 or c == '\0'):
                while re.match('^{{(.*)}}$', nub):
                    tmp = re.sub('^{{(.*)}}$', '{\\1}', nub)
                    td = 0
                    ta = 0
                    for tc in tmp:
                        if tc == '{':
                            td -= 1
                        if tc == '}':
                            td += 1
                        if td >= 0:
                            ta += 1
                        if ta > 1:
                            break
                    if ta == 1:
                        nub = tmp
                    else:
                        break
                nubs.append(nub)
                nub = ''
        return nubs[::-1]

    @staticmethod
    def _close_paren(chars):
        d = 0
        imm = ''
        for c in chars:
            imm += c
            if re.match(NOT_ESCAPED + '{$', imm):
                d += 1
            if re.match(NOT_ESCAPED + '}$', imm):
                d -= 1
        if d > 0:
            chars = chars + ('}' * d)
        if d < 0:
            chars = ('{' * (d * -1)) + chars
        return chars

    @staticmethod
    def _envelop_all(chars):
        tmp = ''
        while tmp != chars:
            tmp = chars
            chars = re.sub('{([^{}]+){', '{{\\1}{', chars)
            chars = re.sub('}([^{}]+)}', '}{\\1}}', chars)
            chars = re.sub('}([^{}]+){', '}{\\1}{', chars)
        return chars

    @staticmethod
    def _cancel_multi_paren(chars):
        rm = []
        for i in range(len(chars) - 1):
            if chars[i] != '{' or chars[i + 1] != '{':
                continue
            dep = [0]
            d = 0
            for j in range(i, len(chars)):
                if chars[j] == '{':
                    d += 1
                if chars[j] == '}':
                    d -= 1
                dep.append(d)
                if d == 0:
                    if chars[j - 1] == '}' or chars[j] == '}':
                        dep.pop(0)
                        dep.pop(0)
                        dep.pop(-1)
                        dep.pop(-1)
                        if 1 not in dep:
                            rm.append(i)
                            rm.append(j)
                    break
        rm.sort()
        rm.reverse()
        u = list(chars)
        for r in rm:
            u.pop(r)
        chars = ''.join(u)
        return chars

    @classmethod
    def _write_math_exp(cls, oe0, chars_state, chars):
        # REMOVE ENCLOSING PARENTHESIS
        if re.match('^{(.*)}$', chars):
            ers = re.sub('^{(.*)}$', '\\1', chars)
            imm = ''
            dep = 0
            for c in ers:
                imm += c
                if re.match(NOT_ESCAPED + '{$', imm):
                    dep += 1
                if re.match(NOT_ESCAPED + '}$', imm):
                    dep -= 1
                if dep < 0:
                    break
            else:
                chars = ers
        tmp = ''
        # ONE NUB
        if re.match('^[^{}]+$', chars):
            cls.__write_nub(oe0, chars_state, chars)
            return ''
        nubs = cls.__get_nubs(chars)
        # FUNCITON
        if False:
            pass
        # INTEGRAL
        elif len(nubs) == 6 and nubs[0] == '{\\int}':
            cls._write_int(oe0, chars_state, '', nubs[2], nubs[4], nubs[5])
        # DOUBLE INTEGRAL
        elif len(nubs) == 6 and nubs[0] == '{\\iint}':
            cls._write_int(oe0, chars_state, '∬', nubs[2], nubs[4], nubs[5])
        # TRIPLE INTEGRAL
        elif len(nubs) == 6 and nubs[0] == '{\\iiint}':
            cls._write_int(oe0, chars_state, '∭', nubs[2], nubs[4], nubs[5])
        # LINE INTEGRAL
        elif len(nubs) == 6 and nubs[0] == '{\\oint}':
            cls._write_int(oe0, chars_state, '∮', nubs[2], nubs[4], nubs[5])
        # SIGMA
        elif len(nubs) == 6 and nubs[0] == '{\\sum}':
            cls._write_sop(oe0, chars_state, '∑', nubs[2], nubs[4], nubs[5])
        # PI
        elif len(nubs) == 6 and nubs[0] == '{\\prod}':
            cls._write_sop(oe0, chars_state, '∏', nubs[2], nubs[4], nubs[5])
        # SUB AND SUP
        elif len(nubs) == 5 and nubs[1] == '{_}' and nubs[3] == '{^}':
            cls._write_bap(oe0, chars_state, nubs[0], nubs[2], nubs[4])
        # CONBINATION AND PERMUTATION
        elif len(nubs) == 5 and nubs[1] == '{_}' and nubs[3] == '{_}':
            cls._write_cop(oe0, chars_state, nubs[0], nubs[2], nubs[4])
        # LOG
        elif len(nubs) == 4 and nubs[0] == '{\\log}':
            if nubs[2] == '{}':
                cls._write_one(oe0, chars_state, 'log', nubs[3])
            else:
                cls._write_two(oe0, chars_state, 'log',
                               nubs[1], nubs[2], nubs[3])
        # LIMIT
        elif len(nubs) == 4 and nubs[0] == '{\\lim}':
            cls._write_lim(oe0, chars_state, nubs[2], nubs[3])
        # SIN
        elif len(nubs) == 4 and nubs[0] == '{\\sin}':
            if nubs[2] == '{}':
                cls._write_one(oe0, chars_state, 'sin', nubs[3])
            else:
                cls._write_two(oe0, chars_state, 'sin',
                               nubs[1], nubs[2], nubs[3])
        # COS
        elif len(nubs) == 4 and nubs[0] == '{\\cos}':
            if nubs[2] == '{}':
                cls._write_one(oe0, chars_state, 'cos', nubs[3])
            else:
                cls._write_two(oe0, chars_state, 'cos',
                               nubs[1], nubs[2], nubs[3])
        # TAN
        elif len(nubs) == 4 and nubs[0] == '{\\tan}':
            if nubs[2] == '{}':
                cls._write_one(oe0, chars_state, 'tan', nubs[3])
            else:
                cls._write_two(oe0, chars_state, 'tan',
                               nubs[1], nubs[2], nubs[3])
        # SUB AND SUP
        elif len(nubs) == 3 and (nubs[1] == '{_}' or nubs[1] == '{^}'):
            cls._write_bop(oe0, chars_state, nubs[1], nubs[0], nubs[2])
        # FRACTION
        elif len(nubs) == 3 and nubs[0] == '{\\frac}':
            cls._write_fra(oe0, chars_state, nubs[1], nubs[2])
        # BINOMIAL
        elif len(nubs) == 3 and nubs[0] == '{\\binom}':
            cls._write_bin(oe0, chars_state, nubs[1], nubs[2])
        # RADICAL ROOT
        elif len(nubs) == 3 and nubs[0] == '{\\sqrt}':
            t = re.sub('^{\\[(.*)\\]}$', '\\1', nubs[1])
            cls._write_rrt(oe0, chars_state, t, nubs[2])
        # LIMIT
        elif len(nubs) == 3 and nubs[0] == '{\\lim}':
            cls._write_lim(oe0, chars_state, nubs[1], nubs[2])
        # EXPONENTIAL
        elif len(nubs) == 2 and nubs[0] == '{\\exp}':
            cls._write_one(oe0, chars_state, 'exp', nubs[1])
        # VECTOR
        elif len(nubs) == 2 and nubs[0] == '{\\vec}':
            cls._write_vec(oe0, chars_state, nubs[1])
        # MATRIX
        elif (len(nubs) >= 2 and
              nubs[0] == '{\\Xbmx}' and nubs[-1] == '{\\Xemx}'):
            c = nubs[1]
            nubs.pop(0)
            nubs.pop(0)
            nubs.pop(-1)
            cls._write_mtx(oe0, chars_state, c, nubs)
        # S PAREN
        elif len(nubs) >= 2 and nubs[0] == '{(-}' and nubs[-1] == '{-)}':
            t = re.sub('{\\(-}(.*){-\\)}', '\\1', chars)
            cls._write_prn(oe0, chars_state, '()', '{' + t + '}')
        # M PAREN
        elif len(nubs) >= 2 and nubs[0] == '{(=}' and nubs[-1] == '{=)}':
            t = re.sub('{\\(=}(.*){=\\)}', '\\1', chars)
            cls._write_prn(oe0, chars_state, '{}', '{' + t + '}')
        # L PAREN
        elif len(nubs) >= 2 and nubs[0] == '{[}' and nubs[-1] == '{]}':
            t = re.sub('{\\[}(.*){\\]}', '\\1', chars)
            cls._write_prn(oe0, chars_state, '[]', '{' + t + '}')
        # LINE BREAK
        elif len(nubs) == 2 and nubs[0] == '{\\\\}':
            chars_state.must_break_line = True
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.must_break_line = False
        # FONT SCALE
        elif len(nubs) == 2 and nubs[0] == '{\\tiny}':
            chars_state.font_scale = 0.2
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.font_scale = 1.0
        elif len(nubs) == 2 and nubs[0] == '{\\scriptsize}':
            chars_state.font_scale = 0.4
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.font_scale = 1.0
        elif len(nubs) == 2 and nubs[0] == '{\\footnotesize}':
            chars_state.font_scale = 0.6
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.font_scale = 1.0
        elif len(nubs) == 2 and nubs[0] == '{\\small}':
            chars_state.font_scale = 0.8
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.font_scale = 1.0
        elif len(nubs) == 2 and nubs[0] == '{\\normalsize}':
            chars_state.font_scale = 1.0
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.font_scale = 1.0
        elif len(nubs) == 2 and nubs[0] == '{\\large}':
            chars_state.font_scale = 1.2
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.font_scale = 1.0
        elif len(nubs) == 2 and nubs[0] == '{\\Large}':
            chars_state.font_scale = 1.4
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.font_scale = 1.0
        elif len(nubs) == 2 and nubs[0] == '{\\LARGE}':
            chars_state.font_scale = 1.6
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.font_scale = 1.0
        elif len(nubs) == 2 and nubs[0] == '{\\huge}':
            chars_state.font_scale = 1.8
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.font_scale = 1.0
        elif len(nubs) == 2 and nubs[0] == '{\\Huge}':
            chars_state.font_scale = 2.0
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.font_scale = 1.0
        # FONT WIDTH
        elif len(nubs) == 3 and nubs[0] == '{\\scalebox}':
            if re.match('{(' + RES_NUMBER + ')}', nubs[1]):
                wid = re.sub('{(' + RES_NUMBER + ')}', '\\1', nubs[1])
                chars_state.font_width = float(wid)
                cls._write_math_exp(oe0, chars_state, nubs[2])
                chars_state.font_width = 1.0
        # ROMAN
        elif len(nubs) == 2 and nubs[0] == '{\\mathrm}':
            chars_state.is_italic = False
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.is_italic = True
        # BOLD
        elif len(nubs) == 2 and nubs[0] == '{\\mathbf}':
            chars_state.is_bold = True
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.is_bold = False
        # STRIKETHROUGH
        elif len(nubs) == 2 and nubs[0] == '{\\sout}':
            chars_state.has_strike = True
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.has_strike = False
        # FRAME
        elif len(nubs) == 2 and nubs[0] == '{\\boxed}':
            chars_state.has_frame = True
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.has_frame = False
        # UNDERLINE
        elif len(nubs) == 2 and nubs[0] == '{\\underline}':
            chars_state.underline = 'single'
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.underline = None
        # FONT COLOR
        elif len(nubs) == 3 and nubs[0] == '{\\textcolor}':
            chars_state.font_color = re.sub('^{(.*)}$', '\\1', nubs[1])
            cls._write_math_exp(oe0, chars_state, nubs[2])
            chars_state.font_color = None
        # HIGHLIGHT COLOR
        elif len(nubs) == 3 and nubs[0] == '{\\colorbox}':
            chars_state.highlight_color = re.sub('^{(.*)}$', '\\1', nubs[1])
            cls._write_math_exp(oe0, chars_state, nubs[2])
            chars_state.highlight_color = None
        # TRACK CHANGES
        elif len(nubs) >= 3 and nubs[0] == '{->}' and nubs[2] == '{<-}':
            chars_state.track_changes = 'del'
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.track_changes = ''
        elif len(nubs) >= 3 and nubs[0] == '{+>}' and nubs[2] == '{<+}':
            chars_state.track_changes = 'ins'
            cls._write_math_exp(oe0, chars_state, nubs[1])
            chars_state.track_changes = ''
        # ERROR
        elif (len(nubs) == 1) and (not re.match('^{.*}$', nubs[0])):
            cls.__write_nub(oe0, chars_state, chars)
        # RECURSION
        else:
            for n in nubs:
                cls._write_math_exp(oe0, chars_state, n)
        return ''

    @classmethod
    def __write_nub(cls, oe0, chars_state, nub):
        if nub == '':
            return
        nub = re.sub('%9', '  ', nub)
        nub = re.sub('%3', ' ', nub)
        nub = re.sub('%2', ' ', nub)
        nub = re.sub('%1', ' ', nub)
        nub = re.sub('%0', '%', nub)
        oe1 = XML.add_tag(oe0, 'm:r', {})
        if chars_state.track_changes == 'del':
            oe2 = XML.add_tag(oe1, 'w:del', {})
        elif chars_state.track_changes == 'ins':
            oe2 = XML.add_tag(oe1, 'w:ins', {})
        else:
            oe2 = oe1
        cls.__decorate_nub(oe2, chars_state)
        oe3 = XML.add_tag(oe2, 'm:t', {}, nub)

    @classmethod
    def __decorate_nub(cls, oe0, chars_state):
        cls.__decorate_nub_m(oe0, chars_state)
        cls.__decorate_nub_w(oe0, chars_state)

    @staticmethod
    def __decorate_nub_m(oe0, chars_state):
        oe1 = XML.add_tag(oe0, 'm:rPr', {})
        # LINE BREAK
        if chars_state.must_break_line:
            oe2 = XML.add_tag(oe1, 'm:brk', {'m:alnAt': '1'})
        # ROMAN AND BOLD
        if chars_state.is_italic and chars_state.is_bold:
            oe2 = XML.add_tag(oe1, 'm:sty', {'m:val': 'bi'})
        elif chars_state.is_bold:
            oe2 = XML.add_tag(oe1, 'm:sty', {'m:val': 'b'})
        elif not chars_state.is_italic:
            oe2 = XML.add_tag(oe1, 'm:sty', {'m:val': 'p'})

    @staticmethod
    def __decorate_nub_w(oe0, chars_state):
        c_size = round(chars_state.font_size * chars_state.font_scale, 1)
        oe1 = XML.add_tag(oe0, 'w:rPr', {})
        # (FONT, ITALIC, BOLD)
        # STRIKETHROUGH (NOT SUPPORTOD BY LIBREOFFICE)
        if chars_state.has_strike:
            oe2 = XML.add_tag(oe1, 'w:strike', {})
        # FRAME (NOT SUPPORTOD BY LIBREOFFICE)
        if chars_state.has_frame:
            # 'w:val': 'single', 'w:sz': '4', 'w:space': '0', 'w:color': 'auto'
            oe2 = XML.add_tag(oe1, 'w:bdr', {'w:val': 'single'})
        # UNDERLINE
        if chars_state.underline is not None:
            oe2 = XML.add_tag(oe1, 'w:u', {'w:val': chars_state.underline})
        # FONT SIZE
        oe2 = XML.add_tag(oe1, 'w:sz', {'w:val': str(c_size * 2)})
        # oe2 = XML.add_tag(oe1, 'w:szCs', {'w:val': str(c_size * 2)})
        # FONT WIDTH
        if chars_state.font_width != 1.00:
            opt = {'w:val': str(int(chars_state.font_width * 100))}
            oe2 = XML.add_tag(oe1, 'w:w', opt)
        # FONT COLOR
        if chars_state.font_color is not None:
            oe2 = XML.add_tag(oe1, 'w:color',
                              {'w:val': chars_state.font_color})
        # HIGHTLIGHT COLOR
        if chars_state.highlight_color is not None:
            opt = {'w:val': chars_state.highlight_color}
            oe2 = XML.add_tag(oe1, 'w:highlight', opt)
        # (SUBSCRIPT, SUPERSCRIPT)

    # INTEGRAL
    @classmethod
    def _write_int(cls, oe0, chars_state, c, t1, t2, t3):
        oe1 = XML.add_tag(oe0, 'm:nary', {})
        oe2 = XML.add_tag(oe1, 'm:naryPr', {})
        if c != '':
            oe3 = XML.add_tag(oe2, 'm:chr', {'m:val': c})
        oe3 = XML.add_tag(oe2, 'm:limLoc', {'m:val': 'subSup'})
        if t1 == '' or t1 == '{}':
            oe3 = XML.add_tag(oe2, 'm:subHide', {'m:val': '1'})
        if t2 == '' or t2 == '{}':
            oe3 = XML.add_tag(oe2, 'm:supHide', {'m:val': '1'})
        #
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe3, chars_state)
        #
        oe2 = XML.add_tag(oe1, 'm:sub', {})
        if not (t1 == '' or t1 == '{}'):
            cls._write_math_exp(oe2, chars_state, t1)
        oe2 = XML.add_tag(oe1, 'm:sup', {})
        if not (t2 == '' or t2 == '{}'):
            cls._write_math_exp(oe2, chars_state, t2)
        oe2 = XML.add_tag(oe1, 'm:e', {})
        cls._write_math_exp(oe2, chars_state, t3)

    # SIGMA, PI
    @classmethod
    def _write_sop(cls, oe0, chars_state, c, t1, t2, t3):
        oe1 = XML.add_tag(oe0, 'm:nary', {})
        oe2 = XML.add_tag(oe1, 'm:naryPr', {})
        oe3 = XML.add_tag(oe2, 'm:chr', {'m:val': c})
        oe3 = XML.add_tag(oe2, 'm:limLoc', {'m:val': 'undOvr'})
        if t1 == '' or t1 == '{}':
            oe3 = XML.add_tag(oe2, 'm:subHide', {'m:val': '1'})
        if t2 == '' or t2 == '{}':
            oe3 = XML.add_tag(oe2, 'm:supHide', {'m:val': '1'})
        #
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe3, chars_state)
        #
        oe2 = XML.add_tag(oe1, 'm:sub', {})
        if not (t1 == '' or t1 == '{}'):
            cls._write_math_exp(oe2, chars_state, t1)
        oe2 = XML.add_tag(oe1, 'm:sup', {})
        if not (t2 == '' or t2 == '{}'):
            cls._write_math_exp(oe2, chars_state, t2)
        oe2 = XML.add_tag(oe1, 'm:e', {})
        cls._write_math_exp(oe2, chars_state, t3)

    # SUB AND SUP
    @classmethod
    def _write_bap(cls, oe0, chars_state, t1, t2, t3):
        oe1 = XML.add_tag(oe0, 'm:sSubSup', {})
        #
        oe2 = XML.add_tag(oe1, 'm:sSubSupPr', {})
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe3, chars_state)
        #
        oe2 = XML.add_tag(oe1, 'm:e', {})
        cls._write_math_exp(oe2, chars_state, t1)
        oe2 = XML.add_tag(oe1, 'm:sub', {})
        cls._write_math_exp(oe2, chars_state, t2)
        oe2 = XML.add_tag(oe1, 'm:sup', {})
        cls._write_math_exp(oe2, chars_state, t3)

    # CONBINATION, PERMUTATION
    @classmethod
    def _write_cop(cls, oe0, chars_state, t1, t2, t3):
        oe1 = XML.add_tag(oe0, 'm:sPre', {})
        #
        oe2 = XML.add_tag(oe1, 'm:sPrePr', {})
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe3, chars_state)
        #
        oe2 = XML.add_tag(oe1, 'm:sub', {})
        cls._write_math_exp(oe2, chars_state, t2)
        oe2 = XML.add_tag(oe1, 'm:sup', {})
        cls._write_math_exp(oe2, chars_state, '{}')
        oe2 = XML.add_tag(oe1, 'm:e', {})
        oe3 = XML.add_tag(oe2, 'm:sSub', {})
        #
        oe4 = XML.add_tag(oe3, 'm:sSubPr', {})
        oe5 = XML.add_tag(oe4, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe5, chars_state)
        #
        oe4 = XML.add_tag(oe3, 'm:e', {})
        cls._write_math_exp(oe4, chars_state, t1)
        oe4 = XML.add_tag(oe3, 'm:sub', {})
        cls._write_math_exp(oe4, chars_state, t3)

    # TWO ARGUMENTS FUNCTION
    @classmethod
    def _write_two(cls, oe0, chars_state, c, s, t1, t2):
        # \sin^2{x}, \log_2{x}
        oe1 = XML.add_tag(oe0, 'm:func', {})
        #
        oe2 = XML.add_tag(oe1, 'm:funcPr', {})
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe3, chars_state)
        #
        oe2 = XML.add_tag(oe1, 'm:fName', {})
        if s == '_' or s == '{_}':
            oe3 = XML.add_tag(oe2, 'm:sSub', {})
        else:
            oe3 = XML.add_tag(oe2, 'm:sSup', {})
        oe4 = XML.add_tag(oe3, 'm:e', {})
        oe5 = XML.add_tag(oe4, 'm:r', {})
        #
        cls.__decorate_nub_m(oe5, chars_state)
        oe6 = XML.add_tag(oe5, 'm:t', {}, c)
        #
        if s == '_' or s == '{_}':
            oe4 = XML.add_tag(oe3, 'm:sub', {})
        else:
            oe4 = XML.add_tag(oe3, 'm:sup', {})
        cls._write_math_exp(oe4, chars_state, t1)
        #
        oe2 = XML.add_tag(oe1, 'm:e', {})
        cls._write_math_exp(oe2, chars_state, t2)

    # SUBSCRIPT OR SUPERSCRIPT
    @classmethod
    def _write_bop(cls, oe0, chars_state, s, t1, t2):
        # x_i, x^2
        if s == '_' or s == '{_}':
            oe1 = XML.add_tag(oe0, 'm:sSub', {})
            oe2 = XML.add_tag(oe1, 'm:sSubPr', {})
        else:
            oe1 = XML.add_tag(oe0, 'm:sSup', {})
            oe2 = XML.add_tag(oe1, 'm:sSupPr', {})
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe3, chars_state)
        oe2 = XML.add_tag(oe1, 'm:e', {})
        cls._write_math_exp(oe2, chars_state, t1)
        if s == '_' or s == '{_}':
            oe2 = XML.add_tag(oe1, 'm:sub', {})
        else:
            oe2 = XML.add_tag(oe1, 'm:sup', {})
        cls._write_math_exp(oe2, chars_state, t2)

    # FRACTION
    @classmethod
    def _write_fra(cls, oe0, chars_state, t1, t2):
        # \frac{2}{3}
        oe1 = XML.add_tag(oe0, 'm:f', {})
        #
        oe2 = XML.add_tag(oe1, 'm:fPr', {})
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe3, chars_state)
        #
        oe2 = XML.add_tag(oe1, 'm:num', {})
        cls._write_math_exp(oe2, chars_state, t1)
        oe2 = XML.add_tag(oe1, 'm:den', {})
        cls._write_math_exp(oe2, chars_state, t2)

    # BINOMIAL
    @classmethod
    def _write_bin(cls, oe0, chars_state, t1, t2):
        # \binom{2}{3}
        oe1 = XML.add_tag(oe0, 'm:d', {})
        oe2 = XML.add_tag(oe1, 'm:dPr', {})
        #
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub(oe3, chars_state)
        #
        oe2 = XML.add_tag(oe1, 'm:e', {})
        oe3 = XML.add_tag(oe2, 'm:f', {})
        #
        oe4 = XML.add_tag(oe3, 'm:fPr', {})
        oe5 = XML.add_tag(oe4, 'm:type', {'m:val': 'noBar'})
        oe5 = XML.add_tag(oe4, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe5, chars_state)
        #
        oe4 = XML.add_tag(oe3, 'm:num', {})
        cls._write_math_exp(oe4, chars_state, t1)
        oe4 = XML.add_tag(oe3, 'm:den', {})
        cls._write_math_exp(oe4, chars_state, t2)

    # RADICAL ROOT
    @classmethod
    def _write_rrt(cls, oe0, chars_state, t1, t2):
        # \sqrt[3]{2}
        oe1 = XML.add_tag(oe0, 'm:rad', {})
        #
        oe2 = XML.add_tag(oe1, 'm:radPr', {})
        if t1 == '' or t1 == '{}':
            oe3 = XML.add_tag(oe2, 'm:degHide', {'m:val': '1'})
        #
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe3, chars_state)
        #
        oe2 = XML.add_tag(oe1, 'm:deg', {})
        cls._write_math_exp(oe2, chars_state, t1)
        oe2 = XML.add_tag(oe1, 'm:e', {})
        cls._write_math_exp(oe2, chars_state, t2)

    # LIMIT
    @classmethod
    def _write_lim(cls, oe0, chars_state, t1, t2):
        # \lim_{x}{y}
        oe1 = XML.add_tag(oe0, 'm:func', {})
        #
        oe2 = XML.add_tag(oe1, 'm:funcPr', {})
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe3, chars_state)
        #
        oe2 = XML.add_tag(oe1, 'm:fName', {})
        oe3 = XML.add_tag(oe2, 'm:limLow', {})
        #
        oe4 = XML.add_tag(oe3, 'm:limLowPr', {})
        oe5 = XML.add_tag(oe4, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe5, chars_state)
        #
        oe4 = XML.add_tag(oe3, 'm:e', {})
        oe5 = XML.add_tag(oe4, 'm:r', {})
        #
        cls.__decorate_nub_m(oe5, chars_state)
        oe6 = XML.add_tag(oe5, 'm:t', {}, 'lim')
        oe4 = XML.add_tag(oe3, 'm:lim', {})
        cls._write_math_exp(oe4, chars_state, t1)
        oe2 = XML.add_tag(oe1, 'm:e', {})
        cls._write_math_exp(oe2, chars_state, t2)

    # ONE ARGUMENT FUNCTION
    @classmethod
    def _write_one(cls, oe0, chars_state, c, t1):
        # \sin{x}, \exp{y}
        oe1 = XML.add_tag(oe0, 'm:func', {})
        #
        oe2 = XML.add_tag(oe1, 'm:funcPr', {})
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub_w(oe3, chars_state)
        #
        oe2 = XML.add_tag(oe1, 'm:fName', {})
        oe3 = XML.add_tag(oe2, 'm:r', {})
        #
        cls.__decorate_nub_m(oe3, chars_state)
        oe4 = XML.add_tag(oe3, 'm:t', {}, c)
        oe2 = XML.add_tag(oe1, 'm:e', {})
        cls._write_math_exp(oe2, chars_state, t1)

    # VECTOR
    @classmethod
    def _write_vec(cls, oe0, chars_state, t1):
        # \vec{x}
        oe1 = XML.add_tag(oe0, 'm:acc', {})
        #
        oe2 = XML.add_tag(oe1, 'm:accPr', {})
        oe3 = XML.add_tag(oe2, 'm:chr', {'m:val': '⃗'})
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub(oe3, chars_state)
        #
        oe2 = XML.add_tag(oe1, 'm:e', {})
        cls._write_math_exp(oe2, chars_state, t1)

    # MATRIX
    @classmethod
    def _write_mtx(cls, oe0, chars_state, c, t1):
        # \begin{pmatrix}a&b\\c&d\end{pmatrix}
        nubs = t1
        nubs.append('{\\Xlmx}')
        mtrx = []
        row = []
        for cel in nubs:
            cel = re.sub('^{(.*)}$', '\\1', cel)
            if cel != '\\Xlmx':
                row.append(cel)
            else:
                mtrx.append(row)
                row = []
        #
        oe1 = XML.add_tag(oe0, 'm:d', {})
        #
        oe2 = XML.add_tag(oe1, 'm:dPr', {})
        if c == '{pmatrix}':
            oe3 = XML.add_tag(oe2, 'm:begChr', {'m:val': '('})
            oe3 = XML.add_tag(oe2, 'm:endChr', {'m:val': ')'})
        elif c == '{bmatrix}':
            oe3 = XML.add_tag(oe2, 'm:begChr', {'m:val': '['})
            oe3 = XML.add_tag(oe2, 'm:endChr', {'m:val': ']'})
        elif c == '{vmatrix}':
            oe3 = XML.add_tag(oe2, 'm:begChr', {'m:val': '|'})
            oe3 = XML.add_tag(oe2, 'm:endChr', {'m:val': '|'})
        elif c == '{Vmatrix}':
            oe3 = XML.add_tag(oe2, 'm:begChr', {'m:val': '‖'})
            oe3 = XML.add_tag(oe2, 'm:endChr', {'m:val': '‖'})
        else:
            oe3 = XML.add_tag(oe2, 'm:begChr', {'m:val': ''})
            oe3 = XML.add_tag(oe2, 'm:endChr', {'m:val': ''})
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub(oe3, chars_state)
        #
        oe2 = XML.add_tag(oe1, 'm:e', {})
        oe3 = XML.add_tag(oe2, 'm:m', {})
        #
        oe4 = XML.add_tag(oe3, 'm:mPr', {})
        oe5 = XML.add_tag(oe4, 'm:ctrlPr', {})
        cls.__decorate_nub(oe5, chars_state)
        #
        for row in mtrx:
            oe4 = XML.add_tag(oe3, 'm:mr', {})
            for cel in row:
                oe5 = XML.add_tag(oe4, 'm:e', {})
                cls._write_math_exp(oe5, chars_state, '{' + cel + '}')

    # PARENTHESIS
    @classmethod
    def _write_prn(cls, oe0, chars_state, t1, t2):
        oe1 = XML.add_tag(oe0, 'm:d', {})
        oe2 = XML.add_tag(oe1, 'm:dPr', {})
        oe3 = XML.add_tag(oe2, 'm:begChr', {'m:val': t1[0]})
        oe3 = XML.add_tag(oe2, 'm:endChr', {'m:val': t1[1]})
        oe3 = XML.add_tag(oe2, 'm:ctrlPr', {})
        cls.__decorate_nub(oe3, chars_state)
        oe2 = XML.add_tag(oe1, 'm:e', {})
        cls._write_math_exp(oe2, chars_state, t2)


class Document:

    """A class to handle document"""

    def __init__(self):
        self.docx_file = ''
        self.formal_md_lines = []
        self.md_lines = []
        self.paragraphs = []

    def get_md_lines(self, formal_md_lines):
        md_lines = []
        for i, rml in enumerate(formal_md_lines):
            ml = MdLine(i + 1, rml)
            md_lines.append(ml)
        # self.md_lines = md_lines
        return md_lines

    def get_raw_paragraphs(self, md_lines):
        raw_paragraphs = []
        block = []
        for ml in md_lines:
            # ISOLATE CONFIGURATIONS
            if 'is_in_configurations' not in locals():
                is_in_configurations = True
            if ml.text != '':
                is_in_configurations = False
            if is_in_configurations:
                block.append(ml)
                continue
            # CONFIRM BLOCK END
            is_block_end = False
            if ml.raw_text == '':
                is_block_end = True
            if len(block) > 0:
                pre_text = block[-1].raw_text
                cur_text = ml.raw_text
                for pc in [ParagraphChapter, ParagraphSection, ParagraphList]:
                    res = '^\\s*' + pc.res_symbol + '\\s+\\S+.*$'
                    if re.match(res, pre_text) and re.match(res, cur_text):
                        is_block_end = True
                    # REMOVED 23.10.23 >
                    # res_r = '^\\s*' + pc.res_reviser + '(\\s.*)?$'
                    # res_s = '^\\s*' + pc.res_symbol + '\\s+\\S+.*$'
                    # if re.match(res_r + '|' + res_s, pre_text):
                    #     if re.match(res_r + '|' + res_s, cur_text):
                    #         is_block_end = True
                    # <
                if re.match(ParagraphRemarks.res_feature, pre_text) and \
                   not re.match(ParagraphRemarks.res_feature, cur_text):
                    is_block_end = True
            # RECORD
            if is_block_end:
                if len(block) == 0:
                    if ml.raw_text != '':
                        block.append(ml)
                    continue
                if re.match('^```.*$', block[0].raw_text):
                    if len(block) == 1:
                        block.append(ml)
                        continue
                    elif not re.match('^.*```$', block[-1].raw_text):
                        block.append(ml)
                        continue
                rp = RawParagraph(block)
                raw_paragraphs.append(rp)
                block = []
            if ml.raw_text != '':
                block.append(ml)
        if len(block) > 0:
            rp = RawParagraph(block)
            raw_paragraphs.append(rp)
            block = []
        # self.raw_paragraphs = raw_paragraphs
        return raw_paragraphs

    def get_paragraphs(self, raw_paragraphs):
        paragraphs = []
        cr = []
        sr = []
        lr = []
        er = []
        hr = []
        sd = []
        res_v = '^v=(' + RES_NUMBER + ')$'
        res_cv = '^V=(' + RES_NUMBER + ')$'
        for rp in raw_paragraphs:
            full_text = rp.full_text
            if rp.paragraph_class == 'empty' or rp.paragraph_class == 'blank':
                cr += rp.chapter_revisers
                sr += rp.section_revisers
                lr += rp.list_revisers
                er += rp.length_revisers
                hr += rp.head_font_revisers + rp.tail_font_revisers
                sd += rp.section_depth_setters
                if rp.paragraph_class == 'blank':
                    nl = full_text.count('\n')
                    er += ['v=' + str(nl)]
            else:
                rp.chapter_revisers = cr + rp.chapter_revisers
                rp.section_revisers = sr + rp.section_revisers
                rp.list_revisers = lr + rp.list_revisers
                for rev in er:
                    if re.match(res_v, rev):
                        rp.length_revisers = [rev] + rp.length_revisers
                    if re.match(res_cv, rev):
                        rev = re.sub('^V=', 'v=', rev)
                        rp.length_revisers = [rev] + rp.length_revisers
                rp.head_font_revisers = hr + rp.head_font_revisers
                rp.section_depth_setters = sd + rp.section_depth_setters
                cr = []
                sr = []
                lr = []
                er = []
                hr = []
                sd = []
                p = rp.get_paragraph()
                paragraphs.append(p)
        # self.paragraphs = paragraphs
        return paragraphs

    def modify_paragraphs(self, paragraphs):
        for i, p in enumerate(paragraphs):
            if i > 0:
                p_prev = paragraphs[i - 1]
            # ARTICLE TITLE (MIMI=EAR)
            if Form.document_style == 'j' and \
               p.paragraph_class == 'section' and \
               p.head_section_depth == 2 and \
               p.tail_section_depth == 2 and \
               i > 0 and \
               p_prev.paragraph_class == 'alignment' and \
               p_prev.alignment == 'left':
                p_prev.length_docx['space before'] \
                    += p.length_conf['space before']
                p.length_docx['space before'] \
                    -= p.length_conf['space before']
        m = len(paragraphs) - 1
        for i, p in enumerate(paragraphs):
            if i > 0:
                p_prev = paragraphs[i - 1]
            if i < m:
                p_next = paragraphs[i + 1]
            # SECTION DEPTH 1
            if p.paragraph_class == 'section' and \
               ParagraphSection._get_section_depths(p.full_text) == (1, 1):
                # BEFORE
                if i > 0:
                    if p_prev.length_docx['space after'] >= 0.1:
                        p_prev.length_docx['space after'] += 0.1
                    elif p_prev.length_docx['space after'] >= 0.0:
                        p_prev.length_docx['space after'] *= 2
                if True:
                    if p.length_docx['space before'] >= 0.1:
                        p.length_docx['space before'] += 0.1
                    elif p.length_docx['space before'] >= 0.0:
                        p.length_docx['space before'] *= 2
                # AFTER
                if True:
                    if p.length_docx['space after'] >= 0.2:
                        p.length_docx['space after'] -= 0.1
                    elif p.length_docx['space after'] >= 0.0:
                        p.length_docx['space after'] /= 2
                if i < m:
                    if p_next.length_docx['space before'] >= 0.2:
                        p_next.length_docx['space before'] -= 0.1
                    elif p_next.length_docx['space before'] >= 0.0:
                        p_next.length_docx['space before'] /= 2
            # TABLE OR IMAGE
            if p.paragraph_class == 'table' or p.paragraph_class == 'image':
                if i > 0:
                    if p.length_docx['space before'] < 0:
                        msg = '※ 警告: ' \
                            + '段落前の余白「v」の値が小さ過ぎます'
                        # msg = 'warning: ' \
                        #     + '"space before" is too small'
                        p.md_lines[0].append_warning_message(msg)
                        p.length_docx['space before'] = 0.0
                    sa = p_prev.length_docx['space after']
                    sb = p.length_docx['space before'] \
                        - p.length_clas['space before']
                    mx = max([0, sa, sb])
                    mn = min([0, sa, sb])
                    if mx > 0:
                        p_prev.length_docx['space after'] \
                            = mx + p.length_clas['space before']
                    else:
                        p_prev.length_docx['space after'] \
                            = mn + p.length_clas['space before']
                    p.length_docx['space before'] = 0.0
                if i < m:
                    if p.length_docx['space after'] < 0:
                        msg = '※ 警告: ' \
                            + '段落前の余白「V」の値が小さ過ぎます'
                        # msg = 'warning: ' \
                        #     + '"space after" is too small'
                        p.md_lines[0].append_warning_message(msg)
                        p.length_docx['space after'] = 0.0
                    sa = p.length_docx['space after'] \
                        - p.length_clas['space after']
                    sb = p_next.length_docx['space before']
                    mx = max([0, sa, sb])
                    mn = min([0, sa, sb])
                    p.length_docx['space after'] = 0.0
                    if mx > 0:
                        p_next.length_docx['space before'] \
                            = mx + p.length_clas['space after']
                    else:
                        p_next.length_docx['space before'] \
                            = mn + p.length_clas['space after']
        # SEARCH FOR ATTACHED PAGEBREAK
        exception = ['empty', 'blank', 'table', 'pagebreak', 'horizontalline',
                     'remarks']
        for i, p in enumerate(paragraphs):
            if i > 0:
                p_prev = paragraphs[i - 1]
            if p.paragraph_class == 'pagebreak':
                if p_prev.paragraph_class not in exception:
                    p.is_attached_pagebreak = True
        return self.paragraphs

    def write_property(self, ms_doc):
        host = socket.gethostname()
        if host is None:
            host = '-'
        hh = self.__get_hash(host)
        user = getpass.getuser()
        if user is None:
            user = '='
        hu = self.__get_hash(user)
        tt = Form.document_title
        if Form.document_style == 'n':
            ct = '（普通）'
        elif Form.document_style == 'k':
            ct = '（契約）'
        elif Form.document_style == 'j':
            ct = '（条文）'
        at = hu + '@' + hh + ' (makdo ' + __version__ + ')'
        ct = self.__get_datetime(Form.created_time)
        mt = self.__get_datetime(Form.modified_time)
        vn = Form.version_number
        cs = Form.content_status
        ms_cp = ms_doc.core_properties
        ddn = datetime.datetime.now(datetime.timezone.utc)
        ms_cp.identifier \
            = 'makdo(' + __version__.split()[0] + ');' \
            + hu + '@' + hh + ';' + ddn.strftime('%Y-%m-%dT%H:%M:%SZ')
        ms_cp.title = tt               # タイトル
        # ms_cp.subject = ''           # 件名
        # ms_cp.keywords = ''          # タグ
        ms_cp.category = ct            # 分類項目
        # ms_cp.comments = ''          # コメント（generated by python-docx）
        ms_cp.author = at              # 作成者
        # ms_cp.last_modified_by = ''  # 前回保存者
        ms_cp.version = vn             # バージョン番号
        # ms_cp.revision = 1           # 改訂番号
        ms_cp.created = ct             # コンテンツの作成日時
        ms_cp.modified = mt            # 前回保存時
        # ms_cp.last_printed = pt      # 前回印刷日
        ms_cp.content_status = cs      # 内容の状態
        # ms_cp.language = ''          # 言語

    @staticmethod
    def __get_hash(st):
        # ''  owicwvnu
        # '-' sojfooxd
        # '=' empzhdhk
        x = 9973
        b = 99999989
        m = 999999999989
        z = int(((4 ** 20) - 1) / (4 - 1) * 2)
        for c in st + ' 2022.05.07 07:31:03':
            x = (x * b + ord(c)) % m
            x = x ^ z
        hs = ''
        for i in range(8):
            hs += chr(x % 26 + 97)
            x = int(x / 26)
        return hs

    @staticmethod
    def __get_datetime(datetime_str):
        # TIMEZONE IS NOT SUPPORTED (UTC ONLY)
        # X datetime_cls = datetime.datetime.now(jst)
        utc = datetime.timezone.utc
        jst = datetime.timezone(datetime.timedelta(hours=+9))
        try:
            datetime_cls = datetime.datetime.fromisoformat(datetime_str)
            if datetime_cls.tzinfo is None:
                datetime_cls = datetime_cls.replace(tzinfo=jst)
            datetime_cls = datetime_cls.astimezone(utc)
        except BaseException:
            datetime_cls = datetime.datetime.utcnow()
            datetime_cls = datetime_cls.replace(tzinfo=utc)
        return datetime_cls

    def write_document(self, ms_doc):
        for p in self.paragraphs:
            p.write_paragraph(ms_doc)

    def print_warning_messages(self):
        for p in self.paragraphs:
            p.print_warning_messages()

    # UNFOLD
    @staticmethod
    def unfold(old_md_lines):
        # |                ->  |
        # |## www...[3]    ->  |## www
        # |                ->  |
        # |...[1]#### yyy  ->  |### xxx
        # |                ->  |
        # |zzz             ->  |#### yyy
        # |                ->  |
        # |...[2]### xxx   ->  |zzz
        # |                ->  |
        # |#### yyy...[1]  ->  |
        # |                ->  |
        # |...[3]## www    ->  |
        # |                ->  |
        # |### xxx...[2]   ->  |
        # |                ->  |
        new_md_lines = []
        remain_md_lines = [True for i in old_md_lines]
        m = len(old_md_lines) - 1
        line_numbers = [0]
        res_mark = '\\.\\.\\.\\[([0-9])+\\]'
        res_from = '^(#+(?:-#+)*(?:\\s.*)?)' + res_mark + '$'
        res_to = '^' + res_mark + '#+(-#+)*(\\s|$)'
        while line_numbers != []:
            i = line_numbers[-1]
            if i > m:
                line_numbers.pop(-1)
                continue
            if not remain_md_lines[i]:
                line_numbers.pop(-1)
                continue
            if re.match(res_to, old_md_lines[i].text):
                line_numbers.pop(-1)
                continue
            if re.match(res_from, old_md_lines[i].text) and \
               re.match(NOT_ESCAPED + res_mark + '$', old_md_lines[i].text):
                folding_number \
                    = re.sub(res_from, '\\2', old_md_lines[i].text)
                old_md_lines[i].text \
                    = re.sub(res_from, '\\1', old_md_lines[i].text)
                # APPEND "FROM LINE"
                new_md_lines.append(old_md_lines[i])
                remain_md_lines[i] = False
                line_numbers[-1] += 1
                if i < m and old_md_lines[i + 1].raw_text == '':
                    # SKIP NEXT EMPTY LINE
                    # new_md_lines.append(old_md_lines[i])
                    remain_md_lines[i + 1] = False
                    line_numbers[-1] += 1
                for j, ml in enumerate(old_md_lines):
                    if not remain_md_lines[j]:
                        continue
                    res = '^\\.\\.\\.\\[' + folding_number + '\\]'
                    if re.match(res, ml.text):
                        line_numbers.append(j)
                        # JUMP TO "TO LINE"
                        # new_md_lines.append(old_md_lines[j])
                        remain_md_lines[j] = False
                        line_numbers[-1] += 1
            else:
                # APPEND "USUAL LINE"
                new_md_lines.append(old_md_lines[i])
                remain_md_lines[i] = False
                line_numbers[-1] += 1
        must_warn = True
        for i, ml in enumerate(old_md_lines):
            if remain_md_lines[i]:
                if must_warn:
                    msg = '※ 警告: ' \
                        + '折り畳まれたセクションが残っています'
                    # msg = 'warning: ' \
                    #     + 'folded sections remain'
                    old_md_lines[i].append_warning_message(msg)
                    must_warn = False
                new_md_lines.append(old_md_lines[i])
        return new_md_lines


class RawParagraph:

    """A class to handle raw paragraph"""

    raw_paragraph_number = 0

    def __init__(self, md_lines):
        # DECLARATION
        self.raw_paragraph_number = -1
        self.md_lines = []
        self.chapter_revisers = []
        self.section_revisers = []
        self.list_revisers = []
        self.length_revisers = []
        self.head_font_revisers = []
        self.tail_font_revisers = []
        self.full_text = ''
        self.full_text_del = ''
        self.full_text_ins = ''
        self.section_depth_setters = []
        self.paragraph_class = ''
        # SUBSTITUTION
        RawParagraph.raw_paragraph_number += 1
        self.raw_paragraph_number = RawParagraph.raw_paragraph_number
        self.md_lines = md_lines
        self.chapter_revisers, \
            self.section_revisers, \
            self.list_revisers, \
            self.length_revisers, \
            self.head_font_revisers, \
            self.tail_font_revisers, \
            self.md_lines \
            = self._get_revisers(self.md_lines)
        self.full_text = self._get_full_text(self.md_lines)
        self.full_text_del = self._get_full_text_del(self.full_text)
        self.full_text_ins = self._get_full_text_ins(self.full_text)
        self.section_depth_setters, self.full_text \
            = self._get_section_depth_setters(self.full_text)
        self.paragraph_class = self._get_paragraph_class()

    @staticmethod
    def _get_revisers(md_lines):
        chapter_revisers = []
        section_revisers = []
        list_revisers = []
        length_revisers = []
        head_font_revisers = []
        tail_font_revisers = []
        res_cr = '^\\s*(' + ParagraphChapter.res_reviser + ')(?:\\s*(.*))?$'
        res_sr = '^\\s*(' + ParagraphSection.res_reviser + ')(?:\\s*(.*))?$'
        res_lr = '^(\\s*' + ParagraphList.res_reviser + ')(?:\\s*(.*))?$'
        res_er = '^\\s*((?:v|V|X|x|<<|<|>)=' + RES_NUMBER + ')(?:\\s*(.*))?$'
        res_fr = '^(' + '|'.join(FONT_DECORATORS) + ')(.*)$'
        res_tr = NOT_ESCAPED + '(' + '|'.join(FONT_DECORATORS) + ')$'
        res_hl = '^' + ParagraphHorizontalLine.res_feature + '$'
        # HEAD REVISERS
        for ml in md_lines:
            if re.match('^' + res_lr, ml.text):
                ml.text = ml.beg_space + ml.text
            if re.match('^.*(  |\t|\u3000)$', ml.spaced_text):
                ml.text = re.sub('<br>$', '  ', ml.text)
            while True:
                if False:
                    pass
                elif re.match(res_cr, ml.text):
                    reviser = re.sub(res_cr, '\\1', ml.text)
                    ml.text = re.sub(res_cr, '\\5', ml.text)
                    chapter_revisers.append(reviser)
                elif re.match(res_sr, ml.text):
                    reviser = re.sub(res_sr, '\\1', ml.text)
                    ml.text = re.sub(res_sr, '\\5', ml.text)
                    section_revisers.append(reviser)
                elif re.match(res_lr, ml.text):
                    reviser = re.sub(res_lr, '\\1', ml.text)
                    ml.text = re.sub(res_lr, '\\3', ml.text)
                    list_revisers.append(reviser)
                elif re.match(res_er, ml.text):
                    reviser = re.sub(res_er, '\\1', ml.text)
                    ml.text = re.sub(res_er, '\\2', ml.text)
                    length_revisers.append(reviser)
                elif (re.match(res_fr, ml.text) and
                      not re.match(res_hl, ml.text)):
                    reviser = re.sub(res_fr, '\\1', ml.text)
                    ml.text = re.sub(res_fr, '\\2', ml.text)
                    head_font_revisers.append(reviser)
                else:
                    break
            if ml.text != '':
                if re.match(NOT_ESCAPED + '  $', ml.text):
                    ml.text = re.sub('  $', '<br>', ml.text)
                break
        # TAIL REVISERS
        for ml in md_lines[::-1]:
            if re.match('^.*(  |\t|\u3000)$', ml.spaced_text):
                ml.text = re.sub('<br>$', '  ', ml.text)
            while True:
                if re.match(NOT_ESCAPED + '\\|:-+$', ml.text):
                    break  # table vertical configuration
                if re.match(res_tr, ml.text) and not re.match(res_hl, ml.text):
                    reviser = re.sub(res_tr, '\\2', ml.text)
                    ml.text = re.sub(res_tr, '\\1', ml.text)
                    tail_font_revisers.insert(0, reviser)
                else:
                    break
            if ml.text != '':
                if re.match('.*  $', ml.text):
                    ml.text = re.sub('  $', '<br>', ml.text)
                break
        # EXAMPLE "# ###=1"
        full_text = ''
        for ml in md_lines:
            full_text += ml.text + ' '
        res = '^\\s*' + \
            '(' + ParagraphSection.res_symbol + ')' + \
            '((\\s+' + ParagraphSection.res_reviser + ')+)' + \
            '\\s*$'
        if re.match(res, full_text):
            symbol = re.sub(res, '\\1', full_text)
            revisers = re.sub(res, '\\4', full_text)
            for ml in md_lines:
                ml.text = ''
            md_lines[0].text = symbol
            for r in revisers.split():
                section_revisers.append(r)
        # self.chapter_revisers = chapter_revisers
        # self.section_revisers = section_revisers
        # self.length_revisers = length_revisers
        # self.head_font_revisers = head_font_revisers
        # self.tail_font_revisers = tail_font_revisers
        # self.md_lines = md_lines
        return chapter_revisers, section_revisers, list_revisers, \
            length_revisers, head_font_revisers, tail_font_revisers, md_lines

    @staticmethod
    def _get_full_text(md_lines):
        full_text = ''
        for ml in md_lines:
            if ml.text != '':
                full_text += ml.text + ' '
        full_text = re.sub('\t', ' ', full_text)
        full_text = re.sub(' +', ' ', full_text)
        full_text = re.sub('^ ', '', full_text)
        full_text = re.sub(' $', '', full_text)
        # FOR PARAGRAPH LIST
        res = '^' + ParagraphList.res_symbol
        if re.match(res, full_text):
            for ml in md_lines:
                if re.match(res, ml.text):
                    full_text = ml.beg_space + full_text
        # self.full_text = full_text
        return full_text

    @classmethod
    def _get_full_text_del(cls, full_text):
        full_text_del \
            = cls._get_full_text_del_or_ins(full_text,
                                            '\\+>', '<\\+', '->', '<-')
        return full_text_del

    @classmethod
    def _get_full_text_ins(cls, full_text):
        full_text_ins \
            = cls._get_full_text_del_or_ins(full_text,
                                            '->', '<-', '\\+>', '<\\+')
        return full_text_ins

    @staticmethod
    def _get_full_text_del_or_ins(full_text,
                                  beg_erase, end_erase,
                                  beg_leave, end_leave):
        full_text_erase = ''
        full_text_leave = ''
        track_changes = ''
        in_to_erase = False
        for c in full_text:
            if in_to_erase:
                full_text_erase += c
                if re.match(NOT_ESCAPED + end_erase + '$', full_text_erase):
                    in_to_erase = False
                full_text_erase = re.sub(end_erase + '$', '', full_text_erase)
            else:
                full_text_leave += c
                if re.match(NOT_ESCAPED + beg_erase + '$', full_text_leave):
                    in_to_erase = True
                full_text_leave = re.sub(beg_erase + '$', '', full_text_leave)
                full_text_leave = re.sub(beg_leave + '$', '', full_text_leave)
                full_text_leave = re.sub(end_leave + '$', '', full_text_leave)
        return full_text_leave

    @staticmethod
    def _get_section_depth_setters(full_text):
        max_depth = len(ParagraphSection.states)
        section_depth_setters = []
        res = '^#{1,' + str(max_depth) + '}$'
        if re.match(res, full_text):
            section_depth_setters = [full_text]
            full_text = ''
        # self.section_depth_setters = depth_setters
        # self.full_text = full_text
        return section_depth_setters, full_text

    def _get_paragraph_class(self):
        ft = self.full_text
        hfrs = self.head_font_revisers
        tfrs = self.tail_font_revisers
        if False:
            pass
        elif ParagraphEmpty.is_this_class(ft, hfrs, tfrs):
            return 'empty'
        elif ParagraphBlank.is_this_class(ft, hfrs, tfrs):
            return 'blank'
        elif ParagraphChapter.is_this_class(ft, hfrs, tfrs):
            return 'chapter'
        elif ParagraphSection.is_this_class(ft, hfrs, tfrs):
            return 'section'
        elif ParagraphList.is_this_class(ft, hfrs, tfrs):
            return 'list'
        elif ParagraphTable.is_this_class(ft, hfrs, tfrs):
            return 'table'
        elif ParagraphImage.is_this_class(ft, hfrs, tfrs):
            return 'image'
        elif ParagraphMath.is_this_class(ft, hfrs, tfrs):
            return 'math'
        elif ParagraphAlignment.is_this_class(ft, hfrs, tfrs):
            return 'alignment'
        elif ParagraphPreformatted.is_this_class(ft, hfrs, tfrs):
            return 'preformatted'
        elif ParagraphPagebreak.is_this_class(ft, hfrs, tfrs):
            return 'pagebreak'
        elif ParagraphHorizontalLine.is_this_class(ft, hfrs, tfrs):
            return 'horizontalline'
        elif ParagraphBreakdown.is_this_class(ft, hfrs, tfrs):
            return 'breakdown'
        elif ParagraphRemarks.is_this_class(ft, hfrs, tfrs):
            return 'remarks'
        else:
            return 'sentence'

    def get_paragraph(self):
        paragraph_class = self.paragraph_class
        if False:
            pass
        elif paragraph_class == 'empty':
            return ParagraphEmpty(self)
        elif paragraph_class == 'blank':
            return ParagraphBlank(self)
        elif paragraph_class == 'chapter':
            return ParagraphChapter(self)
        elif paragraph_class == 'section':
            return ParagraphSection(self)
        elif paragraph_class == 'list':
            return ParagraphList(self)
        elif paragraph_class == 'table':
            return ParagraphTable(self)
        elif paragraph_class == 'image':
            return ParagraphImage(self)
        elif paragraph_class == 'math':
            return ParagraphMath(self)
        elif paragraph_class == 'alignment':
            return ParagraphAlignment(self)
        elif paragraph_class == 'preformatted':
            return ParagraphPreformatted(self)
        elif paragraph_class == 'pagebreak':
            return ParagraphPagebreak(self)
        elif paragraph_class == 'horizontalline':
            return ParagraphHorizontalLine(self)
        elif paragraph_class == 'breakdown':
            return ParagraphBreakdown(self)
        elif paragraph_class == 'remarks':
            return ParagraphRemarks(self)
        else:
            return ParagraphSentence(self)


class Paragraph:

    """A class to handle empty paragraph"""

    paragraph_number = 0

    paragraph_class = None
    res_feature = None

    bridge_head_section_depth = 0
    bridge_tail_section_depth = 0

    bridge_chars_state = None

    @classmethod
    def is_this_class(cls, full_text,
                      head_font_revisers=[], tail_font_revisers=[]):
        if re.match(cls.res_feature, full_text):
            return True
        return False

    def __init__(self, raw_paragraph):
        # RECEIVE
        self.raw_paragraph_number = raw_paragraph.raw_paragraph_number
        self.md_lines = raw_paragraph.md_lines
        self.chapter_revisers = raw_paragraph.chapter_revisers
        self.section_revisers = raw_paragraph.section_revisers
        self.list_revisers = raw_paragraph.list_revisers
        self.length_revisers = raw_paragraph.length_revisers
        self.head_font_revisers = raw_paragraph.head_font_revisers
        self.tail_font_revisers = raw_paragraph.tail_font_revisers
        self.full_text = raw_paragraph.full_text
        self.section_depth_setters = raw_paragraph.section_depth_setters
        self.paragraph_class = raw_paragraph.paragraph_class
        # DECLARE
        self.paragraph_number = -1
        self.head_section_depth = -1
        self.tail_section_depth = -1
        self.proper_depth = -1
        self.length_revi = {}
        self.length_conf = {}
        self.length_clas = {}
        self.length_docx = {}
        self.alignment = ''
        self.text_to_write = ''
        self.text_to_write_with_reviser = ''
        self.beg_chars_state = CharsState()
        self.end_chars_state = CharsState()
        self.chars_state = CharsState()
        self.char_spacing = 0.0
        # SUBSTITUTE
        Paragraph.paragraph_number += 1
        self.paragraph_number = Paragraph.paragraph_number
        self._apply_section_depths_setters(self.section_depth_setters)
        self.head_section_depth, self.tail_section_depth \
            = self._get_section_depths(self.full_text)
        self.proper_depth = self._get_proper_depth(self.full_text)
        self.alignment = self._get_alignment()
        # APPLY REVISERS
        ParagraphChapter._apply_revisers(self.chapter_revisers,
                                         self.md_lines)
        ParagraphSection._apply_revisers(self.section_revisers,
                                         self.md_lines)
        ParagraphList._apply_revisers(self.list_revisers,
                                      self.md_lines)
        ParagraphList._reset_states(self.paragraph_class)
        # GET LENGTH
        self.length_revi = self._get_length_revi()
        self.length_conf = self._get_length_conf()
        self.length_clas = self._get_length_clas()
        self.length_docx = self._get_length_docx()
        # CHECK
        self._check_format()
        # GET TEXT
        self._edit_data()
        self.text_to_write = self._get_text_to_write()
        self.text_to_write_with_reviser \
            = self._get_text_to_write_with_reviser()

    @classmethod
    def _apply_section_depths_setters(cls, section_depth_setters):
        for sds in section_depth_setters:
            depth = len(sds)
            if depth > 0:
                Paragraph.previous_head_section_depth = depth
                Paragraph.bridge_tail_section_depth = depth

    @classmethod
    def _get_section_depths(cls, full_text):
        head_section_depth = 0
        tail_section_depth = 0
        # self.head_section_depth = head_section_depth
        # self.tail_section_depth = tail_section_depth
        return head_section_depth, tail_section_depth

    @classmethod
    def _get_proper_depth(cls, full_text):
        proper_depth = 0
        # self.proper_depth = proper_depth
        return proper_depth

    def _get_alignment(self):
        paragraph_class = self.paragraph_class
        head_section_depth = self.head_section_depth
        full_text = self.full_text
        alignment = ''
        if paragraph_class == 'section' and head_section_depth == 1:
            alignment = 'center'
        if paragraph_class == 'alignment':
            if re.match('^:\\s.*\\s:$', full_text):
                alignment = 'center'
            elif re.match('^:\\s.*$', full_text):
                alignment = 'left'
            elif re.match('^.*\\s:$', full_text):
                alignment = 'right'
        # self.alignment = alignment
        return alignment

    @classmethod
    def _apply_revisers(cls, revisers, md_lines):
        res = '^' + cls.res_reviser + '$'
        if cls.paragraph_class == 'chapter':
            char = '$'
        elif cls.paragraph_class == 'section':
            char = '#'
        else:
            return
        for rev in revisers:
            md_line = md_lines[0]
            res_line = '^(.*\\s)?' \
                + rev.replace(char, '\\' + char) \
                + '(\\s.*)?$'
            for ml in md_lines:
                if re.match(res_line, ml.raw_text):
                    md_line = ml
                    break
            if re.match(res, rev):
                trunk = re.sub(res, '\\1', rev)
                branc = re.sub(res, '\\2', rev)
                chval = re.sub(res, '\\3', rev)
                xdepth = len(trunk) - 1
                ydepth = len(branc.replace(char, ''))
                value = int(chval) - 1
                cls.__set_state(xdepth, ydepth, value, md_line)

    @classmethod
    def __set_state(cls, xdepth, ydepth, value, md_line):
        paragraph_class_ja = cls.paragraph_class_ja
        paragraph_class = cls.paragraph_class
        states = cls.states
        if xdepth >= len(states):
            msg = '※ 警告: ' + paragraph_class_ja \
                + 'の深さが上限を超えています'
            # msg = 'warning: ' + paragraph_class \
            #     + ' depth exceeds limit'
            md_line.append_warning_message(msg)
        elif ydepth >= len(states[xdepth]):
            msg = '※ 警告: ' + paragraph_class_ja \
                + 'の枝が上限を超えています'
            # msg = 'warning: ' + paragraph_class \
            #     + ' branch exceeds limit'
            md_line.append_warning_message(msg)
        for x in range(len(states)):
            for y in range(len(states[x])):
                if x < xdepth:
                    continue
                elif x == xdepth:
                    if y < ydepth:
                        if states[x][y] == 0:
                            msg = '※ 警告: ' + paragraph_class_ja \
                                + 'の枝が"0"を含んでいます'
                            # msg = 'warning: ' + paragraph_class \
                            #     + ' branch has "0"'
                            md_line.append_warning_message(msg)
                    elif y == ydepth:
                        if value is None:
                            states[x][y] += 1  # step state
                        else:
                            states[x][y] = value
                    else:
                        states[x][y] = 0
                else:
                    states[x][y] = 0

    def _get_length_revi(self):
        length_revisers = self.length_revisers
        length_revi \
            = {'space before': 0.0, 'space after': 0.0, 'line spacing': 0.0,
               'first indent': 0.0, 'left indent': 0.0, 'right indent': 0.0}
        self.char_spacing = 0.0
        res_v = '^v=(' + RES_NUMBER + ')$'
        res_cv = '^V=(' + RES_NUMBER + ')$'
        res_cx = '^X=(' + RES_NUMBER + ')$'
        res_x = '^x=(' + RES_NUMBER + ')$'
        res_gg = '^<<=(' + RES_NUMBER + ')$'
        res_g = '^<=(' + RES_NUMBER + ')$'
        res_l = '^>=(' + RES_NUMBER + ')$'
        for lr in length_revisers:
            if re.match(res_v, lr):
                length_revi['space before'] += float(re.sub(res_v, '\\1', lr))
            elif re.match(res_cv, lr):
                length_revi['space after'] += float(re.sub(res_cv, '\\1', lr))
            elif re.match(res_cx, lr):
                length_revi['line spacing'] += float(re.sub(res_cx, '\\1', lr))
            elif re.match(res_x, lr):
                self.char_spacing += float(re.sub(res_x, '\\1', lr))
            elif re.match(res_gg, lr):
                length_revi['first indent'] -= float(re.sub(res_gg, '\\1', lr))
            elif re.match(res_g, lr):
                length_revi['left indent'] -= float(re.sub(res_g, '\\1', lr))
            elif re.match(res_l, lr):
                length_revi['right indent'] -= float(re.sub(res_l, '\\1', lr))
        for ln in length_revi:
            length_revi[ln] = round(length_revi[ln], 3)
        # self.length_revi = length_revi
        return length_revi

    def _get_length_conf(self):
        paragraph_class = self.paragraph_class
        hd = self.head_section_depth
        td = self.tail_section_depth
        sds = self.section_depth_setters
        if paragraph_class != 'section' and len(sds) > 0:
            hd = len(sds[0])
            td = len(sds[-1])
        length_conf \
            = {'space before': 0.0, 'space after': 0.0, 'line spacing': 0.0,
               'first indent': 0.0, 'left indent': 0.0, 'right indent': 0.0}
        sb = (Form.space_before + ',,,,,,,').split(',')
        sa = (Form.space_after + ',,,,,,,').split(',')
        if paragraph_class == 'section':
            if hd <= len(sb) and sb[hd - 1] != '':
                length_conf['space before'] += float(sb[hd - 1])
        if paragraph_class == 'section':
            if td <= len(sa) and sa[td - 1] != '':
                length_conf['space after'] += float(sa[td - 1])
        for ln in length_conf:
            length_conf[ln] = round(length_conf[ln], 3)
        # self.length_conf = length_conf
        return length_conf

    def _get_length_clas(self):
        paragraph_class = self.paragraph_class
        head_section_depth = self.head_section_depth
        tail_section_depth = self.tail_section_depth
        proper_depth = self.proper_depth
        length_revi = self.length_revi
        line_spacing = Form.line_spacing
        length_clas \
            = {'space before': 0.0, 'space after': 0.0, 'line spacing': 0.0,
               'first indent': 0.0, 'left indent': 0.0, 'right indent': 0.0}
        if paragraph_class == 'chapter':
            length_clas['first indent'] = -1.0
            length_clas['left indent'] = proper_depth + 0.0
        elif paragraph_class == 'section':
            if head_section_depth > 1:
                length_clas['first indent'] \
                    = head_section_depth - tail_section_depth - 1.0
            if tail_section_depth > 1:
                length_clas['left indent'] = tail_section_depth - 1.0
        elif paragraph_class == 'list':
            length_clas['first indent'] = -1.0
            length_clas['left indent'] = proper_depth + 0.0
            if tail_section_depth > 0:
                length_clas['left indent'] += tail_section_depth - 1.0
        elif paragraph_class == 'table':
            # ls     sb    sa
            # -0.33  0.19  0.13
            #  0.00  0.45  0.20
            # +0.30  0.67  0.27
            # +0.60  0.91  0.32
            # +0.90  1.16  0.38
            # +1.20  1.40  0.44
            # +1.50  1.66  0.50
            # +1.80  1.76  0.56
            # +2.10  1.78  0.62
            # +2.40  1.78  0.68
            # +2.70  1.78  0.74
            # +3.00  1.78  0.80
            # +3.30  1.78  0.80
            # +3.60  1.78  0.80
            ls = self.length_revi['line spacing']
            if ls <= 1.66:
                length_clas['space before'] += 0.80 * ls + TABLE_SPACE_BEFORE
            else:
                length_clas['space before'] += 1.78
            if ls <= 3.00:
                length_clas['space after'] += 0.20 * ls + TABLE_SPACE_AFTER
            else:
                length_clas['space after'] += 0.80
        elif paragraph_class == 'image':
            length_clas['space before'] += IMAGE_SPACE_BEFORE
            length_clas['space after'] += IMAGE_SPACE_AFTER
        elif paragraph_class == 'preformatted':
            if tail_section_depth > 0:
                length_clas['first indent'] = 0.0
                length_clas['left indent'] = tail_section_depth - 0.0
        elif paragraph_class == 'sentence':
            if tail_section_depth > 0:
                length_clas['first indent'] = 1.0
                length_clas['left indent'] = tail_section_depth - 1.0
        if paragraph_class == 'section' or \
           paragraph_class == 'list' or \
           paragraph_class == 'preformatted' or \
           paragraph_class == 'sentence':
            if ParagraphSection.states[1][0] <= 0 and tail_section_depth > 2:
                length_clas['left indent'] -= 1.0
        if Form.document_style == 'j':
            if ParagraphSection.states[1][0] > 0 and tail_section_depth > 2:
                length_clas['left indent'] -= 1.0
        for ln in length_clas:
            length_clas[ln] = round(length_clas[ln], 3)
        # self.length_clas = length_clas
        return length_clas

    def _get_length_docx(self):
        length_revi = self.length_revi
        length_conf = self.length_conf
        length_clas = self.length_clas
        length_docx \
            = {'space before': 0.0, 'space after': 0.0, 'line spacing': 0.0,
               'first indent': 0.0, 'left indent': 0.0, 'right indent': 0.0}
        for ln in length_docx:
            length_docx[ln] \
                = length_revi[ln] + length_conf[ln] + length_clas[ln]
        # MODIFY SPACE BEFORE AND AFTER
        ls80 = length_docx['line spacing'] * .80
        ls20 = length_docx['line spacing'] * .20
        if length_docx['space before'] >= ls80 * 1.33333:
            length_docx['space before'] -= ls80
        else:
            length_docx['space before'] /= 4
        if length_docx['space after'] >= ls20 * 1.33333:
            length_docx['space after'] -= ls20
        else:
            length_docx['space after'] /= 4
        for ln in length_docx:
            length_docx[ln] = round(length_docx[ln], 3)
        # self.length_docx = length_docx
        return length_docx

    def _check_format(self):
        md_lines = self.md_lines
        is_first_line = True
        for ml in md_lines:
            if is_first_line:
                if re.match('^#+(-#)*$', ml.text):
                    if re.match('^\\s$', ml.end_space):
                        continue
            if re.match('^\\s+$', ml.end_space):
                msg = '※ 警告: ' \
                    + '行末に無意味な空白があります'
                # msg = 'warning: ' \
                #     + 'white spaces at the end of the line'
                ml.append_warning_message(msg)
            if ml.text != '':
                is_first_line = False
        if True:
            if re.match('^.*<br>$', md_lines[-1].text):
                msg = '※ 警告: ' \
                    + '最終行に無意味な改行があります'
                # msg = 'warning: ' \
                #     + 'breaking line at the end of the last line'
                ml.append_warning_message(msg)

    def _edit_data(self):
        md_lines = self.md_lines
        # REMOVE ESCAPING SYMBOL
        for ml in md_lines:
            ml.text = re.sub('^\\\\(\\s+)', '\\1', ml.text)
            ml.text = re.sub('(\\s+)\\\\$', '\\1', ml.text)

    def _edit_data_of_chapter_and_section(self):
        paragraph_class = self.paragraph_class
        paragraph_class_ja = self.paragraph_class_ja
        res = self.res_feature
        md_lines = self.md_lines
        if paragraph_class == 'chapter':
            char = '$'
            paragraph_depth = self.proper_depth
        elif paragraph_class == 'section':
            char = '#'
            paragraph_depth = self.tail_section_depth
        else:
            return
        head_strings = ''
        title = ''
        body = ''
        pdepth = -1
        is_in_body = False
        for ml in md_lines:
            mlt = ml.text
            if not is_in_body:
                while re.match(res, mlt):
                    trunk = re.sub(res, '\\1', mlt)
                    branc = re.sub(res, '\\2', mlt)
                    mlt = re.sub(res, '\\3', mlt)
                    mlt = re.sub('\\\\$', '', mlt)  # remove a escape symbol
                    xdepth = len(trunk) - 1
                    ydepth = len(branc.replace(char, ''))
                    if pdepth > 0 and xdepth != pdepth + 1:
                        msg = '※ 警告: ' + paragraph_class_ja \
                            + 'の深さが飛んでいます'
                        # msg = 'warning: ' + paragraph_class \
                        #     + ' depth is not continuous'
                        ml.append_warning_message(msg)
                    pdepth = xdepth
                    head_strings += self._get_head_string(xdepth, ydepth, ml)
                    self.__step_state(xdepth, ydepth, ml)
                if mlt != ml.text:
                    title = mlt
                    if re.match('^\\s+', title):
                        msg = '※ 警告: ' + paragraph_class_ja \
                            + 'のタイトルの最初に空白があります'
                        # msg = 'warning: ' + paragraph_class \
                        #     + ' title has spaces at the beginning'
                        ml.append_warning_message(msg)
                    ml.text = ''
                if mlt != '':
                    is_in_body = True
            if body == '' and re.match('^\\s+', ml.text):
                msg = '※ 警告: ' + paragraph_class_ja \
                    + 'の本文の最初に空白があります'
                # msg = 'warning: ' + paragraph_class \
                #     + ' body has spaces at the beginning'
                ml.append_warning_message(msg)
            body += ml.text
        if title + body == '':
            return
        if paragraph_class == 'section' and paragraph_depth == 1:
            md_lines[0].text = title
        elif re.match('^.*\\(.*\\)$', head_strings):
            md_lines[0].text = head_strings + ' ' + title
        else:
            md_lines[0].text = head_strings + '\u3000' + title
        return

    @classmethod
    def __step_state(cls, xdepth, ydepth, md_line):
        cls.__set_state(xdepth, ydepth, None, md_line)

    def _get_text_to_write(self):
        md_lines = self.md_lines
        text_to_write = ''
        for ml in md_lines:
            text_to_write = concatenate_text(text_to_write, ml.text)
        # self.text_to_write = text_to_write
        return text_to_write

    def _get_text_to_write_with_reviser(self):
        text_to_write = self.text_to_write
        head_font_revisers = self.head_font_revisers
        tail_font_revisers = self.tail_font_revisers
        text_to_write_with_reviser \
            = ''.join(head_font_revisers) \
            + text_to_write \
            + ''.join(tail_font_revisers)
        # self.text_to_write_with_reviser = text_to_write_with_reviser
        return text_to_write_with_reviser

    def write_paragraph(self, ms_doc):
        # CHARS STATE
        self.beg_chars_state = Paragraph.bridge_chars_state.copy()
        self.chars_state = self.beg_chars_state.copy()
        self.chars_state.char_spacing = self.char_spacing
        paragraph_class = self.paragraph_class
        tail_section_depth = self.tail_section_depth
        alignment = self.alignment
        md_lines = self.md_lines
        chars_state = self.chars_state
        text_to_write_with_reviser = self.text_to_write_with_reviser
        if text_to_write_with_reviser == '':
            return
        if paragraph_class == 'alignment':
            ms_par = self.__get_ms_par(ms_doc)
            # WORD WRAP (英単語の途中で改行する)
            ms_ppr = ms_par._p.get_or_add_pPr()
            XML.add_tag(ms_ppr, 'w:wordWrap', {'w:val': '0'})
        elif paragraph_class == 'preformatted':
            ms_par = self.__get_ms_par(ms_doc, 'makdo-g')
        else:
            ms_par = self.__get_ms_par(ms_doc)
        if alignment == 'left':
            ms_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == 'center':
            ms_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            ms_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif (paragraph_class == 'section' and
              re.match('^\\S*\\s*$', md_lines[0].text) and
              not re.match('^.*<br>', text_to_write_with_reviser)):
            ms_par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif (paragraph_class == 'sentence' and
              not re.match('^.*<br>', text_to_write_with_reviser)):
            ms_par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ms_fmt = ms_par.paragraph_format
        if paragraph_class == 'section' and tail_section_depth == 1:
            chars_state.font_scale = 1.4
            self.write_text(ms_par, chars_state, text_to_write_with_reviser)
            chars_state.font_scale = 1.0
        else:
            self.write_text(ms_par, chars_state, text_to_write_with_reviser)
        # CHARS STATE
        self.end_chars_state = self.chars_state.copy()
        Paragraph.bridge_chars_state = self.end_chars_state.copy()

    def __get_ms_par(self, ms_doc, par_style='makdo'):
        length_docx = self.length_docx
        f_size = Form.font_size
        ms_par = ms_doc.add_paragraph(style=par_style)
        if not Form.auto_space:
            ms_ppr = ms_par._p.get_or_add_pPr()
            # KANJI<->ENGLISH
            XML.add_tag(ms_ppr, 'w:autoSpaceDE', {'w:val': '0'})
            # KANJI<->NUMBER
            XML.add_tag(ms_ppr, 'w:autoSpaceDN', {'w:val': '0'})
        ms_fmt = ms_par.paragraph_format
        ms_fmt.widow_control = False
        if length_docx['space before'] >= 0:
            pt = length_docx['space before'] * Form.line_spacing * f_size
            ms_fmt.space_before = Pt(pt)
        else:
            ms_fmt.space_before = Pt(0)
            msg = '※ 警告: ' \
                + '段落前の余白「v」の値が小さ過ぎます'
            # msg = 'warning: ' \
            #     + '"space before" is too small'
            self.md_lines[0].append_warning_message(msg)
        if length_docx['space after'] >= 0:
            pt = length_docx['space after'] * Form.line_spacing * f_size
            ms_fmt.space_after = Pt(pt)
        else:
            ms_fmt.space_after = Pt(0)
            msg = '※ 警告: ' \
                + '段落後の余白「V」の値が小さ過ぎます'
            # msg = 'warning: ' \
            #     + '"space after" is too small'
            self.md_lines[0].append_warning_message(msg)
        ms_fmt.first_line_indent = Pt(length_docx['first indent'] * f_size)
        ms_fmt.left_indent = Pt(length_docx['left indent'] * f_size)
        ms_fmt.right_indent = Pt(length_docx['right indent'] * f_size)
        # ms_fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        ls = Form.line_spacing * (1 + length_docx['line spacing'])
        if ls >= 1.0:
            ms_fmt.line_spacing = Pt(ls * f_size)
        else:
            ms_fmt.line_spacing = Pt(1.0 * f_size)
            msg = '※ 警告: ' \
                + '行間隔「X」の値が少な過ぎます'
            # msg = 'warning: ' \
            #     + 'too small line spacing'
            self.md_lines[0].append_warning_message(msg)
        ms_fmt.line_spacing = Pt(ls * f_size)
        return ms_par

    def write_text(self, ms_par, chars_state, text, type='normal'):
        text = self.__replace_br_tag(text)
        chars = ''
        for c in text + '\0':
            if not self.__must_continue(chars, c):
                chars = self.__write_chars(ms_par, chars_state, chars, c, type)
            if c != '\0':
                chars += c
        return chars

    @staticmethod
    def __replace_br_tag(text):
        tmp = ''
        res = NOT_ESCAPED + '<br/?>'
        while re.match(res, text):
            text = re.sub(res, '\\1\n', text)
        return text

    @staticmethod
    def __must_continue(tex, c):
        # RELAX
        if re.match(NOT_ESCAPED + RELAX_SYMBOL + '$', tex):
            return True      # "...<>"
        # MATH
        if re.match('^\\\\\\[', tex):
            if not re.match(NOT_ESCAPED + '\\\\\\]$', tex):
                return True  # "\[..."
        # SUB OR SUP
        if re.match('^(_|\\^){', tex):
            if not re.match(NOT_ESCAPED + '}$', tex):
                return True  # "_{..."|"^{..."
            t, d = '', 0
            for c in tex:
                t += c
                d += 1 if re.match(NOT_ESCAPED + '{$', t) else 0
                d -= 1 if re.match(NOT_ESCAPED + '}$', t) else 0
            if d != 0:
                return True  # "_{...{...}..."|"^{...{...}..."
        # ITALIC AND BOLD
        if re.match(NOT_ESCAPED + '\\*$', tex):
            if not re.match(NOT_ESCAPED + '\\*\\*\\*$', tex) and c == '*':
                return True  # "...*" + "*"
        # SMALL
        if re.match(NOT_ESCAPED + '\\-$', tex):
            if not re.match(NOT_ESCAPED + '\\-\\-\\-$', tex) and c == '-':
                return True  # "...-" + "-"
        # LARGE
        if re.match(NOT_ESCAPED + '\\+$', tex):
            if not re.match(NOT_ESCAPED + '\\+\\+\\+$', tex) and c == '+':
                return True  # "...+" + "+"
        # RELAX AND NARROW
        if re.match(NOT_ESCAPED + '>$', tex):
            if re.match(NOT_ESCAPED + RELAX_SYMBOL + '>$', tex):
                return True  # "...<>>"
            if re.match(NOT_ESCAPED + RELAX_SYMBOL + '>>$', tex) and c == '>':
                return True  # "...<>>>" + '>'
        # NARROW
        if re.match(NOT_ESCAPED + '>$', tex):
            if not re.match(NOT_ESCAPED + '>>>$', tex) and c == '>':
                return True  # "...>" + '>'
        # WIDE
        if re.match(NOT_ESCAPED + '<$', tex):
            if not re.match(NOT_ESCAPED + '<<<$', tex) and c == '<':
                return True  # "...<" + '<'
        # ELSE
        return False

    def __write_chars(self, ms_par, chars_state, chars, c, type='normal'):
        res_ivs = '^((?:.|\n)*?)([^0-9\\\\])([0-9]+);$'
        res_foc = NOT_ESCAPED + '\\^([0-9A-Za-z]{0,11})\\^$'
        res_hlc = NOT_ESCAPED + '_([0-9A-Za-z]{1,11})_$'
        if False:
            pass
        elif re.match(NOT_ESCAPED + '<$', chars) and c == '>':
            # '<>' （RELAX） "<<<" (+ ">") = "<<" + "<" (+ ">")
            chars = re.sub('<$', '', chars)
            chars = self.__write_chars(ms_par, chars_state, chars, '', type)
            chars += '<'
        elif re.match(NOT_ESCAPED + '\\\\\\[$', chars):
            # "\[" (BEGINNING OF MATH EXPRESSION) (MUST FIRST)
            chars = re.sub('\\\\\\[$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars = '\\['
        elif (re.match('^\\\\\\[', chars) and
              re.match(NOT_ESCAPED + '\\\\\\]$', chars)):
            # "\]" (END OF MATH EXPRESSION (MUST FIRST)
            chars = re.sub('^\\\\\\[(.*)\\\\\\]$', '\\1', chars)
            is_italic = chars_state.is_italic
            chars_state.is_italic = True
            chars = Math.write_chars(ms_par._p, chars_state, chars)
            chars_state.is_italic = is_italic
        elif re.match(NOT_ESCAPED + '((?:_|\\^){)$', chars):
            # "_{" or "^{" (BEGINNIG OF SUB OR SUP)
            subp = re.sub(NOT_ESCAPED + '((?:_|\\^){)$', '\\2', chars)
            chars = re.sub(NOT_ESCAPED + '((?:_|\\^){)$', '\\1', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars = subp
        elif (re.match('^(?:_|\\^){', chars) and
              re.match(NOT_ESCAPED + '}$', chars)):
            # "}" (END OF SUB OR SUP)
            if re.match('^_{', chars):
                chars_state.apply_sub_or_sup_font_decorator('_{')
            else:
                chars_state.apply_sub_or_sup_font_decorator('^{')
            chars = re.sub('^(?:_|\\^){(.*)}', '\\1', chars)
            chars = self.write_text(ms_par, chars_state, chars, type)
            chars_state.apply_sub_or_sup_font_decorator('}')
        elif re.match(NOT_ESCAPED + RES_IMAGE, chars):
            # "![.*](.+)" (IMAGE)
            path = re.sub(NOT_ESCAPED + RES_IMAGE, '\\3', chars)
            alte = re.sub(NOT_ESCAPED + RES_IMAGE, '\\2', chars)
            chars = re.sub(NOT_ESCAPED + RES_IMAGE, '\\1', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            self.__write_image(ms_par, chars_state, alte, path)
        elif re.match(NOT_ESCAPED + '\\*\\*\\*$', chars):
            # "***" (ITALIC AND BOLD)
            chars = re.sub('\\*\\*\\*$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_is_italic_font_decorator('***')
            chars_state.apply_is_bold_font_decorator('***')
        elif re.match(NOT_ESCAPED + '\\*\\*$', chars):
            # "**" BOLD
            chars = re.sub('\\*\\*$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_is_bold_font_decorator('**')
        elif re.match(NOT_ESCAPED + '\\*$', chars):
            # "*" ITALIC
            chars = re.sub('\\*$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_is_italic_font_decorator('*')
        elif re.match(NOT_ESCAPED + '~~$', chars):
            # "~~" (STRIKETHROUGH)
            chars = re.sub('~~$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_has_strike_font_decorator('~~')
        elif re.match(NOT_ESCAPED + '(?:\\[\\||\\|\\])$', chars):
            # "[|" or "|]" (FRAME)
            if re.match(NOT_ESCAPED + '\\[\\|$', chars):
                fd = '[|'
            elif re.match(NOT_ESCAPED + '\\|\\]$', chars):
                fd = '|]'
            chars = re.sub('(?:\\[\\||\\|\\])$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_has_frame_font_decorator(fd)
        elif re.match(NOT_ESCAPED + '`$', chars):
            # "`" (PREFORMATTED)
            chars = re.sub('`$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_is_preformatted_font_decorator('`')
        elif re.match(NOT_ESCAPED + '//$', chars):
            # "//" (ITALIC)
            if not re.match('^.*[a-z]+://$', chars):
                # not http:// https:// ftp:// ...
                chars = re.sub('//$', '', chars)
                chars = XML.write_chars(ms_par._p, chars_state, chars)
                chars_state.apply_is_italic_font_decorator('//')
        elif re.match(NOT_ESCAPED + '_([\\$=\\.#\\-~\\+]{,4})_$', chars):
            # "_.*_" (UNDERLINE)
            sty = re.sub(NOT_ESCAPED + '_([\\$=\\.#\\-~\\+]{,4})_$', '\\2',
                         chars)
            if sty in UNDERLINE:
                chars = re.sub('_([\\$=\\.#\\-~\\+]{,4})_$', '', chars, 1)
                chars = XML.write_chars(ms_par._p, chars_state, chars)
                chars_state.apply_underline_font_decorator('_' + sty + '_')
        elif re.match(NOT_ESCAPED + '\\-\\-\\-$', chars):
            # "---" (XSMALL)
            chars = re.sub('\\-\\-\\-$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            if chars_state.font_scale == 0.8:
                chars_state.apply_font_scale_font_decorator('--')
                chars += '-'
            else:
                chars_state.apply_font_scale_font_decorator('---')
        elif re.match(NOT_ESCAPED + '\\-\\-$', chars):
            # "--" (SMALL)
            chars = re.sub('\\-\\-$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_font_scale_font_decorator('--')
        elif re.match(NOT_ESCAPED + '\\+\\+\\+$', chars):
            # "+++" (XLARGE)
            chars = re.sub('\\+\\+\\+$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            if chars_state.font_scale == 1.2:
                chars_state.apply_font_scale_font_decorator('++')
                chars += '+'
            else:
                chars_state.apply_font_scale_font_decorator('+++')
        elif re.match(NOT_ESCAPED + '\\+\\+$', chars):
            # "++" (LARGE)
            chars = re.sub('\\+\\+$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_font_scale_font_decorator('++')
        elif re.match(NOT_ESCAPED + '<<<$', chars):
            # "<<<" (XWIDE or RESET)
            if c == '>':
                # "<<<" + ">" = "<<" + "<>"
                chars = re.sub('<<<$', '', chars)
                chars = XML.write_chars(ms_par._p, chars_state, chars)
                chars += '<'
                chars_state.apply_font_width_font_decorator('<<')
            else:
                chars = re.sub('<<<$', '', chars)
                chars = XML.write_chars(ms_par._p, chars_state, chars)
                if chars_state.font_width == 0.8:
                    chars_state.apply_font_width_font_decorator('<<')
                    chars += '<'
                else:
                    chars_state.apply_font_width_font_decorator('<<<')
        elif re.match(NOT_ESCAPED + '<<$', chars):
            chars = re.sub('<<$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_font_width_font_decorator('<<')
        elif re.match(NOT_ESCAPED + '>>>$', chars):
            # ">>>" (XNARROW or RESET)
            if re.match(NOT_ESCAPED + RELAX_SYMBOL + '>>$', chars):
                # "<>>>" + "." = "<>" + ">>" + "."
                chars = re.sub('>>$', '', chars)
                chars = XML.write_chars(ms_par._p, chars_state, chars)
                chars_state.apply_font_width_font_decorator('>>')
            else:
                chars = re.sub('>>>$', '', chars)
                chars = XML.write_chars(ms_par._p, chars_state, chars)
                if chars_state.font_width == 1.2:
                    chars_state.apply_font_width_font_decorator('>>')
                    chars += '>'
                else:
                    chars_state.apply_font_width_font_decorator('>>>')
        elif re.match(NOT_ESCAPED + '>>$', chars):
            # ">>" (NARROW or RESET)
            chars = re.sub('>>$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_font_width_font_decorator('>>')
        elif re.match(NOT_ESCAPED + '@' + RES_NUMBER + '@$', chars):
            # "@.+@" (FONT SCALE)
            c_size = re.sub(NOT_ESCAPED + '@([^@]+)@$', '\\2', chars)
            chars = re.sub(NOT_ESCAPED + '@([^@]+)@$', '\\1', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_font_scale_font_decorator('@' + c_size + '@')
        elif re.match(NOT_ESCAPED + '@([^@]{1,66})@$', chars):
            # "@.+@" (FONT)
            font = re.sub(NOT_ESCAPED + '@([^@]{1,66})@$', '\\2', chars)
            chars = re.sub(NOT_ESCAPED + '@([^@]{1,66})@$', '\\1', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_font_name_font_decorator('@' + font + '@')
        elif re.match(res_ivs, chars):
            # .[0-9]+; (IVS (IDEOGRAPHIC VARIATION SEQUENCE))
            ivsn = re.sub(res_ivs, '\\3', chars)
            ivsc = re.sub(res_ivs, '\\2', chars)
            chars = re.sub(res_ivs, '\\1', chars)
            ivsu = int('0xE0100', 16) + int(ivsn)
            if int(ivsu) <= int('0xE01EF', 16):
                chars = XML.write_chars(ms_par._p, chars_state, chars)
                is_mincho_font = False
                if chars_state.mincho_font == Form.mincho_font:
                    is_mincho_font = True
                    chars_state.mincho_font = chars_state.ivs_font
                chars \
                    = XML.write_chars(ms_par._p, chars_state, ivsc + chr(ivsu))
                if is_mincho_font:
                    chars_state.mincho_font = Form.mincho_font
        elif re.match(res_foc, chars):
            # "^.*^" (FONT COLOR)
            col = re.sub(res_foc, '\\2', chars)
            if col == '':
                col = 'FFFFFF'
            elif re.match('^([0-9A-F])([0-9A-F])([0-9A-F])$', col):
                col = re.sub('^([0-9A-F])([0-9A-F])([0-9A-F])$',
                             '\\1\\1\\2\\2\\3\\3', col)
            elif col in FONT_COLOR:
                col = FONT_COLOR[col]
            if re.match('^[0-9A-F]{6}$', col):
                chars = re.sub('\\^([0-9A-Za-z]*)\\^$', '', chars)
                chars = XML.write_chars(ms_par._p, chars_state, chars)
                chars_state.apply_font_color_font_decorator('^' + col + '^')
        elif re.match(res_hlc, chars):
            # "_.+_" (HIGHLIGHT COLOR)
            col = re.sub(res_hlc, '\\2', chars)
            if col in HIGHLIGHT_COLOR:
                chars = re.sub(res_hlc, '\\1', chars)
                chars = XML.write_chars(ms_par._p, chars_state, chars)
                fd = '_' + col + '_'
                chars_state.apply_highlight_color_font_decorator(fd)
        elif re.match(NOT_ESCAPED + '\\->$', chars):
            # "->" (BEGINNING OF DELETED)
            chars = re.sub('\\->$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_track_changes_font_decorator('->')
        elif re.match(NOT_ESCAPED + '<\\-$', chars):
            # "<-" (END OF DELETED)
            chars = re.sub('<\\-$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_track_changes_font_decorator('<-')
        elif re.match(NOT_ESCAPED + '\\+>$', chars):
            # "+>" (BEGINNING OF INSERTED)
            chars = re.sub('\\+>$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_track_changes_font_decorator('+>')
        elif re.match(NOT_ESCAPED + '<\\+$', chars):
            # "<+" (END OF INSERTED)
            chars = re.sub('<\\+$', '', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            chars_state.apply_track_changes_font_decorator('<+')
        elif re.match(NOT_ESCAPED + '<([^<>]{1,37}?)/([^<>]{1,37}?)>$', chars):
            # "<.+/.+>" (RUBY)
            res = '<([^<>]{1,37}?)/([^<>]{1,37}?)>'
            ruby = re.sub(NOT_ESCAPED + res + '$', '\\3', chars)
            base = re.sub(NOT_ESCAPED + res + '$', '\\2', chars)
            chars = re.sub(NOT_ESCAPED + res + '$', '\\1', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            ms_rb0 = XML.add_tag(ms_par._p, 'w:r', {})
            ms_rb1 = XML.add_tag(ms_rb0, 'w:ruby', {})
            ms_rb2 = XML.add_tag(ms_rb1, 'w:rubyPr', {})
            XML.add_tag(ms_rb2, 'w:rubyAlign', {'w:val': 'center'})
            ms_rb2 = XML.add_tag(ms_rb1, 'w:rt', {})
            chars_state.font_size /= 2
            XML.write_chars(ms_rb2, chars_state, ruby)
            ms_rb2 = XML.add_tag(ms_rb1, 'w:rubyBase', {})
            chars_state.font_size *= 2
            XML.write_chars(ms_rb2, chars_state, base)
        elif re.match(NOT_ESCAPED + '< *((?:[0-9]*\\.)?[0-9]+) *>$', chars):
            # "< *([0-9]*.)?[0-9]+ *>" (SPACE)
            res = '< *((?:[0-9]*\\.)?[0-9]+) *>'
            spac = re.sub(NOT_ESCAPED + res + '$', '\\2', chars)
            chars = re.sub(NOT_ESCAPED + res + '$', '\\1', chars)
            chars = XML.write_chars(ms_par._p, chars_state, chars)
            ori_fw = chars_state.font_width
            tmp_fw = float(spac) * chars_state.font_width
            if tmp_fw >= 0.01:
                if tmp_fw >= 5.00:
                    n = int(tmp_fw / 5.00)
                    chars_state.font_width = 5.00
                    XML.write_chars(ms_par._p, chars_state, '\u3000' * n)
                    tmp_fw -= 5.00 * n
                if tmp_fw >= 0.01:
                    chars_state.font_width = tmp_fw
                    XML.write_chars(ms_par._p, chars_state, '\u3000')
                chars_state.font_width = ori_fw
        elif re.match(NOT_ESCAPED + '(n|N|M)$', chars):
            if type == 'footer':
                # "n|N|M" (PAGE NUMBER)
                char = re.sub(NOT_ESCAPED + '(n|N|M)$', '\\2', chars)
                chars = re.sub(NOT_ESCAPED + '(n|N|M)$', '\\1', chars)
                chars = XML.write_chars(ms_par._p, chars_state, chars)
                chars += XML.write_page_number(ms_par._p, chars_state, char)
        if c == '\0' and chars != '':
            # LAST
            chars = XML.write_chars(ms_par._p, chars_state, chars)
        return chars

    def __write_image(self, ms_par, chars_state, alte, path):
        c_size = chars_state.font_size * chars_state.font_scale
        indent \
            = self.length_docx['first indent'] \
            + self.length_docx['left indent'] \
            + self.length_docx['right indent']
        text_width = PAPER_WIDTH[Form.paper_size] \
            - Form.left_margin - Form.right_margin \
            - (indent * 2.54 / 72)
        text_height = PAPER_HEIGHT[Form.paper_size] \
            - Form.top_margin - Form.bottom_margin
        ms_run = ms_par.add_run()
        res = '^(.*):(' + RES_NUMBER + ')?(?:x(' + RES_NUMBER + ')?)?$'
        cm_w = 0
        cm_h = 0
        if re.match(res, alte):
            st_w = re.sub(res, '\\2', alte)
            if st_w != '':
                cm_w = float(st_w)
                if cm_w < 0:
                    cm_w = text_width * (-1 * cm_w)
            st_h = re.sub(res, '\\3', alte)
            if st_h != '':
                cm_h = float(st_h)
                if cm_h < 0:
                    cm_h = text_height * (-1 * cm_h)
            alte = re.sub(res, '\\1', alte)
        try:
            if cm_w > 0 and cm_h > 0:
                ms_run.add_picture(path, width=Cm(cm_w), height=Cm(cm_h))
            elif cm_w > 0:
                ms_run.add_picture(path, width=Cm(cm_w))
            elif cm_h > 0:
                ms_run.add_picture(path, height=Cm(cm_h))
            else:
                ms_run.add_picture(path, height=Pt(c_size))
        except BaseException:
            ms_run.text = '![' + alte + '](' + path + ')'
            msg = '※ 警告: ' \
                + 'インライン画像「' + path + '」が読み込めません'
            # msg = 'warning: can\'t open "' + path + '"'
            r = '^.*! *\\[.*\\] *\\(' + path + '\\).*$'
            for ml in self.md_lines:
                if re.match(r, ml.text):
                    if msg not in ml.warning_messages:
                        ml.append_warning_message(msg)
                        break
            else:
                self.md_lines[0].append_warning_message(msg)

    def print_warning_messages(self):
        for ml in self.md_lines:
            ml.print_warning_messages()


class ParagraphEmpty(Paragraph):

    """A class to handle empty paragraph"""

    paragraph_class = 'empty'
    res_feature = '^$'


class ParagraphBlank(Paragraph):

    """A class to handle blank paragraph"""

    paragraph_class = 'blank'
    res_feature = '^\n( \n)*$'


class ParagraphChapter(Paragraph):

    """A class to handle chapter paragraph"""

    paragraph_class = 'chapter'
    paragraph_class_ja = 'チャプター'
    res_symbol = '(\\$+)((?:\\-\\$+)*)'
    res_feature = '^' + res_symbol + '(?:\\s((?:.|\n)*))?$'
    # SPACE POLICY
    # res_feature = '^' + res_symbol + '(?:\\s+((?:.|\n)*))?$'
    res_reviser = res_symbol + '=([0-9]+)'
    states = [[0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # 第１編
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # 第１章
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # 第１節
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # 第１款
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0]]  # 第１目
    unit_chars = ['編', '章', '節', '款', '目']

    @classmethod
    def _get_proper_depth(cls, full_text):
        if not re.match(cls.res_feature, full_text):
            return 0
        trunk = re.sub(cls.res_feature, '\\1', full_text)
        proper_depth = len(trunk)
        # self.proper_depth = proper_depth
        return proper_depth

    def _edit_data(self):
        self._edit_data_of_chapter_and_section()
        return

    @classmethod
    def _get_head_string(cls, xdepth, ydepth, md_line):
        xvalue_char = '〓'
        unit_char = '〓'
        if xdepth < len(cls.states):
            if ydepth < len(cls.states[xdepth]):
                value = cls.states[xdepth][0]
                if ydepth == 0:
                    value += 1
                xvalue_char = n2c_n_arab(value, md_line)
            unit_char = cls.unit_chars[xdepth]
        head_string = '第' + xvalue_char + unit_char
        for y in range(1, ydepth + 1):
            if y < len(cls.states[xdepth]):
                value = cls.states[xdepth][y] + 1
                if y == ydepth:
                    value += 1
                yvalue_char = n2c_n_arab(value, md_line)
            else:
                yvalue_char = '〓'
            head_string += 'の' + yvalue_char
        return head_string


class ParagraphSection(Paragraph):

    """A class to handle section paragraph"""

    paragraph_class = 'section'
    paragraph_class_ja = 'セクション'
    res_symbol = '(#+)((?:\\-#+)*)'
    res_feature = '^' + res_symbol + '(?:\\s((?:.|\n)*))?$'
    # SPACE POLICY
    # res_feature = '^' + res_symbol + '(?:\\s+((?:.|\n)*))?$'
    res_reviser = res_symbol + '=([0-9]+)'
    states = [[0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # -
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # 第１
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # １
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # (1)
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # ア
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # (ｱ)
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0],  # ａ
              [0, 0, 0, 0, 0, 0, 0, 0, 0, 0]]  # (a)

    @classmethod
    def is_this_class(cls, full_text,
                      head_font_revisers=[], tail_font_revisers=[]):
        if re.match(cls.res_feature, full_text):
            if not re.match('^#{15,}', full_text):
                return True
        return False

    @classmethod
    def _get_section_depths(cls, full_text):
        ft = full_text
        head_section_depth = 0
        tail_section_depth = 0
        while re.match(cls.res_feature, ft):
            trunk = re.sub(cls.res_feature, '\\1', ft)
            depth = len(trunk)
            if head_section_depth == 0:
                head_section_depth = depth
            tail_section_depth = depth
            ft = re.sub(cls.res_feature, '\\3', ft)
        Paragraph.previous_head_section_depth = head_section_depth
        Paragraph.bridge_tail_section_depth = tail_section_depth
        # self.head_section_depth = head_section_depth
        # self.head_section_depth = tail_section_depth
        return head_section_depth, tail_section_depth

    def _edit_data(self):
        self._edit_data_of_chapter_and_section()
        return

    @classmethod
    def _get_head_string(cls, xdepth, ydepth, md_line):
        # TRUNK
        if xdepth < len(cls.states):
            value = cls.states[xdepth][0]
            if ydepth == 0:
                value += 1
            if xdepth == 0:
                head_string = ''
            elif xdepth == 1:
                if Form.document_style == 'n':
                    head_string = '第' + n2c_n_arab(value, md_line)
                else:
                    head_string = '第' + n2c_n_arab(value, md_line) + '条'
            elif xdepth == 2:
                if Form.document_style != 'j' or cls.states[1][0] == 0:
                    head_string = n2c_n_arab(value, md_line)
                else:
                    head_string = n2c_n_arab(value + 1, md_line)
            elif xdepth == 3:
                head_string = n2c_p_arab(value, md_line)
            elif xdepth == 4:
                head_string = n2c_n_kata(value, md_line)
            elif xdepth == 5:
                head_string = n2c_p_kata(value, md_line)
            elif xdepth == 6:
                head_string = n2c_n_alph(value, md_line)
            elif xdepth == 7:
                head_string = n2c_p_alph(value, md_line)
            else:
                head_string = '〓'
        else:
            head_string = '〓'
        # BRANCH
        for y in range(1, ydepth + 1):
            if y < len(cls.states[xdepth]):
                value = cls.states[xdepth][y] + 1
                if y == ydepth:
                    value += 1
                yvalue_char = n2c_n_arab(value, md_line)
            else:
                yvalue_char = '〓'
            head_string += 'の' + yvalue_char
        return head_string


class ParagraphList(Paragraph):

    """A class to handle list paragraph"""

    paragraph_class = 'list'
    paragraph_class_ja = 'リスト'
    res_symbol = '(\\-|\\+|[0-9]+\\.|[0-9]+\\))()'
    res_feature = '^\\s*' + res_symbol + '\\s(.*)$'
    # SPACE POLICY
    # res_feature = '^\\s*' + res_symbol + '\\s+(.*)$'
    res_reviser = '\\s*(?:[0-9]+\\.|[0-9]+\\))=([0-9]+)'
    states = [[0],  # ①
              [0],  # ㋐
              [0],  # ⓐ
              [0]]  # ㊀

    @classmethod
    def _get_section_depths(cls, full_text):
        head_section_depth = Paragraph.bridge_tail_section_depth
        tail_section_depth = Paragraph.bridge_tail_section_depth
        # self.head_section_depth = head_section_depth
        # self.tail_section_depth = tail_section_depth
        return head_section_depth, tail_section_depth

    @classmethod
    def _get_proper_depth(cls, full_text):
        full_text = re.sub('\u3000', '  ', full_text)
        full_text = re.sub('\t', '  ', full_text)
        full_text = re.sub('  ', ' ', full_text)
        spaces = re.sub('^( *).*$', '\\1', full_text)
        proper_depth = len(spaces) + 1
        # self.proper_depth = proper_depth
        return proper_depth

    @classmethod
    def _apply_revisers(cls, revisers, md_lines):
        for rev in revisers:
            res = '(\\s*).*=([0-9]+)'
            if re.match(res, rev):
                str_d = re.sub(res, '\\1', rev)
                str_v = re.sub(res, '\\2', rev)
                depth = len(re.sub('\\s\\s', ' ', str_d))
                value = int(str_v)
                cls.states[depth][0] = value - 1
                for d in range(depth + 1, len(cls.states)):
                    cls.states[d][0] = 0

    def _edit_data(self):
        full_text = self.full_text
        md_lines = self.md_lines
        res = '^\\s*' + ParagraphList.res_symbol + '\\s*'
        states = ParagraphList.states
        proper_depth = self.proper_depth
        n = 0
        while n < len(md_lines) and md_lines[n].text == '':
            n += 1
        line = md_lines[n].text
        is_numbering = False
        if re.match('\\s*[0-9]+(?:\\.|\\))\\s', full_text):
            is_numbering = True
        line = re.sub(res, '', line)
        if not is_numbering:
            if proper_depth == 1:
                head_strings = '・'
                # head_strings = '• '  # U+2022 Bullet
            elif proper_depth == 2:
                head_strings = '○'
                # head_strings = '◦ '  # U+25E6 White Bullet
            elif proper_depth == 3:
                head_strings = '△'
                # head_strings = '‣ '  # U+2023 Triangular Bullet
            elif proper_depth == 4:
                head_strings = '◇'
                # head_strings = '⁃ '  # U+2043 Hyphen Bullet
            else:
                head_strings = '〓'
        else:
            if proper_depth == 1:
                head_strings = n2c_c_arab(states[0][0] + 1, md_lines[n])
            elif proper_depth == 2:
                head_strings = n2c_c_kata(states[1][0] + 1, md_lines[n])
            elif proper_depth == 3:
                head_strings = n2c_c_alph(states[2][0] + 1, md_lines[n])
            elif proper_depth == 4:
                head_strings = n2c_c_kanj(states[3][0] + 1, md_lines[n])
            else:
                head_strings = '〓'
            if proper_depth <= len(states):
                states[proper_depth - 1][0] += 1
                for d in range(proper_depth, len(states)):
                    states[d][0] = 0
        self.md_lines[n].text = head_strings + '\u3000' + line

    @classmethod
    def _reset_states(cls, paragraph_class):
        if paragraph_class != 'list':
            for s in cls.states:
                s[0] = 0
        return


class ParagraphTable(Paragraph):

    """A class to handle table paragraph"""

    paragraph_class = 'table'
    res_feature = '^(?::\\s+)?\\|.*\\|(:?-*:?)?(\\^+|=+)?(?:\\s+:)?$'

    def write_paragraph(self, ms_doc):
        # CHARS STATE
        self.beg_chars_state = Paragraph.bridge_chars_state.copy()
        self.chars_state = self.beg_chars_state.copy()
        self.chars_state.char_spacing = self.char_spacing
        self.chars_state.apply_font_decorators(self.head_font_revisers)
        md_lines = self.md_lines
        length_docx = self.length_docx
        chars_state = self.chars_state
        c_size = chars_state.font_size * chars_state.font_scale
        # GET DATA
        tab_lines = self.__get_tab_lines(md_lines)
        tab, conf_line_place, table_alignment, \
            col_alig_list, col_widt_list, col_rule_list, \
            row_alig_list, row_heig_list, row_rule_list \
            = self.__get_tab_data(tab_lines)
        cal, cwl, crl = col_alig_list, col_widt_list, col_rule_list
        tab, col_alig_mtrx, col_widt_mtrx, col_rule_mtrx \
            = self.__get_col_data(tab, conf_line_place, cal, cwl, crl)
        ral, rhl, rrl = row_alig_list, row_heig_list, row_rule_list
        tab, row_alig_mtrx, row_heig_mtrx, row_rule_mtrx \
            = self.__get_row_data(tab, conf_line_place, ral, rhl, rrl)
        # hori_alig_list = col_alig_list
        hori_leng_list = col_widt_list
        hori_rule_list = col_rule_list
        # vert_alig_list = row_alig_list
        vert_leng_list = row_heig_list
        vert_rule_list = row_rule_list
        hori_alig_mtrx = col_alig_mtrx
        # hori_leng_mtrx = col_widt_mtrx
        # hori_rule_mtrx = col_rule_mtrx
        vert_alig_mtrx = row_alig_mtrx
        # vert_leng_mtrx = row_heig_mtrx
        # vert_rule_mtrx = row_rule_mtrx
        # MAKE TABLE
        row = len(tab)
        col = len(tab[0])
        ms_tab = ms_doc.add_table(row, col, style='Table Grid')
        if self.length_docx['left indent'] == 0:
            ms_tab.alignment = table_alignment
        else:
            tblind = int(self.length_docx['left indent'] * Form.font_size * 20)
            tblpr = ms_tab._element.xpath('w:tblPr')
            oe = OxmlElement('w:tblInd')
            oe.set(ns.qn('w:w'), str(tblind))
            oe.set(ns.qn('w:type'), 'dxa')  # 1 dxa = 1/1440 inch
            tblpr[0].append(oe)
        # ms_tab.autofit = True
        # SET CELL LENGTH AND ALIGNMENT
        for i in range(len(tab)):
            # ms_tab.rows[i].height_rule = WD_ROW_HEIGHT_RULE.AUTO
            if vert_leng_list[i] == 0:
                ms_tab.rows[i].height = Pt(c_size * BASIC_TABLE_CELL_HEIGHT)
            elif vert_leng_list[i] > 0:
                ms_tab.rows[i].height = Pt(c_size * vert_leng_list[i])
        for j in range(len(tab[0])):
            if hori_leng_list[j] >= 0:
                ms_tab.columns[j].width \
                    = Pt(c_size * (hori_leng_list[j] + BASIC_TABLE_CELL_WIDTH))
        # SET CELLS
        for i in range(len(tab)):
            for j in range(len(tab[i])):
                # CELL ALIGNMENT
                ms_cell = ms_tab.cell(i, j)
                ms_cell.horizontal_alignment = hori_alig_mtrx[i][j]
                ms_cell.vertical_alignment = vert_alig_mtrx[i][j]
                # FOR MS WORD
                if vert_leng_list[i] == 0:
                    ms_cell.height = Pt(c_size * BASIC_TABLE_CELL_HEIGHT)
                elif vert_leng_list[i] > 0:
                    ms_cell.height = Pt(c_size * vert_leng_list[i])
                ms_cell.width \
                    = Pt(c_size * (hori_leng_list[j] + BASIC_TABLE_CELL_WIDTH))
                ms_par = ms_cell.paragraphs[0]
                ms_par.style = 'makdo-t'
                # TEXT
                cell = tab[i][j]
                # WORD WRAP (英単語の途中で改行する)
                ms_ppr = ms_par._p.get_or_add_pPr()
                XML.add_tag(ms_ppr, 'w:wordWrap', {'w:val': '0'})
                cell = re.sub('^\\s+\\\\?', '', cell)
                cell = re.sub('\\\\?\\s+$', '', cell)
                self.write_text(ms_par, chars_state, cell)
                ms_fmt = ms_par.paragraph_format
                ms_fmt.alignment = hori_alig_mtrx[i][j]
                ls = TABLE_LINE_SPACING * (1 + length_docx['line spacing'])
                if ls >= 1.0:
                    ms_fmt.line_spacing = Pt(ls * c_size)
                else:
                    ms_fmt.line_spacing = Pt(1.0 * c_size)
                    msg = '※ 警告: ' \
                        + '行間隔「X」の値が少な過ぎます'
                    # msg = 'warning: ' \
                    #     + 'too small line spacing'
                    self.md_lines[0].append_warning_message(msg)
                # RULE
                ms_tcpr = ms_cell._tc.get_or_add_tcPr()
                ms_tcbr = XML.add_tag(ms_tcpr, 'w:tcBorders')
                if i > 0 and vert_rule_list[i - 1] == '^':
                    XML.add_tag(ms_tcbr, 'w:top', {'w:val': 'nil'})
                if vert_rule_list[i] == '^':
                    XML.add_tag(ms_tcbr, 'w:bottom', {'w:val': 'nil'})
                if i > 0 and vert_rule_list[i - 1] == '=':
                    XML.add_tag(ms_tcbr, 'w:top', {'w:val': 'double'})
                if vert_rule_list[i] == '=':
                    XML.add_tag(ms_tcbr, 'w:bottom', {'w:val': 'double'})
                if j > 0 and hori_rule_list[j - 1] == '^':
                    XML.add_tag(ms_tcbr, 'w:left', {'w:val': 'nil'})
                if hori_rule_list[j] == '^':
                    XML.add_tag(ms_tcbr, 'w:right', {'w:val': 'nil'})
                if j > 0 and hori_rule_list[j - 1] == '=':
                    XML.add_tag(ms_tcbr, 'w:left', {'w:val': 'double'})
                if hori_rule_list[j] == '=':
                    XML.add_tag(ms_tcbr, 'w:right', {'w:val': 'double'})
        # CHARS STATE
        self.chars_state.apply_font_decorators(self.tail_font_revisers)
        self.end_chars_state = self.chars_state.copy()
        Paragraph.bridge_chars_state = self.end_chars_state.copy()

    @staticmethod
    def __get_tab_lines(md_lines):
        tab_lines = []
        tab_line = ''
        for ml in md_lines:
            tab_line += re.sub('^\\s*', '', ml.text)
            if re.match(NOT_ESCAPED + '\\\\$', tab_line):
                tab_line = re.sub('\\s*\\\\$', '', tab_line)
                continue
            tab_lines.append(tab_line)
            tab_line = ''
        return tab_lines

    @staticmethod
    def __get_tab_data(tab_lines):
        tab = []
        conf_line_place = -1.0
        table_alignment = WD_TABLE_ALIGNMENT.CENTER
        col_alig_list, col_widt_list, col_rule_list = [], [], []
        row_alig_list, row_heig_list, row_rule_list = [], [], []
        res_config_row \
            = '^(?::\\s+)?(\\|\\s*(:?-*:?)?(\\^+|=+)?\\s*)+\\|(?:\\s+:)?$'
        for tl in tab_lines:
            if conf_line_place == -1 and \
               re.match(res_config_row, tl) and \
               not re.match('^\\|+$', tl):
                conf_line_place = float(len(tab)) - 0.5
                # TABLE ALIGNMENT
                if re.match('^:\\s+\\|.*\\|\\s+:$', tl):
                    tl = re.sub('^:\\s+', '', tl)
                    tl = re.sub('\\s+:$', '', tl)
                elif re.match('^:\\s+\\|.*$', tl):
                    tl = re.sub('^:\\s+', '', tl)
                    table_alignment = WD_TABLE_ALIGNMENT.LEFT
                elif re.match('^.*\\|\\s+:$', tl):
                    tl = re.sub('\\s+:$', '', tl)
                    table_alignment = WD_TABLE_ALIGNMENT.RIGHT
                tl = re.sub('^\\|(.*)\\|$', '\\1', tl)
                for c in tl.split('|'):
                    # COL RULE
                    if re.match('^.*\\^$', c):
                        col_rule_list.append('^')
                    elif re.match('^.*=$', c):
                        col_rule_list.append('=')
                    else:
                        col_rule_list.append('')
                    c = re.sub('(\\^+|=+)$', '', c)
                    # COL WIDTH
                    col_widt_list.append(float(len(c)) / 2)
                    # COL ALIGN
                    if c == '' or c == '-' or re.match('^:-+$', c):
                        col_alig_list.append(WD_TABLE_ALIGNMENT.LEFT)
                    elif c == ':' or re.match('^:-*:$', c):
                        col_alig_list.append(WD_TABLE_ALIGNMENT.CENTER)
                    elif re.match('^-+:$', c):
                        col_alig_list.append(WD_TABLE_ALIGNMENT.RIGHT)
                    else:
                        col_alig_list.append(WD_TABLE_ALIGNMENT.LEFT)
                continue
            if tl != '':
                if not re.match('^(:?-*:?)?(\\^+|=+)?$', tl):
                    row_rule_list.append('')
                    row_heig_list.append(0.0)
                    row_alig_list.append(WD_ALIGN_VERTICAL.CENTER)
                c = ''
                res = '^(.*?\\|)?\\s*((?::?-+|:-*:|-*:)?(?:\\^+|=+)?)\\s*\\|?$'
                if re.match(res, tl):
                    c = re.sub(res, '\\2', tl)
                    tl = re.sub(res, '\\1', tl)
                # ROW RULE
                if re.match('^.*\\^+$', c):
                    row_rule_list[-1] = '^'
                    c = re.sub('\\^+$', '', c)
                elif re.match('^.*=+$', c):
                    row_rule_list[-1] = '='
                    c = re.sub('=+$', '', c)
                # ROW HEIGHT
                if c != '':
                    row_heig_list[-1] = float(len(c)) / 2
                # ROW ALIGN
                if c == '-' or re.match('^:-+$', c):
                    row_alig_list[-1] = WD_ALIGN_VERTICAL.TOP
                elif c == ':' or re.match('^:-*:$', c):
                    row_alig_list[-1] = WD_ALIGN_VERTICAL.CENTER
                elif re.match('^-+:$', c):
                    row_alig_list[-1] = WD_ALIGN_VERTICAL.BOTTOM
            if tl != '':
                # TAB
                cells = []
                cell = ''
                tl = re.sub('^\\|', '', tl)
                tl = re.sub('\\|$', '', tl)
                for c in tl.split('|'):
                    cell += c
                    if re.match(NOT_ESCAPED + '\\\\$', cell):
                        cell = re.sub('\\\\$', '', cell) + '|'
                    else:
                        cells.append(cell)
                        cell = ''
                tab.append(cells)
        # FOR SHORTAGE
        max_row = 0
        for row in tab:
            if max_row < len(row):
                max_row = len(row)
        for row in tab:
            while len(row) < max_row:
                row.append('')
            while len(col_alig_list) < max_row:
                col_alig_list.append(WD_TABLE_ALIGNMENT.LEFT)
            while len(col_widt_list) < max_row:
                col_widt_list.append(0.0)
            while len(col_rule_list) < max_row:
                col_rule_list.append('')
        # WIDTH
        max_width = [0.0 for j in range(len(tab[0]))]
        for i in range(len(tab)):
            for j in range(len(tab[i])):
                line = ''
                for t in tab[i][j].split('<br>'):
                    if re.match(NOT_ESCAPED + '\\\\$', t):
                        line += re.sub('\\\\$', '', t) + '<br>'
                        continue
                    line += t
                    line = re.sub('^\\s*:\\s(.*)$', '\\1', line)
                    line = re.sub(NOT_ESCAPED + '\\s:\\s*$', '\\1', line)
                    for fd in FONT_DECORATORS + [RELAX_SYMBOL]:
                        while re.match(NOT_ESCAPED + fd, line):
                            line = re.sub(NOT_ESCAPED + fd, '\\1', line)
                    w = get_real_width(line) / 2
                    if max_width[j] < w:
                        max_width[j] = w
                    line = ''
        for j in range(len(col_alig_list)):
            if col_widt_list[j] == 0:
                col_widt_list[j] = max_width[j]
        # RETURN
        return tab, conf_line_place, table_alignment, \
            col_alig_list, col_widt_list, col_rule_list, \
            row_alig_list, row_heig_list, row_rule_list

    @staticmethod
    def __get_col_data(tab, conf_line_place,
                       col_alig_list, col_widt_list, col_rule_list):
        col_alig_mtrx, col_widt_mtrx, col_rule_mtrx = [], [], []
        for i in range(len(tab)):
            ca, cw, cr = [], [], []
            for j in range(len(tab[i])):
                ca.append(col_alig_list[j])
                cw.append(col_widt_list[j])
                cr.append(col_rule_list[j])
            col_alig_mtrx.append(ca)
            col_widt_mtrx.append(cw)
            col_rule_mtrx.append(cr)
        # ALIGNMENT
        for i in range(len(tab)):
            for j in range(len(tab[i])):
                cell = tab[i][j]
                if re.match('^\\s*:\\s', cell) and \
                   re.match(NOT_ESCAPED + '\\s:\\s*$', cell):
                    col_alig_mtrx[i][j] = WD_TABLE_ALIGNMENT.CENTER
                    tab[i][j] = re.sub('^\\s*:\\s(.*)\\s:\\s*$', '\\1', cell)
                elif re.match('^\\s*:\\s', cell):
                    col_alig_mtrx[i][j] = WD_TABLE_ALIGNMENT.LEFT
                    tab[i][j] = re.sub('^\\s*:\\s(.*)$', '\\1', cell)
                elif re.match(NOT_ESCAPED + '\\s:\\s*$', cell):
                    col_alig_mtrx[i][j] = WD_TABLE_ALIGNMENT.RIGHT
                    tab[i][j] = re.sub('^(.*)\\s:\\s*$', '\\1', cell)
                elif conf_line_place > 0 and i < conf_line_place:
                    col_alig_mtrx[i][j] = WD_TABLE_ALIGNMENT.CENTER
        return tab, col_alig_mtrx, col_widt_mtrx, col_rule_mtrx

    @staticmethod
    def __get_row_data(tab, conf_line_place,
                       row_alig_list, row_heig_list, row_rule_list):
        row_alig_mtrx, row_heig_mtrx, row_rule_mtrx = [], [], []
        for i in range(len(tab)):
            ra, rh, rr = [], [], []
            for j in range(len(tab[i])):
                ra.append(row_alig_list[i])
                rh.append(row_heig_list[i])
                rr.append(row_rule_list[i])
            row_alig_mtrx.append(ra)
            row_heig_mtrx.append(rh)
            row_rule_mtrx.append(rr)
        return tab, row_alig_mtrx, row_heig_mtrx, row_rule_mtrx


class ParagraphImage(Paragraph):

    """A class to handle image paragraph"""

    paragraph_class = 'image'
    res_feature = '^(?:\\s*' + RES_IMAGE + '\\s*)+$'

    def write_paragraph(self, ms_doc):
        ttw = self.text_to_write
        ttw = re.sub('\\s*(' + RES_IMAGE + ')\\s*', '\\1\n', ttw)
        ttw = re.sub('\n+', '\n', ttw)
        ttw = re.sub('^\n+', '', ttw)
        ttw = re.sub('\n+$', '', ttw)
        text_width = PAPER_WIDTH[Form.paper_size] \
            - Form.left_margin - Form.right_margin
        text_height = PAPER_HEIGHT[Form.paper_size] \
            - Form.top_margin - Form.bottom_margin
        res = '^(.*):(' + RES_NUMBER + ')?(?:x(' + RES_NUMBER + ')?)?$'
        for text in ttw.split('\n'):
            alte = re.sub(RES_IMAGE, '\\1', text)
            path = re.sub(RES_IMAGE, '\\2', text)
            # CAPTION
            capt = ''
            if re.match('^(.*) "(.*)"$', path):
                capt = re.sub('^(.*) "(.*)"$', '\\2', path)
                path = re.sub('^(.*) "(.*)"$', '\\1', path)
            cm_w = 0
            cm_h = 0
            if re.match(res, alte):
                st_w = re.sub(res, '\\2', alte)
                if st_w != '':
                    cm_w = float(st_w)
                    if cm_w < 0:
                        cm_w = text_width * (-1 * cm_w)
                st_h = re.sub(res, '\\3', alte)
                if st_h != '':
                    cm_h = float(st_h)
                    if cm_h < 0:
                        cm_h = text_height * (-1 * cm_h)
                alte = re.sub(res, '\\1', alte)
            try:
                if cm_w > 0 and cm_h > 0:
                    ms_doc.add_picture(path, width=Cm(cm_w), height=Cm(cm_h))
                elif cm_w > 0:
                    ms_doc.add_picture(path, width=Cm(cm_w))
                elif cm_h > 0:
                    ms_doc.add_picture(path, height=Cm(cm_h))
                else:
                    ms_doc.add_picture(path)
                ms_doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # CAPTION
                if capt != '':
                    ms_run = ms_doc.paragraphs[-1].add_run('\n' + capt)
                    XML.set_font(ms_run, self.chars_state.mincho_font)
                    ms_run.font.size = Pt(self.chars_state.font_size)
            except BaseException:
                e = ms_doc.paragraphs[-1]._element
                e.getparent().remove(e)
                ms_par = self.__get_ms_par(ms_doc)
                ms_par.add_run(text)
                ms_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                msg = '※ 警告: ' \
                    + '画像「' + path + '」が読み込めません'
                # msg = 'warning: can\'t open "' + path + '"'
                r = '^.*! *\\[.*\\] *\\(' + path + '\\).*$'
                for ml in self.md_lines:
                    if re.match(r, ml.text):
                        if msg not in ml.warning_messages:
                            ml.append_warning_message(msg)
                            break
                else:
                    self.md_lines[0].append_warning_message(msg)


class ParagraphMath(Paragraph):

    """A class to handle math paragraph"""

    paragraph_class = 'math'
    res_feature = '^\\\\\\[(.*)\\\\\\]$'

    @classmethod
    def is_this_class(cls, full_text,
                      head_font_revisers=[], tail_font_revisers=[]):
        if re.match(cls.res_feature, full_text):
            if re.match('^\\\\\\[.*$', full_text):
                if re.match(NOT_ESCAPED + '\\\\\\]$', full_text):
                    tmp = re.sub(cls.res_feature, '\\1', full_text)
                    if not re.match(NOT_ESCAPED + '\\\\[\\[\\]].*$', tmp):
                        return True
        return False

    def _get_alignment(self):
        text = re.sub(self.res_feature, '\\1', self.full_text)
        alignment = 'center'
        if re.match('^:\\s.*\\s:$', text):
            alignment = 'center'
        elif re.match('^:\\s.*$', text):
            alignment = 'left'
        elif re.match('^.*\\s:$', text):
            alignment = 'right'
        # self.alignment = alignment
        return alignment

    def _edit_data(self):
        md_lines = self.md_lines
        beg_has_removed = False
        end_has_removed = False
        for ml in md_lines:
            if not beg_has_removed:
                if re.match('^\\s*\\\\\\[(.*)$', ml.text):
                    ml.text = re.sub('^\\s*\\\\\\[(.*)$', '\\1', ml.text)
                    beg_has_removed = True
            if not end_has_removed:
                if re.match('^(.*)\\\\\\]\\s*$', ml.text):
                    ml.text = re.sub('^(.*)\\\\\\]\\s*$', '\\1', ml.text)
                    end_has_removed = True
            if self.alignment == 'left' or self.alignment == 'center':
                ml.text = re.sub('^:\\s', '', ml.text)
                # SPACE POLICY
                # ml.text = re.sub('^:\\s*', '', ml.text)
            if self.alignment == 'center' or self.alignment == 'right':
                ml.text = re.sub('\\s:$', '', ml.text)
                # SPACE POLICY
                # ml.text = re.sub('\\s*:$', '', ml.text)
            if ml.text == ':':
                ml.text = ''
            # COMMENT
            ml.text = re.sub(NOT_ESCAPED + '%.*$', '\\1', ml.text)
        # END OF LINE
        for j in range(len(md_lines)):
            if j == 0:
                continue
            i = j - 1
            if re.match('^.*[0-9A-Za-z]$', md_lines[i].text) and \
               re.match('^[0-9A-Za-z].*$', md_lines[j].text):
                md_lines[i].text += '\\,'

    def write_paragraph(self, ms_doc):
        ttw = self.text_to_write
        hfr = self.head_font_revisers
        tfr = self.tail_font_revisers
        chars_state = self.chars_state
        ms_par = ms_doc.add_paragraph()
        ms_par.style = 'makdo-m'
        ms_mpa = OxmlElement('m:oMathPara')
        self.__set_alignment(ms_par, ms_mpa)
        self.__set_length(ms_par)
        self.__set_font_revisers(hfr)
        chars = Math.write_chars(ms_par._p, chars_state, ttw)
        self.__set_font_revisers(tfr)

    def __set_alignment(self, ms_par, ms_mpa):
        ms_mpp = OxmlElement('m:oMathParaPr')
        ms_mpa.append(ms_mpp)
        oe = OxmlElement('m:jc')
        if self.alignment == 'left':
            ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT    # libre office
            oe.set(ns.qn('m:val'), 'left')                    # ms office
        elif self.alignment == 'right':
            ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT   # libre office
            oe.set(ns.qn('m:val'), 'right')                   # ms office
        else:
            ms_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # libre office
            oe.set(ns.qn('m:val'), 'center')                  # ms office
        ms_mpp.append(oe)

    def __set_length(self, ms_par):
        length_docx = self.length_docx
        f_size = Form.font_size
        ms_fmt = ms_par.paragraph_format
        ms_fmt.widow_control = False
        if length_docx['space before'] >= 0:
            pt = length_docx['space before'] * Form.line_spacing * f_size
            ms_fmt.space_before = Pt(pt)
        else:
            ms_fmt.space_before = Pt(0)
            msg = '※ 警告: ' \
                + '段落前の余白「v」の値が小さ過ぎます'
            # msg = 'warning: ' \
            #     + '"space before" is too small'
            self.md_lines[0].append_warning_message(msg)
        if length_docx['space after'] >= 0:
            pt = length_docx['space after'] * Form.line_spacing * f_size
            ms_fmt.space_after = Pt(pt)
        else:
            ms_fmt.space_after = Pt(0)
            msg = '※ 警告: ' \
                + '段落後の余白「V」の値が小さ過ぎます'
            # msg = 'warning: ' \
            #     + '"space after" is too small'
            self.md_lines[0].append_warning_message(msg)
        ms_fmt.first_line_indent = Pt(length_docx['first indent'] * f_size)
        ms_fmt.left_indent = Pt(length_docx['left indent'] * f_size)
        ms_fmt.right_indent = Pt(length_docx['right indent'] * f_size)
        # ms_fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        ls = Form.line_spacing * (1 + length_docx['line spacing'])
        if ls >= 1.0:
            ms_fmt.line_spacing = Pt(ls * f_size)
        else:
            ms_fmt.line_spacing = Pt(1.0 * f_size)
            msg = '※ 警告: ' \
                + '行間隔「X」の値が少な過ぎます'
            # msg = 'warning: ' \
            #     + 'too small line spacing'
            self.md_lines[0].append_warning_message(msg)
        ms_fmt.line_spacing = Pt(ls * f_size)

    def __set_font_revisers(self, font_revisers):
        chars_state = self.chars_state
        for fr in font_revisers:
            if False:
                pass
            # elif re.match('^@(' + RES_NUMBER + ')@$', fr):
            #     c_size = float(re.sub('^@(' + RES_NUMBER + ')@$', '\\1', fr))
            #     if c_size > 0:
            #         chars_state.font_scale = c_size / chars_state.font_size
            elif fr == '---':
                chars_state.font_scale = 0.6
            elif fr == '--':
                chars_state.font_scale = 0.8
            elif fr == '++':
                chars_state.font_scale = 1.2
            elif fr == '+++':
                chars_state.font_scale = 1.4
            elif fr == '**':
                chars_state.is_bold = not chars_state.is_bold
            elif fr == '~~':
                chars_state.has_strike = not chars_state.has_strike
            elif fr == '[|' or fr == '|]':
                if fr == '[|':
                    chars_state.has_frame = True
                elif fr == '|]':
                    chars_state.has_frame = False
            elif re.match('^_([\\$=\\.#\\-~\\+]{,4})_$', fr):
                sty = re.sub('^_([\\$=\\.#\\-~\\+]{,4})_$', '\\1', fr)
                if sty in UNDERLINE:
                    if chars_state.underline is None:
                        chars_state.underline = sty
                    elif chars_state.underline != sty:
                        chars_state.underline = sty
                    else:
                        chars_state.underline = None
            elif re.match('^\\^([0-9A-Za-z]{0,11})\\^$', fr):
                col = re.sub('^\\^([0-9A-Za-z]{0,11})\\^$', '\\1', fr)
                if col == '':
                    col = 'FFFFFF'
                elif re.match('^([0-9A-F])([0-9A-F])([0-9A-F])$', col):
                    col = re.sub('^([0-9A-F])([0-9A-F])([0-9A-F])$',
                                 '\\1\\1\\2\\2\\3\\3', col)
                elif col in FONT_COLOR:
                    col = FONT_COLOR[col]
                if re.match('^[0-9A-F]{6}$', col):
                    if chars_state.font_color is None:
                        chars_state.font_color = col
                    elif chars_state.font_color is col:
                        chars_state.font_color = col
                    else:
                        chars_state.font_color = None
            elif re.match('^_([0-9A-Za-z]{1,11})_$', fr):
                col = re.sub('^_([0-9A-Za-z]{1,11})_$', '\\1', fr)
                if col in HIGHLIGHT_COLOR:
                    hc = HIGHLIGHT_COLOR[col]
                    if chars_state.highlight_color is None:
                        chars_state.highlight_color = hc
                    elif chars_state.highlight_color != hc:
                        chars_state.highlight_color = hc
                    else:
                        chars_state.highlight_color = None


class ParagraphAlignment(Paragraph):

    """A class to handle alignment paragraph"""

    paragraph_class = 'alignment'
    res_feature = '^(?::|:\\s+.*|.*\\s+:)$'

    def _check_format(self):
        super()._check_format()
        md_lines = self.md_lines
        alignment = self.alignment
        for ml in md_lines:
            if alignment == 'left' or alignment == 'center':
                if re.match('^:\\s{2,}.*$', ml.text):
                    msg = '※ 警告: ' \
                        + 'テキストの最初に空白があります' \
                        + '（必要な場合は先頭に"\\"を入れてください）'
                    # msg = 'warning: ' \
                    #     + ' spaces at the beginning' \
                    #     + ' (if necessary, insert "\\")'
                    ml.append_warning_message(msg)
            if alignment == 'center' or alignment == 'right':
                if re.match('^.*\\s{2,}:$', ml.text):
                    msg = '※ 警告: ' \
                        + 'テキストの最後に空白があります'
                    # msg = 'warning: ' \
                    #     + ' spaces at the end'
                    ml.append_warning_message(msg)

    def _get_text_to_write(self):
        md_lines = self.md_lines
        alignment = self.alignment
        text_to_write = ''
        for ml in md_lines:
            if ml.text == '':
                continue
            if alignment == 'left':
                if re.match('^:\\s(.|\n)*$', ml.text):
                    ml.text = re.sub('^:\\s', '\n', ml.text)
                    # ml.text = re.sub('^:\\s+', '\n', ml.text)
            elif alignment == 'center':
                if re.match('^:\\s(.|\n)*$', ml.text):
                    ml.text = re.sub('^:\\s', '\n', ml.text)
                    # ml.text = re.sub('^:\\s+', '\n', ml.text)
                if re.match('^(.|\n)*\\s:$', ml.text):
                    ml.text = re.sub('\\s:$', '\n', ml.text)
                    # ml.text = re.sub('\\s+:$', '\n', ml.text)
            elif alignment == 'right':
                if re.match('^(.|\n)*\\s:$', ml.text):
                    ml.text = re.sub('\\s:$', '\n', ml.text)
                    # ml.text = re.sub('\\s+:$', '\n', ml.text)
            text_to_write += ml.text
        if alignment == 'center':
            text_to_write = text_to_write.replace('\n\n', '\n')
        if alignment == 'left' or alignment == 'center':
            text_to_write = re.sub('^\n', '', text_to_write)
        if alignment == 'center' or alignment == 'right':
            text_to_write = re.sub('\n$', '', text_to_write)
        return text_to_write


class ParagraphPreformatted(Paragraph):

    """A class to handle preformatted paragraph"""

    paragraph_class = 'preformatted'

    @classmethod
    def is_this_class(cls, full_text,
                      head_font_revisers, tail_font_revisers):
        if re.match('^```.*$', ''.join(head_font_revisers)) and \
           re.match('^.*```$', ''.join(tail_font_revisers)):
            return True
        return False

    @classmethod
    def _get_section_depths(cls, full_text):
        head_section_depth = Paragraph.bridge_tail_section_depth
        tail_section_depth = Paragraph.bridge_tail_section_depth
        # self.head_section_depth = head_section_depth
        # self.tail_section_depth = tail_section_depth
        return head_section_depth, tail_section_depth

    def _edit_data(self):
        self.head_font_revisers.pop(0)
        self.head_font_revisers.pop(0)
        self.head_font_revisers.pop(0)
        self.tail_font_revisers.pop(-1)
        self.tail_font_revisers.pop(-1)
        self.tail_font_revisers.pop(-1)
        self.md_lines[0].text = re.sub('\\s', '', self.md_lines[0].text)
        return

    def _get_text_to_write(self):
        md_lines = self.md_lines
        text_to_write = ''
        for i in range(len(md_lines)):
            if i == 0:
                if md_lines[i].text != '':
                    text_to_write += '[' + md_lines[i].text + ']\n'
            else:
                text_to_write += md_lines[i].text + '\n'
        text_to_write = re.sub('\n$', '', text_to_write)
        text_to_write = '`' + text_to_write + '`'
        return text_to_write


class ParagraphPagebreak(Paragraph):

    """A class to handle preformatted paragraph"""

    paragraph_class = 'pagebreak'
    res_feature = '^(?:<div style="break-.*: page;"></div>|<pgbr/?>|<Pgbr/?>)$'

    def __init__(self, raw_paragraph):
        super().__init__(raw_paragraph)
        self.is_attached_pagebreak = False

    def write_paragraph(self, ms_doc):
        is_attached_pagebreak = self.is_attached_pagebreak
        ttw = self.text_to_write
        if is_attached_pagebreak:
            ms_run = XML.add_tag(ms_doc.paragraphs[-1]._p, 'w:r')
            XML.add_tag(ms_run, 'w:br', {'w:type': 'page'})
        else:
            ms_doc.add_page_break()
        if re.match('<Pgbr/?>', ttw):
            ms_doc.add_section(WD_SECTION.NEW_PAGE)
            XML.add_tag(ms_doc.sections[-1]._sectPr,
                        'w:pgNumType', {'w:start': '1'})


class ParagraphHorizontalLine(Paragraph):

    """A class to handle Horizontalline paragraph"""

    paragraph_class = 'horizontalline'
    res_feature = '^(?:(?:\\s*\\-\\s*){4,}|(?:\\s*\\*\\s*){4,})$'

    def write_paragraph(self, ms_doc):
        length_revi = self.length_revi
        length_conf = self.length_conf
        length_clas = self.length_clas
        line_spacing = Form.line_spacing
        length_docx = self.length_docx
        f_size = Form.font_size
        ms_par = ms_doc.add_paragraph(style='makdo-h')
        length_docx \
            = {'space before': 0.0, 'space after': 0.0, 'line spacing': 0.0,
               'first indent': 0.0, 'left indent': 0.0, 'right indent': 0.0}
        for ln in length_docx:
            length_docx[ln] \
                = length_revi[ln] + length_conf[ln] + length_clas[ln]
        ms_fmt = ms_par.paragraph_format
        ms_fmt.line_spacing = 0
        ms_fmt.first_line_indent = Pt(length_docx['first indent'] * f_size)
        ms_fmt.left_indent = Pt(length_docx['left indent'] * f_size)
        ms_fmt.right_indent = Pt(length_docx['right indent'] * f_size)
        sb = (((line_spacing - 1) * 0.75 + 0.5) * f_size) \
            + (0.5 * length_docx['line spacing'] * line_spacing * f_size) \
            + length_docx['space before'] * line_spacing * f_size
        sa = (((line_spacing - 1) * 0.25 + 0.5) * f_size) \
            + (0.5 * length_docx['line spacing'] * line_spacing * f_size) \
            + length_docx['space after'] * line_spacing * f_size
        if sb < 0:
            msg = '※ 警告: ' \
                + '段落前の余白「v」の値が小さ過ぎます'
            # msg = 'warning: ' \
            #     + '"space before" is too small'
            self.md_lines[0].append_warning_message(msg)
            sb = 0
        if sa < 0:
            msg = '※ 警告: ' \
                + '段落前の余白「V」の値が小さ過ぎます'
            # msg = 'warning: ' \
            #     + '"space after" is too small'
            self.md_lines[0].append_warning_message(msg)
            sa = 0
        ms_fmt.space_before = Pt(sb)
        ms_fmt.space_after = Pt(sa)
        opts = {}
        opts['w:val'] = 'single'
        opts['w:sz'] = '6'
        # opts['w:space'] = '1'
        # opts['w:color'] = 'auto'
        ms_ppr = ms_par._p.get_or_add_pPr()
        ms_bdr = XML.add_tag(ms_ppr, 'w:pBdr', {})
        XML.add_tag(ms_bdr, 'w:bottom', opts)


class ParagraphBreakdown(Paragraph):

    """A class to handle breakdown paragraph"""

    paragraph_class = 'breakdown'
    res_feature = NOT_ESCAPED + '!.*!$'


class ParagraphRemarks(Paragraph):

    """A class to handle remarks paragraph"""

    paragraph_class = 'remarks'
    res_feature = '^""\\s+.*$'

    def write_paragraph(self, ms_doc):
        if Form.has_completed:
            return
        md_lines = self.md_lines
        ms_par = ms_doc.add_paragraph(style='makdo-r')
        for i, ml in enumerate(md_lines):
            if ml.text == '':
                continue
            text = '●' + re.sub('^""\\s+', '', ml.text)
            if i < len(md_lines) - 1:
                text += '\n'
            ms_run = ms_par.add_run(text)


class ParagraphSentence(Paragraph):

    """A class to handle sentence paragraph"""

    paragraph_class = 'sentence'

    @classmethod
    def _get_section_depths(cls, full_text):
        head_section_depth = Paragraph.bridge_tail_section_depth
        tail_section_depth = Paragraph.bridge_tail_section_depth
        # self.head_section_depth = head_section_depth
        # self.tail_section_depth = tail_section_depth
        return head_section_depth, tail_section_depth


class MdLine:

    """A class to handle markdown line"""

    is_in_comment = False

    def __init__(self, line_number, raw_text):
        self.line_number = line_number
        self.raw_text = raw_text
        self.spaced_text, self.comment = self._separate_comment(self.raw_text)
        self.beg_space, self.text, self.end_space \
            = self._separate_spaces(self.spaced_text)
        self.warning_messages = []

    @staticmethod
    def _separate_comment(raw_text):
        com_sep = ' / '
        spaced_text = ''
        comment = ''
        tmp = ''
        for i, c in enumerate(raw_text):
            tmp += c
            if not MdLine.is_in_comment:
                if re.match(NOT_ESCAPED + '<!--$', tmp):
                    tmp = re.sub('<!--$', '', tmp)
                    spaced_text += tmp
                    tmp = ''
                    MdLine.is_in_comment = True
            else:
                if re.match(NOT_ESCAPED + '-->$', tmp):
                    tmp = re.sub('-->$', '', tmp)
                    comment += tmp + com_sep
                    tmp = ''
                    MdLine.is_in_comment = False
            if MdLine.is_in_comment:
                continue
        else:
            if tmp != '':
                if MdLine.is_in_comment:
                    comment += tmp + com_sep
                else:
                    spaced_text += tmp
                tmp = ''
        comment = re.sub(com_sep + '$', '', comment)
        # self.spaced_text = spaced_text
        return spaced_text, comment

    @staticmethod
    def _separate_spaces(spaced_text):
        text = spaced_text
        res = '^(\\s+)(.*?)$'
        beg_space = ''
        if re.match(res, text):
            beg_space = re.sub(res, '\\1', text)
            text = re.sub(res, '\\2', text)
        res = '^(.*?)(\\s+)$'
        end_space = ''
        if re.match(res, text):
            end_space = re.sub(res, '\\2', text)
            text = re.sub(res, '\\1', text)
        if text == ':' and re.match('^( |\t|\u3000)$', end_space):
            text += end_space
            end_space = ''
        if re.match('^.*(  |\t|\u3000)$', end_space):
            text += '<br>'
            end_space = re.sub('(  |\t|\u3000)$', '', end_space)
        # self.beg_space = beg_space
        # self.text = text
        # self.end_space = end_space
        return beg_space, text, end_space

    def append_warning_message(self, warning_message):
        self.warning_messages.append(warning_message)

    def print_warning_messages(self):
        for wm in self.warning_messages:
            msg = wm + '\n' \
                + '  (line ' + str(self.line_number) + ') ' + self.raw_text
            sys.stderr.write(msg + '\n\n')


class Script:

    constant = {'pi': '3.141592653589793', 'e': '2.718281828459045'}

    def __init__(self, md_lines):
        self.md_lines = md_lines

    def execute(self):
        md_lines = self.md_lines
        for i in range(0, 10):
            md_lines = self.__execute_at_level(md_lines, i)
        # self.md_lines = md_lines
        return md_lines

    def __execute_at_level(self, md_lines, n):
        is_in_math = False
        is_in_script = False
        tmp_text = ''
        for ml in md_lines:
            old_text = ml.text
            new_text = ''
            for c in old_text:
                tmp_text += c
                if re.match(NOT_ESCAPED + '\\\\\\[', tmp_text):
                    is_in_math = True
                if re.match(NOT_ESCAPED + '\\\\\\]', tmp_text):
                    is_in_math = False
                if not is_in_script:
                    if self.__is_script_beginning(tmp_text, n):
                        if (not is_in_math) or \
                           (not re.match('^(.|\n)*{{', tmp_text)):
                            tmp_text = re.sub('{.?{$', '', tmp_text)
                            new_text += tmp_text
                            tmp_text = ''
                            is_in_script = True
                else:
                    if self.__is_script_end(tmp_text, n):
                        if (not is_in_math) or \
                           (not re.match('^(.|\n)*}}', tmp_text)):
                            tmp_text = re.sub('}.?}$', '', tmp_text)
                            new_text += self.__execute_script(tmp_text, ml)
                            tmp_text = ''
                            is_in_script = False
            else:
                if tmp_text != '':
                    if not is_in_script:
                        new_text += tmp_text
                        tmp_text = ''
                    else:
                        new_text += self.__execute_script(tmp_text, ml)
                        tmp_text = ''
            ml.text = new_text
        return md_lines

    @staticmethod
    def __is_script_beginning(text, n):
        if n == 1:
            if re.match(NOT_ESCAPED + '{1?{$', text):
                return True
        else:
            if re.match(NOT_ESCAPED + '{' + str(n) + '{$', text):
                return True
        return False

    @staticmethod
    def __is_script_end(text, n):
        if n == 1:
            if re.match(NOT_ESCAPED + '}1?}$', text):
                return True
        else:
            if re.match(NOT_ESCAPED + '}' + str(n) + '}$', text):
                return True
        return False

    def __execute_script(self, script, md_line):
        text_to_print = ''
        scr = script
        scr = re.sub('<br>$', '', scr)
        scr += ';'
        while scr != '':
            one = re.sub('^(.*?);(.*)$', '\\1', scr)
            scr = re.sub('^(.*?);(.*)$', '\\2', scr)
            if re.match('^\\s*(.*?)\\s*(/|\\*|%|-|\\+)=\\s*(.*?)\\s*$', one):
                # TRANSFORM ("x ?= y" -> "x = x ? y")
                one = re.sub('^\\s*(.*?)\\s*(/|\\*|%|-|\\+)=\\s*(.*?)\\s*$',
                             '\\1 = \\1 \\2 (\\3)', one)
            if one == '':
                pass
            elif re.match('^\\s*(.*?)\\s*=\\s*(.*?)\\s*$', one):
                # SUBSTITUTE ("x = y ? z")
                var = re.sub('^\\s*(.*?)\\s*=\\s*(.*?)\\s*$', '\\1', one)
                val = re.sub('^\\s*(.*?)\\s*=\\s*(.*?)\\s*$', '\\2', one)
                cal = self.__calc_value(val, md_line)
                self.constant[var] = cal
            elif re.match('^\\s*print\\s*\\((.*)\\)\\s*$', one):
                # PRINT ("print(x ? y)")
                val = re.sub('^\\s*print\\s*\\((.*)\\)\\s*$', '\\1', one)
                opt = ''
                if re.match('^(.*),\\s*["\'](.*)["\']$', val):
                    opt = re.sub('^(.*),\\s*["\'](.*)["\']$', '\\2', val)
                    val = re.sub('^(.*),\\s*["\'](.*)["\']$', '\\1', val)
                val = re.sub('^\\s*str\\s*\\((.*)\\)\\s*$', '\\1', val)
                cal = self.__calc_value(val, md_line)
                if re.match('^\\.[0-9]+$', cal):
                    cal = '0' + cal
                adj = cal
                # ADJUST
                if opt != '' and opt != '3' and opt != '4' and opt != '4s':
                    msg = '※ 警告: ' \
                        + '「' + opt + '」は不正なオプションです'
                    # msg = 'warning: ' \
                    #    '"' + opt + '" is a bad option'
                    md_line.append_warning_message(msg)
                if re.match('^([0-9]+\\.)?([0-9]+)$', cal) and (opt != ''):
                    if not re.match('^([0-9]+)\\.([0-9]+)$', cal):
                        inp = cal
                        dep = ''
                    else:
                        inp = re.sub('^([0-9]+)\\.([0-9]+)$', '\\1', cal)
                        dep = re.sub('^([0-9]+)\\.([0-9]+)$', '\\2', cal)
                    if opt == '3':
                        inp = '{:,}'.format(int(inp))
                    elif opt == '4' or opt == '4s':
                        if float(cal) >= 10000000000000000:
                            inp = re.sub('^(.*)(.{16})$', '\\1京\\2', inp)
                        if float(cal) >= 1000000000000:
                            inp = re.sub('^(.*)(.{12})$', '\\1兆\\2', inp)
                        if float(cal) >= 100000000:
                            inp = re.sub('^(.*)(.{8})$', '\\1億\\2', inp)
                        if float(cal) >= 10000:
                            inp = re.sub('^(.*)(.{4})$', '\\1万\\2', inp)
                    if opt == '4s':
                        inp = re.sub('0000$', '', inp)
                        inp = re.sub('0000万$', '', inp)
                        inp = re.sub('0000億$', '', inp)
                    if dep == '':
                        adj = inp
                    else:
                        adj = inp + '.' + dep
                # STRING
                adj = re.sub("^'((?:.|\n)*)'$", '\\1', adj)
                adj = re.sub('^"((?:.|\n)*)"$', '\\1', adj)
                text_to_print += adj
            else:
                msg = '※ 警告: ' \
                    + '「' + one + '」は不正なスクリプトです'
                # msg = 'warning: ' \
                #     + 'bad script'
                md_line.append_warning_message(msg)
        return text_to_print

    def __calc_value(self, value, md_line):
        val = value
        # NEW LINE
        val = re.sub('\\\\n', '\n', val)
        # NUMBER
        if re.match('^-?(?:[0-9,]*\\.)?[0-9,]+$', val):
            val = val.replace(',', '')
        # FUNCTIONS AND PARENTHESES
        new = ''
        tmp = ''
        dep = 0
        for c in val:
            if dep == 0 and c == '(':
                new += tmp
                tmp = ''
                dep = 1
            elif dep == 1 and c == ')':
                cal = self.__calc_value(tmp, md_line)
                if re.match('(^|[^a-zA-Z0-9])int\\s*$', new):
                    new = re.sub('int\\s*$', '', new) + str(int(float(cal)))
                else:
                    new += cal
                tmp = ''
                dep = 0
            else:
                tmp += c
                if c == '(':
                    dep += 1
                elif c == ')':
                    dep -= 1
        else:
            if tmp != '':
                new += tmp
        val = new
        # SUBSTITUTE VARIABLE
        tmp = ''
        par = ''
        for c in val + '\0':
            res = '(.*?)([_a-zA-Z][_a-zA-Z0-9]*)$'
            if par == '':
                if re.match(res, tmp) and re.match('^[^_a-zA-Z0-9]$', c):
                    var = re.sub(res, '\\2', tmp)
                    if var in self.constant:
                        tmp = re.sub(res, '\\g<1>' + self.constant[var], tmp)
                    else:
                        msg = '※ 警告: ' \
                            + '「' + var + '」という変数は未定義です'
                        # msg = 'warning: ' \
                        #     + 'variable "' + t + '" is undefined'
                        md_line.append_warning_message(msg)
            if re.match(NOT_ESCAPED + "'$", tmp + c):
                if par == '':
                    par = "'"
                elif par == "'":
                    par = ''
            if re.match(NOT_ESCAPED + '"$', tmp + c):
                if par == '':
                    par = '"'
                elif par == '"':
                    par = ''
            if c != '\0':
                tmp += c
        val = tmp
        # STRING AND STRING
        res = '^\\s*\'((?:.|\n)*?)\'\\s*\\+\\s*\'((?:.|\n)*)\'\\s*$'
        while re.match(res, val):
            val = re.sub(res, "'\\1\\2'", val)
        res = '^\\s*\'((?:.|\n)*?)\'\\s*\\+\\s*"((?:.|\n)*)"\\s*$'
        while re.match(res, val):
            val = re.sub(res, "'\\1\\2'", val)
        res = '^\\s*"((?:.|\n)*?)"\\s*\\+\\s*\'((?:.|\n)*)\'\\s*$'
        while re.match(res, val):
            val = re.sub(res, "'\\1\\2'", val)
        res = '^\\s*"((?:.|\n)*?)"\\s*\\+\\s*"((?:.|\n)*)"\\s*$'
        while re.match(res, val):
            val = re.sub(res, "'\\1\\2'", val)
        # BINARY OPERATE (x^y, x**y, x/y, x//y, x%y, x*y, x-y, x+y)
        val = self.__binary_operate('\\^|\\*\\*', val, md_line)
        val = self.__binary_operate('/|//|%|\\*', val, md_line)
        val = self.__binary_operate('\\-|\\+', val, md_line)
        # RETURN
        return val

    def __binary_operate(self, res_ope, val, md_line):
        res = '^((?:.*?\\s+)?)' + \
            '(-?(?:[0-9,]*\\.)?[0-9,]+|〓)' + \
            '\\s*(' + res_ope + ')\\s*' + \
            '(-?(?:[0-9,]*\\.)?[0-9,]+|〓)' + \
            '((?:\\s+.*)?)$'
        while re.match(res, val):
            s1 = re.sub(res, '\\2', val)
            s1 = s1.replace(',', '')
            op = re.sub(res, '\\3', val)
            s2 = re.sub(res, '\\4', val)
            s2 = s2.replace(',', '')
            if s1 == '〓' or s2 == '〓':
                return '〓'
            if ('.' not in s1) and ('.' not in s2):
                v1 = int(s1)
                v2 = int(s2)
            else:
                v1 = float(s1)
                v2 = float(s2)
            if False:
                pass
            elif op == '^' or op == '**':
                if v1 < 0 and type(v2) == float:
                    cal = '〓'
                    msg = '※ 警告: ' \
                        + '「' + val + '」は負数の小数乗です'
                    # msg = 'warning: ' \
                    #     + 'operation "' + val + '" is a decimal power'
                    md_line.append_warning_message(msg)
                elif v1 == 0 and v2 < 0:
                    cal = '〓'
                    msg = '※ 警告: ' \
                        + '「' + val + '」はゼロの負数乗です'
                    # msg = 'warning: ' \
                    #     + 'operation "' + val + '" is a negative power'
                    md_line.append_warning_message(msg)
                else:
                    cal = str(v1 ** v2)
            elif op == '/':
                if v2 == 0:
                    cal = '〓'
                    msg = '※ 警告: ' \
                        + '「' + val + '」はゼロで割っています'
                    # msg = 'warning: ' \
                    #     + 'operation "' + val + '" is division by zero'
                    md_line.append_warning_message(msg)
                else:
                    cal = str(v1 / v2)
            elif op == '//':
                if v2 == 0:
                    cal = '〓'
                    msg = '※ 警告: ' \
                        + '「' + val + '」はゼロで割っています'
                    # msg = 'warning: ' \
                    #     + 'operation "' + val + '" is division by zero'
                    md_line.append_warning_message(msg)
                else:
                    cal = str(v1 // v2)
            elif op == '%':
                if v2 == 0:
                    cal = '〓'
                    msg = '※ 警告: ' \
                        + '「' + val + '」はゼロで割っています'
                    # msg = 'warning: ' \
                    #     + 'operation "' + val + '" is modulo by zero'
                    md_line.append_warning_message(msg)
                else:
                    cal = str(v1 % v2)
            elif op == '*':
                cal = str(v1 * v2)
            elif op == '-':
                cal = str(v1 - v2)
            elif op == '+':
                cal = str(v1 + v2)
            val = re.sub(res, '\\g<1>' + cal + '\\g<5>', val)
        return val


class Md2Docx:

    """A class to make a MS Word file from a Markdown file"""

    def __init__(self, inputed_md_file, args=None):
        self.io = IO()
        io = self.io
        self.doc = Document()
        doc = self.doc
        self.frm = Form()
        frm = self.frm
        # READ MARKDOWN FILE
        io.set_md_file(inputed_md_file)
        formal_md_lines = io.read_md_file()
        doc.md_lines = doc.get_md_lines(formal_md_lines)
        # CONFIGURE
        frm.md_lines = doc.md_lines
        frm.args = args
        frm.configure()
        # UNFOLD
        doc.md_lines = Document().unfold(doc.md_lines)
        # EXECUTE SCRIPT
        doc.md_lines = Script(doc.md_lines).execute()
        # GET RAW PARAGRAPHS
        doc.raw_paragraphs = doc.get_raw_paragraphs(doc.md_lines)

    def make_docx(self):
        doc = self.doc
        frm = self.frm
        # GET PARAGRAPHS
        Paragraph.bridge_chars_state = CharsState()
        doc.paragraphs = doc.get_paragraphs(doc.raw_paragraphs)
        doc.paragraphs = doc.modify_paragraphs(doc.paragraphs)

    def save(self, inputed_docx_file):
        io = self.io
        doc = self.doc
        # MAKE DOCX
        self.make_docx()
        # WRITE DOCUMENT
        io.ms_doc = io.get_ms_doc()
        doc.write_property(io.ms_doc)
        doc.write_document(io.ms_doc)
        # SAVE MS WORD FILE
        io.set_docx_file(inputed_docx_file)
        io.save_docx_file()
        # PRINT WARNING MESSAGES
        doc.print_warning_messages()

    @staticmethod
    def set_document_title(value):
        return Form.set_document_title(value)

    @staticmethod
    def get_document_title():
        return Form.document_title

    @staticmethod
    def set_document_style(value):
        return Form.set_document_style(value)

    @staticmethod
    def get_document_style():
        return Form.document_style

    @staticmethod
    def set_paper_size(value):
        return Form.set_paper_size(value)

    @staticmethod
    def get_paper_size():
        return Form.paper_size

    @staticmethod
    def set_top_margin(value):
        return Form.set_top_margin(str(value))

    @staticmethod
    def get_top_margin():
        return Form.top_margin

    @staticmethod
    def set_bottom_margin(value):
        return Form.set_bottom_margin(str(value))

    @staticmethod
    def get_bottom_margin():
        return Form.bottom_margin

    @staticmethod
    def set_left_margin(value):
        return Form.set_left_margin(str(value))

    @staticmethod
    def get_left_margin():
        return Form.left_margin

    @staticmethod
    def set_right_margin(value):
        return Form.set_right_margin(str(value))

    @staticmethod
    def get_right_margin():
        return Form.right_margin

    @staticmethod
    def set_header_string(value):
        return Form.set_header_string(value)

    @staticmethod
    def get_header_string():
        return Form.header_string

    @staticmethod
    def set_page_number(value):
        return Form.set_page_number(value)

    @staticmethod
    def get_page_number():
        return Form.page_number

    @staticmethod
    def set_line_number(value):
        return Form.set_line_number(value)

    @staticmethod
    def get_line_number():
        return Form.line_number

    @staticmethod
    def set_mincho_font(value):
        return Form.set_mincho_font(value)

    @staticmethod
    def get_mincho_font():
        return Form.mincho_font

    @staticmethod
    def set_gothic_font(value):
        return Form.set_gothic_font(value)

    @staticmethod
    def get_gothic_font():
        return Form.gothic_font

    @staticmethod
    def set_ivs_font(value):
        return Form.set_ivs_font(value)

    @staticmethod
    def get_ivs_font():
        return Form.ivs_font

    @staticmethod
    def set_font_size(value):
        return Form.set_font_size(str(value))

    @staticmethod
    def get_font_size():
        return Form.font_size

    @staticmethod
    def set_line_spacing(value):
        return Form.set_line_spacing(str(value))

    @staticmethod
    def get_line_spacing():
        return Form.line_spacing

    @staticmethod
    def set_space_before(value):
        return Form.set_space_before(value)

    @staticmethod
    def get_space_before():
        return Form.space_before

    @staticmethod
    def set_space_after(value):
        return Form.set_space_after(value)

    @staticmethod
    def get_space_after():
        return Form.space_after

    @staticmethod
    def set_auto_space(value):
        return Form.set_auto_space(str(value))

    @staticmethod
    def get_auto_space():
        return Form.auto_space

    @staticmethod
    def set_version_number(value):
        return Form.set_version_number(value)

    @staticmethod
    def get_version_number():
        return Form.version_number

    @staticmethod
    def set_content_status(value):
        return Form.set_content_status(value)

    @staticmethod
    def get_content_status():
        return Form.content_status

    @staticmethod
    def set_has_completed(value):
        return Form.set_has_completed(str(value))

    @staticmethod
    def get_has_completed():
        return Form.has_completed


############################################################
# MAIN


def main():
    args = get_arguments()
    m2d = Md2Docx(args.md_file, args)
    m2d.save(args.docx_file)
    sys.exit(0)


if __name__ == '__main__':
    main()
